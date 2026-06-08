# -*- coding: utf-8 -*-
"""
MILA OFFICE — веб-версия (чат с агентами в браузере).

Запуск:
    cd "E:\\MILA GOLD\\mila-office"
    python webapp.py
Откроется http://127.0.0.1:5000 — по агенту на вкладку, как в макете.

Бэкенд переиспользует тех же 8 агентов (base.run_agent + их SYSTEM/TOOLS/handle).
История диалога хранится на сервере по агенту (локальное приложение, один пользователь).
"""
import sys
# Принудительно UTF-8 для консоли Windows — до любого вывода/логирования
# (иначе русский и ✓/эмодзи в логах и stdout превращаются в кракозябры).
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stdin.reconfigure(encoding="utf-8")
except Exception:
    pass

import base64
import importlib
import json
import logging
import mimetypes
import os
import re
import secrets
import threading
import uuid
import webbrowser
from collections import OrderedDict
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path
from urllib.parse import urlencode

import requests
from flask import Flask, request, jsonify, redirect, session, Response, abort, send_file

import base
import memory  # общая память офиса (профиль/фаза/события) — для дашборда
import error_monitor  # Централизованное логирование ошибок с Telegram alerts
import data_sanitizer  # Удаление конфиденциальных данных из логов

# ─── Логирование ─────────────────────────────────────────
# Полный traceback пишем в файл, клиенту отдаём безопасное сообщение.
_LOG_FILE = base.MILA_FOLDER / "logs" / "webapp.log"
_LOG_FILE.parent.mkdir(parents=True, exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    handlers=[
        logging.FileHandler(_LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(),
    ],
)


class _WerkzeugLocalNoiseFilter(logging.Filter):
    def filter(self, record):
        if record.name != "werkzeug":
            return True
        msg = record.getMessage()
        # Chrome/DevTools can probe https://localhost while the local app is
        # intentionally running plain HTTP. Werkzeug logs the TLS handshake bytes
        # as ERROR 400; it is transport noise, not an application failure.
        if (
            "code 400, message Bad request" in msg
            or "code 400, message Bad HTTP/0.9 request type" in msg
            or "\x16\x03\x01" in msg
            or "\\x16\\x03\\x01" in msg
        ):
            return False
        if "GET /.well-known/appspecific/com.chrome.devtools.json" in msg:
            return False
        return True


for _handler in logging.getLogger().handlers:
    _handler.addFilter(_WerkzeugLocalNoiseFilter())

logger = logging.getLogger("mila.webapp")

# ─── Реестр агентов ──────────────────────────────────────
# Марина (office/agent.py) — со своим run_agent(msg, history).
marina = importlib.import_module("agent")

_mods = {
    "victoria": importlib.import_module("victoria"),
    "alina":    importlib.import_module("alina"),
    "dima":     importlib.import_module("dima"),
    "tyoma":    importlib.import_module("tyoma"),
    "olya":     importlib.import_module("olya"),
    "vasya":    importlib.import_module("vasya"),
    "lera":     importlib.import_module("lera"),
    "rita":     importlib.import_module("rita"),
    "manager":  importlib.import_module("manager"),
    "producer": importlib.import_module("producer"),
}

_client = base.get_client()


def _office_responder(mod, key):
    # compose_system подмешивает улучшения Стаса (prompt_overrides/<key>.md) к SYSTEM.
    def respond(msg, history):
        return base.run_agent(_client, base.compose_system(key, mod.SYSTEM),
                              mod.TOOLS, mod.handle, msg, history, agent_key=key)
    return respond


# Marina теперь использует стандартный handle() как остальные агенты
# (не нужен специальный _marina_responder)


def _chips(mod):
    """Берёт быстрые команды агента → [{label, prompt}] (без /помощь, /выход)."""
    quick = getattr(mod, "QUICK", None) or getattr(mod, "QUICK_COMMANDS", {}) or {}
    out = []
    for k, v in quick.items():
        if k in ("/помощь", "/выход"):
            continue
        out.append({"label": k.lstrip("/").capitalize(), "prompt": v})
    return out[:10]  # у Стаса 8 команд — лимит 6 их резал; держим запас


AGENTS = {
    "marina": {
        "name": "Марина", "role": "Маркетолог", "emoji": "📣", "color": "#C4614A",
        "intro": "Привет! 👋 Я Марина. Давай сегодня: прочитаю комментарии и подготовлю ответы, "
                 "соберу аналитику за неделю, напишу Reels или карусель, спланирую стратегию. "
                 "Выбери внизу или скажи, что нужно.",
        "responder": _office_responder(marina, "marina"), "chips": _chips(marina),
    },
    "victoria": {
        "name": "Виктория", "role": "Редактор", "emoji": "✍️", "color": "#4A7A5E",
        "intro": "Я Виктория. Хочешь, я прочитаю последние посты и оценю их? Или пришлёшь текст "
                 "на редактуру — проверю голос, хук, CTA и дам финальную версию.",
        "responder": _office_responder(_mods["victoria"], "victoria"), "chips": _chips(_mods["victoria"]),
    },
    "alina": {
        "name": "Алина", "role": "Клиенты", "emoji": "👩", "color": "#2B7A8B",
        "intro": "Я Алина, менеджер клиентов. Помогу разобрать анкеты, определить "
                 "паттерн, подготовить тебя к сессии и вести истории клиенток.",
        "responder": _office_responder(_mods["alina"], "alina"), "chips": _chips(_mods["alina"]),
    },
    "dima": {
        "name": "Дима", "role": "Финансы", "emoji": "💰", "color": "#2C5F3A",
        "intro": "Я Дима, финансы. Посчитаю доход, сравню с целями, построю прогноз. "
                 "Для продаж с Gumroad нужен GUMROAD_ACCESS_TOKEN в .env.",
        "responder": _office_responder(_mods["dima"], "dima"), "chips": _chips(_mods["dima"]),
    },
    "tyoma": {
        "name": "Тёма", "role": "Telegram", "emoji": "💬", "color": "#2B5278",
        "intro": "Я Тёма, Telegram-менеджер. Напишу посты для канала и welcome-цепочку. "
                 "Для публикации нужен TELEGRAM_BOT_TOKEN в .env.",
        "responder": _office_responder(_mods["tyoma"], "tyoma"), "chips": _chips(_mods["tyoma"]),
    },
    "olya": {
        "name": "Оля", "role": "Тренды", "emoji": "🔍", "color": "#6B3FA0",
        "intro": "Я Оля, тренды. Найду вирусные темы, разберу конкурентов и дам "
                 "конкретные хуки для Reels.",
        "responder": _office_responder(_mods["olya"], "olya"), "chips": _chips(_mods["olya"]),
    },
    "vasya": {
        "name": "Вася", "role": "Планировщик", "emoji": "📅", "color": "#8B4513",
        "intro": "Я Вася, планировщик. Составлю расписание публикаций на неделю/месяц "
                 "и напомню, что нужно снять.",
        "responder": _office_responder(_mods["vasya"], "vasya"), "chips": _chips(_mods["vasya"]),
    },
    "lera": {
        "name": "Лера", "role": "Продажи", "emoji": "🎯", "color": "#A84F3C",
        "intro": "Я Лера, продажи. Хочешь, я напишу продающий текст для практикума? Или посмотрим "
                 "воронку, конверсию и какие акции запустить? Выбери внизу или скажи свою идею.",
        "responder": _office_responder(_mods["lera"], "lera"), "chips": _chips(_mods["lera"]),
    },
    "manager": {
        "name": "Стас", "role": "Офис-менеджер", "emoji": "🗂️", "color": "#5C6B7A",
        "intro": "Я Стас, офис-менеджер и бизнес-стратег. Делаю ревью работы агентов, считаю "
                 "метрики офиса, веду задачи — и мыслю воронкой, юнит-экономикой и ставками "
                 "роста. Дам приоритеты, KPI и измеримый план.\n\nНачни с быстрого действия "
                 "внизу — «Обзор за 24ч» или «Стратегия».",
        "responder": _office_responder(_mods["manager"], "manager"), "chips": _chips(_mods["manager"]),
    },
    "producer": {
        "name": "Кирилл", "role": "Продюсер", "emoji": "🎬", "color": "#C8962C",
        "intro": "Я Кирилл, продюсер эксперта. Отвечаю за бизнес самой Людмилы: продуктовую "
                 "линейку, запуски, позиционирование и рост дохода — не «пост на завтра», а "
                 "«как вырасти за квартал».\n\nНачни с быстрого действия — «Линейка» или «Запуск».",
        "responder": _office_responder(_mods["producer"], "producer"), "chips": _chips(_mods["producer"]),
    },
    "rita": {
        "name": "Рита", "role": "Продуктовый архитектор", "emoji": "📚", "color": "#9C5BA8",
        "intro": "Я Рита, архитектор цифровых продуктов. Превращаю идею в ясную структуру "
                 "(главы, упражнения, поток) — основу для PDF-воркбука, который напишет Марина "
                 "и отредактирует Виктория.\n\nНачни с «Воркбук» или опиши идею продукта.",
        "responder": _office_responder(_mods["rita"], "rita"), "chips": _chips(_mods["rita"]),
    },
}

# Состояние чата хранится по браузерной сессии, чтобы вкладки/пользователи
# не читали и не сбрасывали историю друг друга.
_histories = {}
_locks = {k: threading.Lock() for k in AGENTS}  # свой замок на агента → разные агенты параллельны
_jobs_lock = threading.Lock()

# Логирование истории в файлы для персистентности
_SESSION_LOGS_DIR = base.MILA_FOLDER / "logs" / "sessions"
_SESSION_LOGS_DIR.mkdir(parents=True, exist_ok=True)

# Фоновые задачи: агент может думать десятки секунд. Не держим HTTP-запрос
# открытым — кладём вызов в пул, фронтенд опрашивает /api/result.
_pool = ThreadPoolExecutor(max_workers=4)
_uploads = OrderedDict() # upload_id -> {"sid": ..., "name": ..., "text": ...}
_jobs = OrderedDict()    # job_id -> {"status": ...}; ограничен по размеру, читается один раз
MAX_JOBS = 200           # бэкстоп против утечки незабранных задач
MAX_HISTORY_MSGS = 40    # ~20 последних реплик user/assistant на агента (защита памяти/токенов)

MAX_UPLOADS = 80
MAX_UPLOAD_BYTES = 12 * 1024 * 1024
MAX_EXTRACTED_CHARS = 30000
_DOCUMENTS_DIR = base.MILA_FOLDER / "reports" / "documents"
_DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)


def _safe_doc_id(doc_id: str) -> str:
    doc_id = (doc_id or "").strip()
    if not re.fullmatch(r"[A-Za-z0-9_-]{4,80}", doc_id):
        abort(404)
    return doc_id


def _doc_path(doc_id: str) -> Path:
    return _DOCUMENTS_DIR / f"{_safe_doc_id(doc_id)}.json"


def _doc_export_path(doc_id: str) -> Path:
    return _DOCUMENTS_DIR / f"{_safe_doc_id(doc_id)}.txt"


def _plain_msg_text(item) -> str:
    if isinstance(item, dict):
        return str(item.get("content") or item.get("text") or item.get("message") or "")
    return str(item or "")


def _load_document_record(doc_id: str):
    path = _doc_path(doc_id)
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, ValueError):
            logger.exception("document record is unreadable: %s", path)
    return _document_from_history(doc_id)


def _document_from_history(doc_id: str):
    doc_id = _safe_doc_id(doc_id)
    hits = []
    for sid, histories in list(_histories.items()):
        for agent_key, hist in list((histories or {}).items()):
            for idx, item in enumerate(hist or []):
                text = _plain_msg_text(item)
                if doc_id in text:
                    prev = _plain_msg_text(hist[idx - 1]) if idx else ""
                    hits.append({
                        "agent": agent_key,
                        "timestamp": datetime.now().isoformat(timespec="seconds"),
                        "verdict": "done" if "VERDICT: done" in text or "готов" in text.lower() else "ready_next",
                        "input": prev[-8000:],
                        "output": text[-20000:],
                    })
    if hits:
        return {
            "id": doc_id,
            "file_name": f"mila-document-{doc_id}.txt",
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "status": "ready",
            "original_content": hits[0].get("input", ""),
            "stages": hits,
            "feedback_chain": [],
        }
    return {
        "id": doc_id,
        "file_name": f"mila-document-{doc_id}.txt",
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "status": "missing_source",
        "original_content": "",
        "stages": [{
            "agent": "office",
            "timestamp": datetime.now().isoformat(timespec="seconds"),
            "verdict": "needs_revision",
            "input": "",
            "output": (
                f"Документ {doc_id} был упомянут в чате, но исходный файл не был сохранен "
                "в реестре документов. Отправьте файл агенту еще раз или скачайте текст "
                "из сообщения через кнопку в чате."
            ),
        }],
        "feedback_chain": [],
    }


def _save_document_record(doc_id: str, agent_key: str, user_msg: str, reply: str, attachment=None):
    doc_id = _safe_doc_id(doc_id)
    now = datetime.now().isoformat(timespec="seconds")
    doc = _load_document_record(doc_id)
    if not doc or doc.get("status") == "missing_source":
        doc = {
            "id": doc_id,
            "file_name": f"mila-document-{doc_id}.txt",
            "created_at": now,
            "status": "ready",
            "original_content": (attachment or {}).get("text") or user_msg or "",
            "stages": [],
            "feedback_chain": [],
        }
    doc["status"] = "ready"
    # Если агент вернул готовый документ в маркерах — это и есть новый файл на
    # скачивание; в транскрипт-этап кладём только комментарий (без полотна текста).
    final = _extract_final_document(reply)
    stage_output = _strip_doc_block(reply) if final else (reply or "")
    if final:
        doc["final_content"] = final
        doc["final_by"] = agent_key
        doc["final_at"] = now
    doc.setdefault("stages", []).append({
        "agent": agent_key,
        "timestamp": now,
        "verdict": "done" if "VERDICT: done" in reply or "готов" in reply.lower() else "ready_next",
        "input": user_msg or "",
        "output": stage_output or "",
    })
    _doc_path(doc_id).write_text(json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")
    _doc_export_path(doc_id).write_text(_document_download_text(doc), encoding="utf-8")
    return doc


def _extract_doc_ids(text: str):
    text = text or ""
    ids = set(re.findall(r"\[doc_id:([A-Za-z0-9_-]{4,80})\]", text))
    ids.update(re.findall(r"/api/document/([A-Za-z0-9_-]{4,80})", text))
    return sorted(ids)


# Маркеры, которыми агент оборачивает ГОТОВЫЙ (уже исправленный) документ — чтобы
# приложение отдало чистый файл, а не транскрипт обсуждения с комментариями.
# Конвенция задаётся всем агентам в base.compose_system.
_DOC_BLOCK_RE = re.compile(r"\[ДОКУМЕНТ\](.*?)\[/ДОКУМЕНТ\]", re.S | re.I)


def _extract_final_document(reply: str):
    """Чистый исправленный текст документа из ответа агента, либо None."""
    if not reply:
        return None
    m = _DOC_BLOCK_RE.search(reply)
    if not m:
        return None
    return m.group(1).strip() or None


def _strip_doc_block(reply: str) -> str:
    """Убирает блок [ДОКУМЕНТ]…[/ДОКУМЕНТ] из видимого в чате текста: сам документ
    уезжает в скачиваемый файл, в чате остаётся только комментарий агента."""
    if not reply:
        return reply
    return _DOC_BLOCK_RE.sub("", reply).strip()


def _document_download_text(doc: dict) -> str:
    """Текст для скачивания. Если агент пометил готовый документ — отдаём ЧИСТЫЙ
    исправленный документ; иначе (старый сценарий) — полный транскрипт этапов."""
    final = (doc or {}).get("final_content")
    if final:
        return final.strip() + "\n"
    return _document_to_text(doc)


def _document_to_text(doc: dict) -> str:
    lines = [
        f"MILA OFFICE DOCUMENT: {doc.get('id', '')}",
        f"Статус: {doc.get('status', '')}",
        f"Создан: {doc.get('created_at', '')}",
        "",
    ]
    original = doc.get("original_content") or ""
    if original:
        lines += ["ИСХОДНЫЙ МАТЕРИАЛ", original, ""]
    for idx, stage in enumerate(doc.get("stages") or [], 1):
        lines += [
            f"ЭТАП {idx}: {stage.get('agent', '')} / {stage.get('verdict', '')}",
            f"Время: {stage.get('timestamp', '')}",
        ]
        if stage.get("input"):
            lines += ["", "Вход:", stage.get("input", "")]
        if stage.get("output"):
            lines += ["", "Результат:", stage.get("output", "")]
        lines.append("")
    return "\n".join(lines).strip() + "\n"

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_BYTES
# Нужен для session (CSRF + OAuth-state). Ключ ДОЛЖЕН переживать перезапуск, иначе
# каждый рестарт инвалидирует session-cookie открытой вкладки → старый CSRF-токен не
# сходится с новым → POST /api/chat падает 403 (фронт ловит как «Сеть недоступна»).
# Берём из .env, иначе генерим один раз и кладём в .secret_key рядом с webapp.
def _persistent_secret():
    env = os.getenv("FLASK_SECRET_KEY")
    if env:
        return env
    f = Path(__file__).resolve().parent / ".secret_key"
    try:
        if f.exists():
            return f.read_text(encoding="utf-8").strip()
        key = secrets.token_hex(32)
        f.write_text(key, encoding="utf-8")
        return key
    except OSError:
        return secrets.token_hex(32)  # фолбэк: хотя бы не падаем

app.secret_key = _persistent_secret()
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=os.getenv("MILA_HTTPS", "").lower() in ("1", "true", "yes"),
)


@app.after_request
def _clear_hsts(resp):
    """Когда приложение на HTTP (MILA_HTTPS=0), явно гасим HSTS, чтобы браузер,
    запомнивший https с прошлого MILA_HTTPS=1, не упирался в TLS-на-HTTP-порт
    (поток '400 Bad request version', страница/агенты не грузятся). max-age=0
    сбрасывает закэшированную политику. На самом https это не отправляется."""
    if not _HTTPS:
        resp.headers["Strict-Transport-Security"] = "max-age=0"
    return resp


def _session_id():
    sid = session.get("sid")
    if not sid:
        sid = secrets.token_urlsafe(24)
        session["sid"] = sid
    return sid


def _csrf_token():
    token = session.get("csrf")
    if not token:
        token = secrets.token_urlsafe(32)
        session["csrf"] = token
    return token


def _same_origin_request():
    origin = request.headers.get("Origin") or request.headers.get("Referer")
    if not origin:
        return True
    host_url = request.host_url.rstrip("/")
    return origin.startswith(host_url)


@app.before_request
def _protect_post_routes():
    if request.method != "POST":
        return
    if not _same_origin_request():
        abort(403)
    if not secrets.compare_digest(request.headers.get("X-CSRF-Token", ""), _csrf_token()):
        abort(403)


def _save_session_message(sid, agent_key, role, content):
    """Сохранить сообщение в JSONL файл для персистентности."""
    import datetime as _dt
    log_dir = _SESSION_LOGS_DIR / sid
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"{agent_key}.jsonl"
    try:
        with open(log_file, "a", encoding="utf-8") as f:
            json.dump({
                "timestamp": _dt.datetime.utcnow().isoformat(),
                "role": role,
                "content": content
            }, f, ensure_ascii=False)
            f.write("\n")
    except Exception as e:
        logger.warning(f"Failed to save session message: {e}")


def _load_session_history(sid, agent_key):
    """Загрузить историю агента из JSONL файла."""
    log_file = _SESSION_LOGS_DIR / sid / f"{agent_key}.jsonl"
    if not log_file.exists():
        return []
    history = []
    try:
        with open(log_file, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    msg = json.loads(line)
                    # Восстанавливаем исходный формат истории
                    history.append({
                        "role": msg.get("role", "user"),
                        "content": msg.get("content", "")
                    })
                except json.JSONDecodeError:
                    pass
    except Exception as e:
        logger.warning(f"Failed to load session history: {e}")
        return []
    return history


def _session_histories(sid):
    """Получить историю по сессии, загружая из файлов если нужно."""
    if sid not in _histories:
        _histories[sid] = {}
        for k in AGENTS:
            _histories[sid][k] = _load_session_history(sid, k)
    return _histories[sid]


def _trim(history):
    """Обрезает историю до последних MAX_HISTORY_MSGS сообщений.

    base.run_agent кладёт в history только пары user/assistant (tool-блоки живут
    в локальном messages), поэтому срез по границе сохраняет порядок и не рвёт пары.
    """
    return history[-MAX_HISTORY_MSGS:] if len(history) > MAX_HISTORY_MSGS else history


def _safe_upload_name(name: str) -> str:
    name = os.path.basename(name or "upload")
    name = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip(" ._")
    return name[:120] or "upload"


def _clip_text(text: str) -> str:
    text = (text or "").replace("\x00", "").strip()
    return text[:MAX_EXTRACTED_CHARS]


def _decode_text_file(raw: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1251", "cp1252"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _looks_garbled(text: str) -> bool:
    """Эвристика «битый текстовый слой» (типично для PDF из GAMMA с subset-шрифтами
    без корректного ToUnicode): в извлечённом тексте кириллические слова перемешаны
    с цифрами/латиницей ВНУТРИ слова — «Поч<G», «6ы5иD4N», «A989?O@». В нормальном
    русском тексте такого почти не бывает (исключения вроде «30-дневная» редки)."""
    tokens = re.findall(r"\S+", text or "")
    cyr_tokens = mixed = 0
    for t in tokens:
        if not re.search(r"[а-яёА-ЯЁ]", t):
            continue  # чисто латинские/числовые токены (@liudmyla, 1–7) не считаем
        cyr_tokens += 1
        if re.search(r"[A-Za-z0-9]", t):  # цифра/латиница внутри слова с кириллицей
            mixed += 1
    if cyr_tokens < 10:
        return False  # слишком мало данных, чтобы судить
    return (mixed / cyr_tokens) > 0.30


def _ocr_one_page(args) -> str:
    """Распознать одну отрендеренную страницу через Gemini. args=(idx, png_bytes,
    url, key, prompt). Возвращает текст или пометку об ошибке. С ретраями на 429/5xx."""
    import base64, time as _t
    idx, png, url, key, prompt = args
    payload = {
        "contents": [{"role": "user", "parts": [
            {"text": prompt},
            {"inline_data": {"mime_type": "image/png",
                             "data": base64.b64encode(png).decode()}},
        ]}],
        "generationConfig": {"maxOutputTokens": 2048,
                             "thinkingConfig": {"thinkingBudget": 0}},
    }
    for attempt in range(3):
        try:
            r = requests.post(url, params={"key": key}, json=payload, timeout=90)
            if r.status_code in (429, 500, 502, 503, 504):
                _t.sleep(1.5 * (attempt + 1))
                continue
            if r.status_code != 200:
                return f"[стр. {idx+1}: OCR ошибка {r.status_code}]"
            parts = r.json().get("candidates", [{}])[0].get("content", {}).get("parts", [])
            return ("\n".join(p.get("text", "") for p in parts if p.get("text")).strip()
                    or f"[стр. {idx+1}: пусто]")
        except Exception as e:
            if attempt == 2:
                return f"[стр. {idx+1}: OCR сбой {type(e).__name__}]"
            _t.sleep(1.5 * (attempt + 1))
    return f"[стр. {idx+1}: OCR не удался]"


def _pdf_ocr_via_gemini(raw: bytes) -> tuple[str, str]:
    """Фолбэк для PDF с битым текстовым слоем (GAMMA-экспорт и т.п.): визуальный
    слой страниц чистый, поэтому рендерим страницы в PNG (PyMuPDF) и распознаём
    текст через Gemini (мультимодальный), параллельно. Возвращает (text, note);
    пусто — если нет PyMuPDF / GEMINI_KEY или распознать не удалось."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        return "", "OCR недоступен: не установлен PyMuPDF (pip install pymupdf)."
    key = getattr(base, "GEMINI_KEY", "")
    if not key:
        return "", "OCR недоступен: не задан GEMINI_KEY для распознавания страниц."
    max_pages = int(os.getenv("MILA_PDF_OCR_PAGES", "15"))
    model = getattr(base, "GEMINI_MODEL", "gemini-2.5-flash")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    prompt = ("Перед тобой страница PDF — рабочая тетрадь по психологии на русском. "
              "Извлеки ВЕСЬ читаемый текст со страницы дословно: заголовки, списки, "
              "вопросы, подписи. Сохрани порядок. Ничего не добавляй от себя и не "
              "комментируй — только текст страницы.")
    try:
        doc = fitz.open(stream=raw, filetype="pdf")
    except Exception as e:
        return "", f"OCR: не удалось открыть PDF ({type(e).__name__})."
    total = len(doc)
    n = min(total, max_pages)
    # Рендер делаем последовательно (быстро, CPU), сетевые вызовы — параллельно.
    jobs = []
    for i in range(n):
        try:
            png = doc[i].get_pixmap(dpi=130).tobytes("png")
            jobs.append((i, png, url, key, prompt))
        except Exception as e:
            jobs.append((i, b"", url, key, prompt))
    doc.close()
    from concurrent.futures import ThreadPoolExecutor
    results = [""] * n
    with ThreadPoolExecutor(max_workers=4) as ex:
        for i, txt in zip(range(n), ex.map(_ocr_one_page, jobs)):
            results[i] = txt
    text = "\n\n".join(results).strip()
    if not text:
        return "", "OCR не дал результата (страницы не распознались)."
    if total > n:
        text += f"\n\n[…распознаны первые {n} из {total} страниц; лимит MILA_PDF_OCR_PAGES]"
    note = ("Текст распознан со страниц (OCR через Gemini) — у PDF повреждён "
            "текстовый слой. Возможны мелкие неточности.")
    return text, note


def _extract_pdf_text(raw: bytes) -> tuple[str, str]:
    import io
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader
        except Exception:
            return "", "PDF загружен, но библиотека pypdf/PyPDF2 не установлена. Установи pypdf, чтобы читать текст PDF."
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages[:30]:
            pages.append(page.extract_text() or "")
        text = "\n\n".join(pages).strip()
        garbled = _looks_garbled(text)
        # Нет текста ИЛИ битый слой → пробуем OCR по картинкам страниц (если включён).
        if (not text or garbled) and os.getenv("MILA_PDF_OCR", "1").lower() in ("1", "true", "yes"):
            ocr_text, ocr_note = _pdf_ocr_via_gemini(raw)
            if ocr_text and not _looks_garbled(ocr_text):
                return ocr_text, ocr_note
        if not text:
            return "", "PDF загружен, но текст не найден. Возможно, это скан; загрузи картинку или установи OCR (PyMuPDF + GEMINI_KEY)."
        if garbled:
            # OCR не вышел/выключен — честно объясняем и даём образец каши.
            note = (
                "⚠️ Текстовый слой этого PDF повреждён (похоже на экспорт из GAMMA с "
                "subset-шрифтами без ToUnicode): при извлечении получается нечитаемая "
                "кодировочная каша, оценивать её бессмысленно.\n"
                "Что делать: пришли воркбук как Markdown/.docx или текстом — тогда "
                "редактор сможет оценить содержание. Сам PDF визуально в порядке, "
                "ломается только машинно извлекаемый текст.\n\n"
                "Образец того, что извлеклось (для наглядности):\n"
                + text[:400]
            )
            return note, ("Битый текстовый слой PDF — прислан образец вместо содержимого. "
                          "Загрузи Markdown/.docx/текст для оценки.")
        return text, ""
    except Exception as e:
        return "", f"Не удалось прочитать PDF: {type(e).__name__}"


def _extract_docx_text(raw: bytes) -> tuple[str, str]:
    import io
    try:
        from docx import Document
    except Exception:
        return "", "DOCX uploaded, but python-docx is not installed. Install python-docx to read Word files."
    try:
        doc = Document(io.BytesIO(raw))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if not text:
            return "", "DOCX uploaded, but no text was found in the document."
        return text, ""
    except Exception as e:
        return "", f"Could not read DOCX: {type(e).__name__}"


def _image_feedback_prompt() -> str:
    return (
        "Ты читаешь скриншот для последующего фидбека от агента MILA Office.\n"
        "1. Сначала сделай OCR: перепиши весь видимый текст максимально полно, включая сообщения чата, "
        "ошибки консоли, кнопки, статусы и короткие фразы пользователя.\n"
        "2. Отдельно выпиши, какая проблема видна на скриншоте и где она находится.\n"
        "3. Если это интерфейс чата, отметь, что пользователь хотел сделать и что ответил агент.\n"
        "4. Не отвечай, что не можешь прочитать изображение, если текст хотя бы частично виден. "
        "Лучше перепиши видимые фрагменты и явно пометь неразборчивые места как [неразборчиво]."
    )


def _describe_image_with_claude(raw: bytes, mime: str) -> tuple[str, str]:
    if not (base.ANTHROPIC_KEY or base.ANTHROPIC_AUTH_TOKEN):
        return "", "Claude Vision недоступен: ANTHROPIC_API_KEY не настроен."
    client = _client or base.get_client()
    if client is None:
        return "", "Claude Vision недоступен: клиент Anthropic не настроен."
    try:
        resp = client.messages.create(
            model=base.MODEL,
            max_tokens=2200,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": _image_feedback_prompt()},
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": mime,
                            "data": base64.b64encode(raw).decode("ascii"),
                        },
                    },
                ],
            }],
        )
        text = "\n".join(
            getattr(part, "text", "") for part in getattr(resp, "content", [])
            if getattr(part, "type", "") == "text" and getattr(part, "text", "")
        ).strip()
        return text, "" if text else "Claude Vision не вернул текстовое описание изображения."
    except Exception as e:
        return "", f"Claude Vision не смог прочитать изображение: {type(e).__name__}"


def _describe_image_with_gemini(raw: bytes, mime: str) -> tuple[str, str]:
    if not getattr(base, "GEMINI_KEY", ""):
        return _describe_image_with_claude(raw, mime)
    payload = {
        "contents": [{
            "role": "user",
            "parts": [
                {"text": _image_feedback_prompt()},
                {"inline_data": {"mime_type": mime, "data": base64.b64encode(raw).decode("ascii")}},
            ],
        }],
        "generationConfig": {"maxOutputTokens": 2200},
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{base.GEMINI_MODEL}:generateContent"
    try:
        r = requests.post(url, params={"key": base.GEMINI_KEY}, json=payload, timeout=60)
        if r.status_code != 200:
            text, warning = _describe_image_with_claude(raw, mime)
            if text:
                return text, f"Gemini Vision вернул HTTP {r.status_code}; использован Claude Vision."
            return "", f"Картинка загружена, но Gemini Vision вернул HTTP {r.status_code}. {warning}"
        parts = r.json().get("candidates", [{}])[0].get("content", {}).get("parts", [])
        text = "\n".join(p.get("text", "") for p in parts if p.get("text")).strip()
        if text:
            return text, ""
        return _describe_image_with_claude(raw, mime)
    except Exception as e:
        text, warning = _describe_image_with_claude(raw, mime)
        if text:
            return text, f"Gemini Vision не сработал ({type(e).__name__}); использован Claude Vision."
        return "", f"Картинка загружена, но описание не удалось получить: {type(e).__name__}. {warning}"


def _extract_upload(filename: str, raw: bytes, mime: str) -> tuple[str, str]:
    suffix = Path(filename).suffix.lower()
    mime = (mime or mimetypes.guess_type(filename)[0] or "").lower()
    if suffix in {".txt", ".md", ".csv", ".json", ".html", ".htm", ".rtf"} or mime.startswith("text/"):
        return _decode_text_file(raw), ""
    if suffix == ".pdf" or mime == "application/pdf":
        return _extract_pdf_text(raw)
    if suffix == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx_text(raw)
    if mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}:
        return _describe_image_with_gemini(raw, mime or "image/png")
    return "", "Поддерживаются текстовые файлы, PDF и изображения."


def _set_upload(upload_id: str, value: dict):
    _uploads[upload_id] = value
    _uploads.move_to_end(upload_id)
    while len(_uploads) > MAX_UPLOADS:
        _uploads.popitem(last=False)


def _set_job(job_id, value):
    """Кладёт результат задачи и держит размер _jobs под MAX_JOBS (FIFO-вытеснение)."""
    with _jobs_lock:
        _jobs[job_id] = value
        _jobs.move_to_end(job_id)
        while len(_jobs) > MAX_JOBS:
            _jobs.popitem(last=False)


_ACTIVITY_LOG = base.MILA_FOLDER / "logs" / "user_activity.jsonl"


def _log_activity(rec: dict):
    """Пишет одно событие активности пользователя (JSONL) для анализа Стасом.
    Лог в .gitignore — содержит текст запросов. Не путать с session-notes клиентов:
    это запросы Людмилы к агентам офиса."""
    try:
        _ACTIVITY_LOG.parent.mkdir(parents=True, exist_ok=True)
        rec["ts"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        with open(_ACTIVITY_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass  # логирование активности не должно ронять чат


def _run_job(job_id, sid, key, msg, attachment=None):
    """Выполняется в фоновом потоке: вызывает агента, обновляет историю и job."""
    import time as _t
    _t0 = _t.time()
    try:
        with _locks[key]:  # сериализуем вызовы одного агента; разные агенты идут параллельно
            histories = _session_histories(sid)
            reply, new_hist = AGENTS[key]["responder"](msg, histories[key])
            histories[key] = _trim(new_hist)
            # Сохранить обновлённую историю в файлы для персистентности
            for item in new_hist:
                if isinstance(item, dict):
                    _save_session_message(sid, key, item.get("role", "user"), item.get("content", ""))
        _log_activity({"agent": key, "message": (msg or "")[:1000],
                       "msg_len": len(msg or ""), "has_attachment": bool(attachment),
                       "response_ms": int((_t.time() - _t0) * 1000), "ok": True})
        # Парсим VERDICT и очищаем его из видимого текста
        verdict = "ready_next"  # default
        reply_clean = reply
        verdict_match = re.search(r"\[VERDICT:\s*(\w+)\]", reply)
        if verdict_match:
            verdict = verdict_match.group(1)
            reply_clean = reply.replace(verdict_match.group(0), "").strip()

        # Document workflow tracking: если был attachment (загруженный файл), сохраняем этап
        doc_id = None
        if attachment:
            doc_id = attachment.get("_doc_id")
            if not doc_id:
                result = memory.start_document_workflow(
                    attachment.get("name", "document"),
                    attachment.get("text", "")
                )
                doc_id = result.get("doc_id")
                attachment["_doc_id"] = doc_id
            memory.add_workflow_stage(doc_id, agent=key, input_text=msg, output_text=reply_clean, verdict=verdict)
            try:
                _save_document_record(doc_id, key, msg, reply_clean, attachment)
            except Exception:
                logger.exception("could not save document record: %s", doc_id)

        for mentioned_doc_id in _extract_doc_ids(reply_clean):
            try:
                _save_document_record(mentioned_doc_id, key, msg, reply_clean, attachment)
            except Exception:
                logger.exception("could not save mentioned document record: %s", mentioned_doc_id)

        # Агент вернул готовый документ в маркерах → в чат отдаём комментарий без
        # полотна текста + кнопку скачать чистый файл (через [doc_id:…]).
        visible = reply_clean
        if _extract_final_document(reply_clean):
            visible = _strip_doc_block(reply_clean)
            if doc_id and doc_id not in visible:
                visible = (visible + "\n\n📄 Готовый документ с правками — кнопкой ниже."
                           f"\n[doc_id:{doc_id}]").strip()

        _set_job(job_id, {"status": "done", "sid": sid, "reply": visible, "doc_id": doc_id, "verdict": verdict})
    except Exception as e:
        logger.exception("Ошибка агента %s (job %s)", key, job_id)
        _log_activity({"agent": key, "message": (msg or "")[:1000],
                       "msg_len": len(msg or ""), "has_attachment": bool(attachment),
                       "response_ms": int((_t.time() - _t0) * 1000),
                       "ok": False, "error": str(e)[:120]})
        _set_job(job_id, {"status": "done", "sid": sid, "error": "Внутренняя ошибка — см. логи (logs/webapp.log)"})


# Короткое описание агента для тултипа (1-2 строки). Имена/роли — как в AGENTS.
_TAGLINES = {
    "marina":   "Контент, стратегия роста, аналитика Instagram.",
    "victoria": "Проверяет тексты перед публикацией: голос, хук, CTA.",
    "alina":    "Анкеты, подготовка к сессиям, CRM клиенток.",
    "dima":     "Доход, Gumroad, прогнозы и цели.",
    "tyoma":    "Telegram-канал, бот, welcome-цепочка.",
    "olya":     "Вирусные темы, тренды, конкуренты.",
    "vasya":    "Расписание публикаций и что нужно снять.",
    "lera":     "Воронка, офферы, конверсия в продажи.",
    "manager":  "Система, метрики офиса, самообучение.",
    "producer": "Продуктовая линейка, запуски, масштаб дохода.",
    "rita":     "Структура цифровых продуктов: воркбуки, гайды.",
}


@app.get("/api/meta")
def meta():
    _session_id()
    return jsonify({
        "csrf": _csrf_token(),
        "agents": [
            {"key": k, "name": a["name"], "role": a["role"], "emoji": a["emoji"],
             "color": a["color"], "intro": a["intro"], "chips": a["chips"],
             "tagline": _TAGLINES.get(k, "")}
            for k, a in AGENTS.items()
        ]
    })


@app.post("/api/upload")
def upload_file():
    sid = _session_id()
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "Файл не выбран"}), 400
    raw = f.read(MAX_UPLOAD_BYTES + 1)
    if len(raw) > MAX_UPLOAD_BYTES:
        return jsonify({"ok": False, "error": "Файл слишком большой. Максимум 12 МБ."}), 413
    name = _safe_upload_name(f.filename)
    mime = (f.mimetype or mimetypes.guess_type(name)[0] or "application/octet-stream").lower()
    suffix = Path(name).suffix.lower()
    kind = "image" if (mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}) else "file"
    text, warning = _extract_upload(name, raw, mime)
    text = _clip_text(text)
    if not text and warning:
        text = warning
    upload_id = uuid.uuid4().hex
    _set_upload(upload_id, {"sid": sid, "name": name, "mime": mime, "kind": kind, "text": text, "warning": warning})
    return jsonify({
        "ok": True,
        "upload_id": upload_id,
        "name": name,
        "mime": mime,
        "chars": len(text),
        "warning": warning,
        "excerpt": text[:700],
    })


@app.post("/api/chat")
def chat():
    sid = _session_id()
    data = request.get_json(force=True)
    key = data.get("agent")
    msg = (data.get("message") or "").strip()
    upload_id = (data.get("upload_id") or "").strip()
    if key not in AGENTS:
        return jsonify({"error": "Неизвестный агент"}), 400
    attachment = None
    if upload_id:
        attachment = _uploads.get(upload_id)
        if not attachment or attachment.get("sid") != sid:
            return jsonify({"error": "Файл не найден. Загрузи его ещё раз."}), 400
    if not msg and not attachment:
        return jsonify({"error": "Пустое сообщение"}), 400
    # Не блокируем запрос на время раздумий агента — ставим задачу в пул.
    if attachment:
        ask = msg or "Дай фидбек по загруженному файлу."
        if attachment.get("kind") == "image":
            msg = (
                f"{ask}\n\n"
                f"Пользователь загрузил скриншот/изображение: {attachment.get('name')} ({attachment.get('mime')}).\n"
                "Ниже уже дан OCR и визуальный разбор скриншота. Используй этот текст как содержимое изображения. "
                "Не отвечай, что не можешь прочитать скриншот, если в блоке ниже есть извлечённые фрагменты. "
                "Найди проблему, объясни её по делу и предложи следующий практический шаг.\n\n"
                "--- OCR / визуальный разбор скриншота ---\n"
                f"{attachment.get('text', '')}"
            )
        else:
            warn = attachment.get("warning") or ""
            warn_block = f"⚠️ Предупреждение при чтении файла: {warn}\n\n" if warn else ""
            msg = (
                f"{ask}\n\n"
                f"Пользователь загрузил файл: {attachment.get('name')} ({attachment.get('mime')}).\n"
                f"{warn_block}"
                "Дай конкретный фидбек на основе содержимого файла: что работает, что улучшить, "
                "какие правки внести и следующий практический шаг. Если содержимое не читается "
                "(битый текстовый слой PDF/каша) — не выдумывай оценку, а честно скажи об этом "
                "и попроси прислать текстовую версию (Markdown/.docx).\n\n"
                "--- Содержимое / описание файла ---\n"
                f"{attachment.get('text', '')}"
            )
    job_id = uuid.uuid4().hex
    _set_job(job_id, {"status": "pending", "sid": sid})
    _pool.submit(_run_job, job_id, sid, key, msg, attachment)
    return jsonify({"job": job_id}), 202


@app.get("/api/result")
def result():
    sid = _session_id()
    job_id = request.args.get("job", "")
    with _jobs_lock:
        job = _jobs.get(job_id)
        if job is not None and job.get("sid") != sid:
            job = None
        if job is not None and job.get("status") != "pending":
            _jobs.pop(job_id, None)  # результат забран — освобождаем память (читается один раз)
    if job is None:
        return jsonify({"error": "Задача не найдена"}), 404
    if job["status"] == "pending":
        return jsonify({"status": "pending"})
    # Готово: либо reply, либо error — фронтенд рендерит как раньше.
    return jsonify({k: v for k, v in job.items() if k != "status"})


@app.get("/api/document/<doc_id>/download")
def api_document_download_text(doc_id):
    doc = _load_document_record(doc_id)
    if doc.get("status") == "missing_source":
        abort(404)
    # Всегда регенерируем из актуальной записи — чтобы отдать последнюю чистую
    # версию (final_content), а не закэшированный транскрипт.
    path = _doc_export_path(doc_id)
    path.write_text(_document_download_text(doc), encoding="utf-8")
    # Если есть готовый документ — имя «mila-готовый-…», иначе прежнее имя.
    name = doc.get("file_name") or f"mila-document-{_safe_doc_id(doc_id)}.txt"
    if doc.get("final_content"):
        name = f"mila-готовый-{_safe_doc_id(doc_id)}.txt"
    return send_file(
        path,
        mimetype="text/plain; charset=utf-8",
        as_attachment=True,
        download_name=name,
    )




@app.post("/api/reset")
def reset():
    sid = _session_id()
    data = request.get_json(force=True) or {}
    histories = _session_histories(sid)
    if data.get("all"):
        # «Очистить сессию» — сбрасываем историю ВСЕХ агентов этой сессии.
        for k in list(histories.keys()):
            lock = _locks.get(k)
            if lock:
                with lock:  # не трогаем историю, пока агент думает
                    histories[k] = []
            else:
                histories[k] = []
        return jsonify({"ok": True, "cleared": "all"})
    key = data.get("agent")
    if key in histories:
        with _locks[key]:  # не сбрасываем историю, пока агент думает
            histories[key] = []
    return jsonify({"ok": True})


# ─── Настройки и подключения (статус + Instagram OAuth) ──
# Instagram API использует OAuth (в отличие от Anthropic Messages API — там только
# ключ). Здесь — реальный flow «Instagram API with Instagram Login»:
# authorize → callback → обмен кода на short-lived → обмен на long-lived → запись в .env.
IG_OAUTH_AUTHORIZE = "https://www.instagram.com/oauth/authorize"
IG_OAUTH_TOKEN_URL = "https://api.instagram.com/oauth/access_token"
IG_OAUTH_LONGLIVED = "https://graph.instagram.com/access_token"
IG_OAUTH_SCOPES = ("instagram_business_basic,instagram_business_manage_comments,"
                   "instagram_business_manage_messages,instagram_business_content_publish,"
                   "instagram_business_manage_insights")
# Meta требует HTTPS-redirect, точно зарегистрированный в кабинете приложения
# (плоский http и 127.0.0.1 отклоняются — отсюда «Invalid redirect_uri»).
# При MILA_HTTPS=1 приложение поднимается по https и redirect — https://localhost.
_HTTPS = os.getenv("MILA_HTTPS", "").lower() in ("1", "true", "yes")
IG_REDIRECT_URI = os.getenv("IG_OAUTH_REDIRECT_URI") or (
    "https://localhost:5000/auth/instagram/callback" if _HTTPS
    else "http://127.0.0.1:5000/auth/instagram/callback")


def _ig_app_creds():
    return (os.getenv("IG_APP_ID") or os.getenv("META_APP_ID"),
            os.getenv("IG_APP_SECRET") or os.getenv("META_APP_SECRET"))


def _integration_status():
    """Статусы подключений без раскрытия секретов — только флаги/режимы."""
    cid, _ = _ig_app_creds()
    return {
        "claude": {"configured": bool(base.ANTHROPIC_KEY or base.ANTHROPIC_AUTH_TOKEN),
                   "mode": "auth_token (Bearer)" if base.ANTHROPIC_AUTH_TOKEN else "api_key (x-api-key)",
                   "oauth": False,
                   "note": "У Messages API нет публичного OAuth — авторизация по ключу."},
        "gemini": {"configured": bool(base.GEMINI_KEY),
                   "model": base.GEMINI_MODEL,
                   "provider": base.LLM_PROVIDER,
                   "heavy_lifting": base.LLM_PROVIDER in ("gemini", "google"),
                   "anthropic_agents": sorted(base.ANTHROPIC_AGENT_KEYS)},
        "instagram": {"configured": bool(base.INSTAGRAM_TOKEN), "flow": base.IG_FLOW,
                      "node": base.IG_NODE or "", "oauth": bool(cid),
                      "redirect_uri": IG_REDIRECT_URI},
        "telegram": {"configured": bool(base.TELEGRAM_TOKEN), "oauth": False},
        "gumroad": {"configured": bool(base.GUMROAD_TOKEN), "oauth": False},
    }


def _update_env(updates):
    """In-place обновление ключей в tools/.env, сохраняя остальные строки."""
    env_path = base.MILA_FOLDER / "tools" / ".env"
    lines = env_path.read_text(encoding="utf-8").splitlines() if env_path.exists() else []
    for key, val in updates.items():
        if not val:
            continue
        for i, ln in enumerate(lines):
            if ln.strip().startswith(f"{key}="):
                lines[i] = f"{key}={val}"
                break
        else:
            lines.append(f"{key}={val}")
    env_path.parent.mkdir(parents=True, exist_ok=True)
    env_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _oauth_done(title, msg, show_test=False):
    test_html = ""
    if show_test:
        # Кнопка живой проверки + «что дальше» — чтобы после OAuth не было вопроса «и что теперь?».
        test_html = (
            "<p><button onclick='testIg()' style=\"font-family:inherit;background:#C4614A;color:#fff;"
            "border:0;border-radius:8px;padding:9px 16px;font-size:14px;cursor:pointer\">"
            "Проверить Instagram API</button> <span id='tr' style='color:#7A5E54;font-size:13px'></span></p>"
            "<p style='color:#7A5E54;font-size:13px'>Дальше: перезапусти приложение, чтобы агенты "
            "подхватили новый токен и доступ к комментариям.</p>"
            "<script>async function testIg(){var s=document.getElementById('tr');s.textContent='Проверяю…';"
            "try{var d=await (await fetch('/api/test/instagram')).json();"
            "s.innerHTML=d.ok?('\\u2713 @'+(d.username||'?')+' (id '+(d.user_id||'?')+')'):('\\u26A0\\uFE0F '+(d.error||'ошибка'));}"
            "catch(e){s.textContent='\\u26A0\\uFE0F сеть';}}</script>"
        )
    return Response(
        f"<!doctype html><meta charset=utf-8><body style='font-family:Georgia,serif;"
        f"max-width:560px;margin:60px auto;color:#1E140F'>"
        f"<h2 style='color:#C4614A'>{title}</h2><p>{msg}</p>{test_html}"
        f"<p><a href='/settings'>← Настройки</a></p></body>", mimetype="text/html")


@app.get("/api/settings")
def api_settings():
    _session_id()
    status = _integration_status()
    status["csrf"] = _csrf_token()
    return jsonify(status)


def _probe(url, timeout=2.5):
    """Живой GET-пробник локального сервиса: up/down + код. Без секретов."""
    try:
        r = requests.get(url, timeout=timeout)
        return {"up": r.status_code < 500, "status": r.status_code}
    except Exception as e:
        return {"up": False, "error": type(e).__name__}


@app.get("/api/health")
def api_health():
    """Единая health-сводка: конфигурация LLM/каналов (флаги) + ЖИВЫЕ пробы
    локальных сервисов (n8n, bridge) + доступность Supabase. Безопасно —
    только статусы, без ключей."""
    cfg = _integration_status()
    n8n_url = os.getenv("N8N_BASE_URL", "http://127.0.0.1:5678").rstrip("/")
    bridge_port = os.getenv("N8N_BRIDGE_PORT", "5051")
    # Supabase: проверяем только наличие конфига (живой запрос требует supa+сети).
    try:
        import sys as _sys
        tdir = str(base.MILA_FOLDER / "tools")
        if tdir not in _sys.path:
            _sys.path.insert(0, tdir)
        import supa
        supa_state = {"configured": supa.available(), "can_write": supa.can_write()}
    except Exception:
        supa_state = {"configured": False, "can_write": False}

    health = {
        "gemini":   {"configured": cfg["gemini"]["configured"], "model": cfg["gemini"]["model"]},
        "claude":   {"configured": cfg["claude"]["configured"]},
        "telegram": {"configured": cfg["telegram"]["configured"]},
        "instagram": {"configured": cfg["instagram"]["configured"], "flow": cfg["instagram"]["flow"]},
        "supabase": supa_state,
        "n8n":      _probe(f"{n8n_url}/healthz"),
        "bridge":   _probe(f"http://127.0.0.1:{bridge_port}/health"),
    }
    # Сводный флаг: всё ли критичное поднято (LLM + хотя бы один канал).
    health["ok"] = bool(
        (health["gemini"]["configured"] or health["claude"]["configured"])
        and health["telegram"]["configured"]
    )

    # Очередь: что требует руки человека (для статус-бара Healthy/Degraded/Action).
    try:
        st = memory.office_status(limit=1)
        counts = (st.get("tasks") or {}).get("counts") or {}
        approvals = st.get("approvals") or {}
    except Exception:
        counts, approvals = {}, {}
    approvals_waiting = sum(
        1 for v in approvals.values()
        if isinstance(v, dict) and v.get("status") in {"pending", "changes_requested"}
    )
    attention = {
        "failed": int(counts.get("failed", 0) or 0),
        "awaiting_approval": int(counts.get("awaiting_approval", 0) or 0) + approvals_waiting,
        "running": int(counts.get("running", 0) or 0),
        "pending": int(counts.get("pending", 0) or 0),
    }
    health["attention"] = attention

    # Уровень: action_needed (критично/ждёт человека) > degraded (автоматика хромает) > healthy.
    core_ok = health["gemini"]["configured"] or health["claude"]["configured"]
    critical, degraded = [], []
    if not core_ok:
        critical.append("нет LLM (Gemini/Claude)")
    if attention["failed"]:
        critical.append(f"{attention['failed']} задач с ошибкой")
    if attention["awaiting_approval"]:
        critical.append(f"{attention['awaiting_approval']} ждут одобрения")
    if not health["telegram"]["configured"]:
        degraded.append("Telegram не настроен")
    if not health["instagram"]["configured"]:
        degraded.append("Instagram не подключён")
    if not health["supabase"]["configured"]:
        degraded.append("Supabase не настроен")
    if not health["bridge"].get("up"):
        degraded.append("n8n-мост недоступен")
    if not health["n8n"].get("up"):
        degraded.append("n8n недоступен")

    if critical:
        health["level"], health["reasons"] = "action_needed", critical
    elif degraded:
        health["level"], health["reasons"] = "degraded", degraded
    else:
        health["level"], health["reasons"] = "healthy", []
    return jsonify(health)


@app.get("/settings")
def settings_page():
    return Response(SETTINGS_HTML, mimetype="text/html")


@app.get("/auth/instagram/login")
def ig_login():
    cid, csec = _ig_app_creds()
    if not cid or not csec:
        return ("Instagram OAuth не настроен: задай IG_APP_ID и IG_APP_SECRET "
                "(или META_APP_ID/META_APP_SECRET) в tools/.env и зарегистрируй redirect URI "
                f"({IG_REDIRECT_URI}) в кабинете Meta."), 400
    state = secrets.token_urlsafe(16)
    session["ig_state"] = state
    params = {"client_id": cid, "redirect_uri": IG_REDIRECT_URI, "response_type": "code",
              "scope": IG_OAUTH_SCOPES, "state": state}
    return redirect(IG_OAUTH_AUTHORIZE + "?" + urlencode(params))


@app.get("/auth/instagram/callback")
def ig_callback():
    if request.args.get("error"):
        return f"OAuth ошибка: {request.args.get('error_description') or request.args.get('error')}", 400
    code = request.args.get("code")
    if not code or request.args.get("state") != session.pop("ig_state", None):
        return "Неверный ответ OAuth (нет code или несовпадение state).", 400
    cid, csec = _ig_app_creds()
    try:
        r = requests.post(IG_OAUTH_TOKEN_URL, data={
            "client_id": cid, "client_secret": csec, "grant_type": "authorization_code",
            "redirect_uri": IG_REDIRECT_URI, "code": code}, timeout=15)
        r.raise_for_status()
        short = r.json()
        token, user_id = short.get("access_token"), str(short.get("user_id") or "")
        # short-lived → long-lived (~60 дней)
        rl = requests.get(IG_OAUTH_LONGLIVED, params={
            "grant_type": "ig_exchange_token", "client_secret": csec, "access_token": token}, timeout=15)
        long_token = rl.json().get("access_token", token) if rl.ok else token
        _update_env({"IG_API_FLOW": "instagram_login",
                     "IG_ACCESS_TOKEN": long_token, "IG_USER_ID": user_id})
        logger.info("Instagram OAuth: токен сохранён (user_id=%s)", user_id)
        return _oauth_done(
            "Instagram подключён ✓",
            f"Аккаунт ID: <code>{user_id or '—'}</code><br>"
            "Долгоживущий токен (~60 дней) сохранён в tools/.env.",
            show_test=True)
    except Exception:
        logger.exception("Instagram OAuth token exchange failed")
        return "Обмен кода на токен не удался — подробности в logs/webapp.log.", 500


def _current_ig_token() -> str:
    """Свежий IG-токен из tools/.env (после OAuth/ручного сохранения) — base.* мог
    загрузиться со старым значением до перезапуска. Fallback — то, что в base."""
    for envf in (base.MILA_FOLDER / "tools" / ".env", base.MILA_FOLDER / ".env"):
        try:
            for line in envf.read_text(encoding="utf-8").splitlines():
                s = line.strip()
                for key in ("IG_ACCESS_TOKEN=", "INSTAGRAM_ACCESS_TOKEN="):
                    if s.startswith(key):
                        return s.split("=", 1)[1].strip().strip('"').strip("'")
        except OSError:
            pass
    return getattr(base, "INSTAGRAM_TOKEN", "") or ""


@app.get("/api/test/<service>")
def api_test(service):
    """Живая проверка одного подключения по кнопке «Проверить». Read-only (GET):
    ничего не меняет, только дёргает соответствующий API и возвращает {ok, detail}."""
    service = (service or "").lower()
    try:
        if service == "instagram":
            # .env могли только что обновить (новый токен/flow/IDs) — перечитываем,
            # чтобы тест отражал актуальные значения, а не загруженные при старте.
            from dotenv import load_dotenv
            load_dotenv(base.MILA_FOLDER / "tools" / ".env", override=True)
            load_dotenv(base.MILA_FOLDER / ".env", override=True)
            token = _current_ig_token()
            if not token:
                return jsonify({"ok": False, "error": "Токен Instagram не задан"}), 400
            # Facebook-токен (EAA…) ходит через graph.facebook.com + node=IG_USER_ID;
            # instagram_login-токен — через graph.instagram.com + node=me. Берём по flow.
            flow = (os.getenv("IG_API_FLOW", "facebook") or "facebook").strip().lower()
            ver = os.getenv("GRAPH_API_VERSION", "v21.0")
            if flow == "instagram_login":
                host, node = "https://graph.instagram.com", "me"
            else:
                host, node = "https://graph.facebook.com", (os.getenv("IG_USER_ID") or "").strip() or "me"
            r = requests.get(f"{host}/{ver}/{node}",
                             params={"fields": "id,username,name", "access_token": token}, timeout=10)
            info = r.json()
            if "error" in info:
                return jsonify({"ok": False, "error": info["error"].get("message", "невалидный токен"),
                                "flow": flow}), 400
            uname = info.get("username") or info.get("name")
            return jsonify({"ok": True, "username": uname, "user_id": str(info.get("id") or ""),
                            "flow": flow, "detail": ("@" + uname) if uname else ("id " + str(info.get("id") or "?"))})
        if service == "telegram":
            token = (getattr(base, "TELEGRAM_TOKEN", "") or "").strip()
            if not token:
                return jsonify({"ok": False, "error": "TELEGRAM_BOT_TOKEN не задан"}), 400
            r = requests.get(f"https://api.telegram.org/bot{token}/getMe", timeout=10)
            j = r.json()
            if not j.get("ok"):
                return jsonify({"ok": False, "error": j.get("description", "невалидный токен")}), 400
            return jsonify({"ok": True, "detail": "@" + (j.get("result", {}).get("username") or "?")})
        if service in ("bridge", "n8n"):
            if service == "bridge":
                url = f"http://127.0.0.1:{os.getenv('N8N_BRIDGE_PORT', '5051')}/health"
            else:
                url = os.getenv("N8N_BASE_URL", "http://127.0.0.1:5678").rstrip("/") + "/healthz"
            p = _probe(url)
            return jsonify({"ok": bool(p.get("up")), "detail": f"HTTP {p.get('status', p.get('error', '?'))}"})
        if service == "supabase":
            try:
                import sys as _sys
                tdir = str(base.MILA_FOLDER / "tools")
                if tdir not in _sys.path:
                    _sys.path.insert(0, tdir)
                import supa
                if not supa.available():
                    return jsonify({"ok": False, "error": "Supabase не настроен (нет URL/ключа)"}), 400
                supa.select("purchases", columns="id", limit=1)  # лёгкий ping
                return jsonify({"ok": True, "detail": "запись" if supa.can_write() else "только чтение"})
            except Exception as e:
                return jsonify({"ok": False, "error": str(e)[:200]}), 400
        if service in ("claude", "gemini"):
            cfg = _integration_status().get(service, {})
            return jsonify({"ok": bool(cfg.get("configured")),
                            "detail": cfg.get("model") or ("настроен" if cfg.get("configured") else "не настроен")})
    except Exception:
        logger.exception("api_test failed for %s", service)
        return jsonify({"ok": False, "error": "Ошибка проверки — см. логи"}), 500
    return jsonify({"ok": False, "error": f"неизвестный сервис: {service}"}), 404


@app.post("/api/settings/instagram-token")
def ig_save_token():
    """Ручное подключение Instagram: проверяем вставленный токен через
    graph.instagram.com/me (без redirect URI!) и сохраняем в tools/.env.
    Надёжный путь на localhost, когда Meta отклоняет redirect_uri."""
    data = request.get_json(force=True) or {}
    token = (data.get("token") or "").strip()
    if not token:
        return jsonify({"ok": False, "error": "Пустой токен"}), 400
    try:
        r = requests.get("https://graph.instagram.com/v21.0/me",
                         params={"fields": "user_id,username", "access_token": token}, timeout=10)
        info = r.json()
        if "error" in info:
            return jsonify({"ok": False, "error": info["error"].get("message", "невалидный токен")}), 400
        user_id = str(info.get("user_id") or info.get("id") or (data.get("user_id") or ""))
        _update_env({"IG_API_FLOW": "instagram_login",
                     "IG_ACCESS_TOKEN": token, "IG_USER_ID": user_id})
        logger.info("IG token сохранён вручную (user=%s)", info.get("username"))
        return jsonify({"ok": True, "username": info.get("username"), "user_id": user_id})
    except Exception:
        logger.exception("manual IG token save failed")
        return jsonify({"ok": False, "error": "Не удалось проверить/сохранить токен — см. логи"}), 500


# ─── Дашборд Людмилы: «готовят → она одобряет» ───────────
# Собирает всё, что ждёт её решения, и даёт одну кнопку «Одобрить всё».
_POST_QUEUE = base.MILA_FOLDER / "MILA-BUSINESS" / "02-content" / "post_queue.json"
_OFFICE_ACTIONS = base.MILA_FOLDER / "reports" / "office_actions.json"
_OVERRIDES_DIR = base.MILA_FOLDER / "MILA-BUSINESS" / "05-analytics" / "prompt_overrides"


def _read_json_safe(path, default):
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except (FileNotFoundError, ValueError, OSError):
        return default


def _pending_posts():
    """Посты в очереди, ждущие одобрения (не approved/published/needs_media)."""
    q = _read_json_safe(_POST_QUEUE, [])
    out = []
    for it in q:
        if it.get("status") in ("draft", "pending", "review", None):
            out.append({"id": it.get("id"), "when": it.get("when"),
                        "caption": (it.get("caption") or "")[:160],
                        "has_media": bool((it.get("media_url") or "").strip())})
    return out


def _open_actions():
    """Открытые задачи офиса от Стаса (office_actions.json)."""
    acts = _read_json_safe(_OFFICE_ACTIONS, [])
    return [{"id": a.get("id"), "title": a.get("title"), "assignee": a.get("assignee"),
             "priority": a.get("priority"), "due": a.get("due")}
            for a in acts if a.get("status") == "open"]


def _recent_events(limit=12):
    """Последние события из общей памяти (что агенты делали)."""
    f = base.MILA_FOLDER / "mila-office" / "memory" / "events.jsonl"
    try:
        lines = f.read_text(encoding="utf-8").splitlines()[-limit:]
    except (FileNotFoundError, OSError):
        return []
    out = []
    for ln in reversed(lines):
        try:
            e = json.loads(ln)
            out.append({"ts": e.get("ts", "")[:16].replace("T", " "),
                        "kind": e.get("kind", "")})
        except ValueError:
            continue
    return out


def _pending_improvements():
    """Активные улучшения Стаса (overrides) — что он предлагает агентам."""
    out = []
    if _OVERRIDES_DIR.exists():
        for f in sorted(_OVERRIDES_DIR.glob("*.md")):
            if f.name in ("README.md", "improvement_log.md"):
                continue
            txt = f.read_text(encoding="utf-8")
            topics = re.findall(r"^##\s+(.+?)\s+—", txt, re.MULTILINE)
            if topics:
                out.append({"agent": f.stem, "topics": topics})
    return out


def _latest_report(prefix):
    """Свежий reports/<prefix>_*.json (по mtime) или None."""
    rep = base.MILA_FOLDER / "reports"
    files = sorted(rep.glob(f"{prefix}_*.json"), key=lambda p: p.stat().st_mtime)
    if not files:
        return None
    return _read_json_safe(files[-1], None)


def _kpi():
    """Реальные метрики для дашборда. Берём из свежих отчётов + БД; ничего не
    выдумываем — чего нет, отдаём None (фронт покажет «—»)."""
    kpi = {"followers": None, "avg_reach": None, "er": None,
           "sales": None, "phase": _safe_phase(), "goal": None}
    try:
        prof = memory.read_profile().get("business", {})
        kpi["goal"] = prof.get("goal")
        kpi["followers"] = prof.get("ig_followers")
    except Exception:
        pass
    # followers — из свежего account-отчёта, точнее профиля
    acc = _latest_report("account")
    if isinstance(acc, dict) and acc.get("followers_count"):
        kpi["followers"] = acc["followers_count"]
    # avg reach / ER — из свежего posts-отчёта
    pr = _latest_report("posts")
    posts = (pr.get("posts") if isinstance(pr, dict) else pr) or []
    posts = [p for p in posts if isinstance(p, dict)]
    if posts:
        reaches = [p.get("reach") or 0 for p in posts]
        avg_reach = round(sum(reaches) / len(reaches)) if reaches else 0
        kpi["avg_reach"] = avg_reach
        # ER = средняя вовлечённость / охват
        if avg_reach:
            eng = [p.get("engagement") or (p.get("likes", 0) + p.get("comments", 0)) for p in posts]
            kpi["er"] = round(100 * (sum(eng) / len(eng)) / avg_reach, 1)
    # продажи — из БД (purchases), как в фазе
    try:
        kpi["sales"] = memory.sales_count()
    except Exception:
        pass
    return kpi


@app.get("/api/dashboard")
def api_dashboard():
    _session_id()
    prof = {}
    try:
        prof = memory.read_profile().get("business", {})
    except Exception:
        pass
    return jsonify({
        "csrf": _csrf_token(),
        "pending_posts": _pending_posts(),
        "open_actions": _open_actions(),
        "improvements": _pending_improvements(),
        "events": _recent_events(),
        "kpi": _kpi(),
        "profile": {"phase": (lambda: _safe_phase())(),
                    "followers": prof.get("ig_followers"),
                    "goal": prof.get("goal")},
    })


@app.post("/api/dashboard/post/<post_id>/<action>")
def api_dashboard_post(post_id, action):
    """Действие по ОДНОМУ посту из очереди: approve | reject.
    approve → status=approved (опубликует publish_due); reject → status=rejected."""
    if action not in ("approve", "reject"):
        return jsonify({"ok": False, "error": "unknown action"}), 400
    q = _read_json_safe(_POST_QUEUE, [])
    found = None
    for it in q:
        if str(it.get("id")) == str(post_id):
            if action == "approve":
                if not (it.get("media_url") or "").strip():
                    it["status"] = "needs_media"
                    found = "needs_media"
                else:
                    it["status"] = "approved"; found = "approved"
            else:
                it["status"] = "rejected"; found = "rejected"
            break
    if found is None:
        return jsonify({"ok": False, "error": "post not found"}), 404
    try:
        Path(_POST_QUEUE).write_text(json.dumps(q, ensure_ascii=False, indent=2), encoding="utf-8")
        memory.log_event("dashboard:post", {"id": post_id, "result": found})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True, "result": found})


@app.post("/api/dashboard/action/<action_id>/close")
def api_dashboard_action_close(action_id):
    """Закрыть ОДНУ задачу офиса (status open → done)."""
    acts = _read_json_safe(_OFFICE_ACTIONS, [])
    found = False
    for a in acts:
        if str(a.get("id")) == str(action_id) and a.get("status") == "open":
            a["status"] = "done"
            a["closed"] = datetime.now().strftime("%Y-%m-%d %H:%M")
            found = True
            break
    if not found:
        return jsonify({"ok": False, "error": "action not found"}), 404
    try:
        Path(_OFFICE_ACTIONS).write_text(json.dumps(acts, ensure_ascii=False, indent=2), encoding="utf-8")
        memory.log_event("dashboard:action_close", {"id": action_id})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    return jsonify({"ok": True})


def _safe_phase():
    try:
        return memory.current_phase()
    except Exception:
        return "—"


@app.post("/api/approve-all")
def api_approve_all():
    """Главная кнопка: одобрить все черновики постов из очереди (status→approved).
    Безопасно: только посты, у которых есть медиа; публикует их потом publish_due
    по расписанию (не здесь). Улучшения Стаса и задачи — отдельными кнопками."""
    q = _read_json_safe(_POST_QUEUE, [])
    approved = 0
    for it in q:
        if it.get("status") in ("draft", "pending", "review", None):
            if (it.get("media_url") or "").strip():
                it["status"] = "approved"
                approved += 1
            else:
                it["status"] = "needs_media"
    if approved or q:
        try:
            Path(_POST_QUEUE).write_text(json.dumps(q, ensure_ascii=False, indent=2),
                                         encoding="utf-8")
        except OSError as e:
            return jsonify({"ok": False, "error": str(e)}), 500
    try:
        memory.log_event("dashboard:approve_all", {"approved": approved})
    except Exception:
        pass
    return jsonify({"ok": True, "approved": approved})


@app.post("/api/pipeline/publish_due")
def api_publish_due():
    """Webhook для n8n: опубликовать все посты которые пора (status=approved и время наступило).
    Вызывается по расписанию из n8n вместо Windows Task Scheduler.
    Примеры вызова:
      curl -X POST http://localhost:5000/api/pipeline/publish_due
      (из n8n webhook каждый час)
    """
    try:
        # Импортируем publish_due из tools/pipeline.py
        import sys
        from pathlib import Path
        tools_dir = Path(os.getenv("MILA_FOLDER", r"E:\MILA GOLD")) / "tools"
        if str(tools_dir) not in sys.path:
            sys.path.insert(0, str(tools_dir))

        import pipeline as tools_pipeline

        # Запускаем publish_due из tools
        result = tools_pipeline.publish_due()

        logger.info(f"API publish_due: {result}")
        try:
            memory.log_event("api:publish_due", {"result": result})
        except Exception:
            pass

        return jsonify({
            "ok": True,
            "status": "published",
            "published": result.get("published", 0),
            "changed": result.get("changed", 0)
        })
    except Exception as e:
        logger.error(f"API publish_due error: {e}", exc_info=True)
        return jsonify({
            "ok": False,
            "error": str(e)
        }), 500


@app.get("/api/errors/stats")
def api_error_stats():
    """Получить статистику ошибок за последние 24 часа.

    Response:
      {
        "ok": true,
        "period": "last 24 hours",
        "total_errors": 5,
        "by_type": {"ValueError": 2, "TimeoutError": 1},
        "by_level": {"ERROR": 4, "CRITICAL": 1},
        "by_context": {"webapp": 3, "pipeline": 2}
      }
    """
    try:
        hours = request.args.get("hours", 24, type=int)
        stats = error_monitor.get_error_stats(hours=hours)
        return jsonify({"ok": True, **stats})
    except Exception as e:
        logger.error(f"Error stats error: {e}", exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/api/errors/recent")
def api_recent_errors():
    """Получить последние N ошибок с деталями.

    Query params:
      limit=10 (по умолчанию)

    Response:
      {
        "ok": true,
        "errors": [
          {
            "timestamp": "2026-06-08T14:05:32Z",
            "level": "ERROR",
            "error_type": "ValueError",
            "error_message": "invalid literal",
            "context": {"agent": "lera"}
          }
        ]
      }
    """
    try:
        limit = request.args.get("limit", 10, type=int)
        errors = error_monitor.get_recent_errors(limit=limit)
        return jsonify({"ok": True, "count": len(errors), "errors": errors})
    except Exception as e:
        logger.error(f"Recent errors error: {e}", exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.post("/api/check-sensitive-data")
def api_check_sensitive_data():
    """Проверить файл или текст на наличие конфиденциальных данных.

    POST data:
      {
        "type": "text" | "file",
        "content": "some text here",    # если type=text
        "file_path": "03-clients/..."   # если type=file
      }

    Response:
      {
        "ok": true,
        "has_sensitive": false,
        "patterns_found": {},
        "total_matches": 0
      }
    """
    try:
        data = request.get_json() or {}
        check_type = data.get("type", "text")

        if check_type == "text":
            text = data.get("content", "")
            # Проверяем текст на паттерны
            patterns_found = {}
            total_matches = 0
            for pattern_name, pattern in data_sanitizer.PATTERNS.items():
                matches = re.findall(pattern, text, flags=re.IGNORECASE)
                if matches:
                    patterns_found[pattern_name] = len(matches)
                    total_matches += len(matches)

            return jsonify({
                "ok": True,
                "has_sensitive": total_matches > 0,
                "patterns_found": patterns_found,
                "total_matches": total_matches
            })

        elif check_type == "file":
            file_path = data.get("file_path", "")
            result = data_sanitizer.check_file_for_sensitive_data(file_path)
            return jsonify({"ok": True, **result})

        else:
            return jsonify({"ok": False, "error": "Invalid type (use 'text' or 'file')"}), 400

    except Exception as e:
        logger.error(f"Check sensitive data error: {e}", exc_info=True)
        return jsonify({"ok": False, "error": str(e)}), 500


@app.get("/dashboard")
def dashboard_page():
    return Response(DASHBOARD_HTML, mimetype="text/html")


@app.get("/operator")
def operator_page():
    return Response(OPERATOR_HTML, mimetype="text/html")


@app.get("/api/operator")
def api_operator():
    status = request.args.get("status") or None
    tasks = memory.list_tasks(status)
    approvals = memory.office_status().get("approvals", {})
    pending_approvals = {
        k: v for k, v in approvals.items()
        if v.get("status") in {"pending", "missing", "changes_requested"}
    }
    supervisor = memory.read_supervisor_status()
    try:
        n8n_url = os.getenv("N8N_BASE_URL", "http://127.0.0.1:5678").rstrip("/")
        bridge_port = os.getenv("N8N_BRIDGE_PORT", "5051")
        services = {
            "webapp": {"up": True, "status": 200},
            "bridge": _probe(f"http://127.0.0.1:{bridge_port}/health"),
            "n8n": _probe(f"{n8n_url}/healthz"),
        }
        supervisor = {
            **supervisor,
            "ok": all(bool(v.get("up")) for v in services.values()),
            "status": "ok" if all(bool(v.get("up")) for v in services.values()) else "degraded",
            "services": services,
        }
    except Exception:
        logger.exception("operator live supervisor status failed")
    # Reply queue status — для виджета очереди ответов на Instagram-комментарии.
    try:
        reply_queue = memory.reply_queue_status(limit=10)
    except Exception:
        reply_queue = {"pending": 0, "sent": 0, "failed": 0, "items_pending": []}
    return jsonify({
        "ok": True,
        "csrf": _csrf_token(),
        "tasks": tasks,
        "status": memory.office_status(limit=30),
        "supervisor": supervisor,
        "events": memory.recent_events(30),
        "pending_approvals": pending_approvals,
        "reply_queue": reply_queue,
    })


@app.post("/api/operator/task/<task_id>/<action>")
def api_operator_task(task_id: str, action: str):
    body = request.get_json(silent=True) or {}
    if action == "retry":
        rec = memory.retry_task(task_id, reset_attempts=bool(body.get("reset_attempts")))
    elif action == "cancel":
        rec = memory.cancel_task(task_id, reason=body.get("reason", "operator"))
    elif action == "unblock":
        rec = memory.unblock_task(task_id)
    else:
        return jsonify({"ok": False, "error": f"unknown action: {action}"}), 400
    return jsonify({"ok": bool(rec.get("id")), "task": rec}), (200 if rec.get("id") else 400)


@app.post("/api/reply-send-one")
def api_reply_send_one():
    """Отправить один ответ из очереди в Instagram (вручную из оператора)."""
    # Динамически импортируем reply_sender чтобы избежать циклических импортов.
    import sys as _sys
    _tools_dir = str(base.MILA_FOLDER / "mila-office")
    if _tools_dir not in _sys.path:
        _sys.path.insert(0, _tools_dir)
    try:
        import reply_sender
    except ImportError:
        return jsonify({"ok": False, "error": "reply_sender модуль не найден"}), 500

    rep = memory.dequeue_reply()
    if not rep:
        return jsonify({"ok": False, "error": "В очереди нет ответов"}), 400

    ok, err, resp_id = reply_sender.post_reply(rep["comment_id"], rep["message"])
    if ok:
        memory.mark_reply(rep["id"], "sent", response_id=resp_id)
        logger.info("Manual reply sent: %s → %s", rep["id"], rep["comment_id"])
        return jsonify({"ok": True, "detail": f"@{rep.get('username', '?')}"})
    else:
        memory.mark_reply(rep["id"], "failed", error=err)
        logger.warning("Manual reply failed: %s: %s", rep["id"], err)
        return jsonify({"ok": False, "error": err}), 400


@app.post("/api/reply-send-all")
def api_reply_send_all():
    """Отправить все ответы из очереди подряд (оператор)."""
    import sys as _sys
    _tools_dir = str(base.MILA_FOLDER / "mila-office")
    if _tools_dir not in _sys.path:
        _sys.path.insert(0, _tools_dir)
    try:
        import reply_sender
    except ImportError:
        return jsonify({"ok": False, "error": "reply_sender модуль не найден"}), 500

    sent, failed = 0, 0
    while True:
        rep = memory.dequeue_reply()
        if not rep:
            break
        ok, err, resp_id = reply_sender.post_reply(rep["comment_id"], rep["message"])
        if ok:
            memory.mark_reply(rep["id"], "sent", response_id=resp_id)
            sent += 1
        else:
            memory.mark_reply(rep["id"], "failed", error=err)
            failed += 1
        # Небольшая пауза между ответами чтобы не выглядеть спамом.
        import time
        time.sleep(2)
    logger.info("Manual send-all: %d sent, %d failed", sent, failed)
    return jsonify({"ok": True, "sent": sent, "failed": failed})


@app.post("/api/reply-delete/<reply_id>")
def api_reply_delete(reply_id: str):
    """Удалить ответ из очереди (оператор)."""
    # Просто помечаем как cancelled.
    rec = memory.mark_reply(reply_id, "cancelled", error="Удалён оператором")
    if rec.get("id"):
        logger.info("Manual reply delete: %s", reply_id)
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "Ответ не найден"}), 400


@app.get("/api/documents")
def api_documents():
    """Список активных и завершённых документ-workflows."""
    in_progress = memory.list_workflows("in_progress", limit=20)
    completed = memory.list_workflows("completed", limit=20)
    return jsonify({
        "ok": True,
        "in_progress": in_progress,
        "completed": completed
    })


@app.get("/api/document/<doc_id>")
def api_document(doc_id: str):
    """Получить историю документа через все этапы обработки."""
    doc = memory.get_document_workflow(doc_id)
    if doc.get("ok") is False:
        return jsonify({"ok": False, "error": "Document not found"}), 404
    return jsonify({"ok": True, "document": doc})


@app.post("/api/document/<doc_id>/feedback")
def api_document_feedback(doc_id: str):
    """Отправить правки от одного агента к другому."""
    data = request.get_json() or {}
    from_agent = data.get("from_agent")
    to_agent = data.get("to_agent")
    feedback = data.get("feedback", "")

    if not (from_agent and to_agent and feedback):
        return jsonify({"ok": False, "error": "Missing from_agent, to_agent, or feedback"}), 400

    result = memory.add_backward_feedback(doc_id, from_agent, to_agent, feedback)
    return jsonify(result)


@app.post("/api/document/<doc_id>/archive")
def api_document_archive(doc_id: str):
    """Архивировать документ."""
    result = memory.archive_document(doc_id)
    return jsonify(result)


@app.post("/api/document/<doc_id>/export")
def api_document_export(doc_id: str):
    """Экспортировать документ со всей историей (JSON)."""
    result = memory.export_document(doc_id)
    if not result.get("ok"):
        return jsonify(result), 404

    # Отдаём как JSON-файл для скачивания
    export = result.get("export", {})
    filename = f"{export.get('file_name', 'document')}_history.json"
    response = Response(json.dumps(export, ensure_ascii=False, indent=2),
                       mimetype="application/json")
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@app.post("/api/agent-message")
def api_agent_message():
    """Сохранить сообщение от пользователя или агента в историю."""
    data = request.get_json() or {}
    agent = data.get("agent", "").strip()
    text = data.get("text", "").strip()
    is_user = data.get("is_user", False)
    verdict = data.get("verdict")

    # Validation
    if not agent:
        return jsonify({"ok": False, "error": "Missing agent"}), 400
    if not text:
        return jsonify({"ok": False, "error": "Empty message"}), 400
    if len(text) > 50000:
        return jsonify({"ok": False, "error": "Message too long (max 50000 chars)"}), 400
    if len(agent) > 50 or not agent.replace("_", "").replace("-", "").isalnum():
        return jsonify({"ok": False, "error": "Invalid agent name"}), 400

    # Validate verdict if provided
    if verdict and verdict not in ("ready_next", "needs_revision", "done"):
        return jsonify({"ok": False, "error": "Invalid verdict"}), 400

    try:
        result = memory.save_agent_message(agent, text, is_user, verdict)
        return jsonify(result)
    except Exception as e:
        logger.error(f"Error saving agent message: {e}", exc_info=True)
        return jsonify({"ok": False, "error": "Server error"}), 500


@app.get("/api/agent-history/<agent>")
def api_get_agent_history(agent: str):
    """Получить историю переписки с агентом."""
    result = memory.get_agent_history(agent)
    return jsonify(result)


@app.get("/api/agent-histories")
def api_list_agent_histories():
    """Список всех историй с информацией о количестве сообщений."""
    histories = memory.list_agent_histories()
    return jsonify({"ok": True, "histories": histories})


@app.post("/api/agent-history/<agent>/clear")
def api_clear_agent_history(agent: str):
    """Очистить историю конкретного агента."""
    result = memory.clear_agent_history(agent)
    return jsonify(result)


@app.post("/api/agent-histories/clear-all")
def api_clear_all_histories():
    """Очистить все истории."""
    result = memory.clear_all_histories()
    return jsonify(result)


@app.get("/favicon.ico")
def favicon():
    # Браузер всегда просит /favicon.ico — без маршрута это 404 в консоли.
    # Отдаём маленькую SVG-иконку в фирменном цвете (буква M), 404 уходит.
    svg = ("<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 32 32'>"
           "<rect width='32' height='32' rx='7' fill='#C4614A'/>"
           "<text x='16' y='23' font-size='19' text-anchor='middle' fill='#fff'"
           " font-family='Georgia,serif'>M</text></svg>")
    return Response(svg, mimetype="image/svg+xml")


@app.get("/")
def index():
    return Response(INDEX_HTML, mimetype="text/html")


# ─── Фронтенд (один файл, без внешних зависимостей/CDN) ───
INDEX_HTML = r"""<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>MILA OFFICE</title>
<style>
  :root{
    --t:#C4614A; --n:#1E140F; --c:#FAF6F1; --m:#F2EAE2;
    --u:#7A5E54; --b:#E0D0C8; --w:#fff;
  }
  *{box-sizing:border-box}
  body{margin:0;font-family:Georgia,'Times New Roman',serif;background:var(--c);color:var(--n);height:100vh;display:flex;overflow:hidden}
  #side{width:84px;background:var(--n);display:flex;flex-direction:column;align-items:center;padding:14px 0;gap:8px;flex-shrink:0;overflow-y:auto}
  #side .logo{color:var(--t);font-size:10px;letter-spacing:2px;margin-bottom:8px;text-align:center;line-height:1.3}
  .apill{width:54px;height:54px;border-radius:16px;border:2px solid transparent;background:rgba(255,255,255,.06);display:flex;flex-direction:column;align-items:center;justify-content:center;cursor:pointer;transition:.15s;color:#cbb}
  .apill:hover{background:rgba(255,255,255,.12)}
  .apill.active{background:rgba(255,255,255,.10)}
  .apill .em{font-size:20px;line-height:1}
  .apill .nm{font-size:9px;margin-top:3px;color:#cbb}
  /* Тултип агента — position:fixed (через JS), чтобы не обрезался overflow сайдбара.
     Появляется справа от кнопки при hover с задержкой 300ms. */
  .atip{position:fixed;z-index:30;width:220px;max-width:220px;
        background:#1E140F;border:1px solid #c08;border-radius:10px;padding:12px 14px;
        box-shadow:0 4px 16px rgba(0,0,0,.3);pointer-events:none;
        opacity:0;transform:translateX(-4px);transition:opacity .12s,transform .12s}
  .atip.show{opacity:1;transform:translateX(0)}
  .atip .h{font-size:12px;font-weight:bold;margin-bottom:6px}
  .atip .d{font-size:11px;color:#c0a898;line-height:1.5;margin-bottom:8px}
  .atip .cmds{display:flex;flex-wrap:wrap;gap:5px}
  .atip .cmd{font-size:10px;padding:2px 7px;border-radius:10px;color:#e6d7cc}
  #main{flex:1;display:flex;flex-direction:column;min-width:0}
  header{background:var(--n);padding:16px 22px;display:flex;align-items:center;gap:14px}
  header .av{width:46px;height:46px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:20px;color:#fff;flex-shrink:0}
  header .ttl{flex:1;min-width:0}
  header .ttl .nm{font-size:19px;font-weight:bold;color:#fff}
  header .ttl .sub{font-size:12px;color:#9a8278;margin-top:2px}
  header .toggle{background:rgba(196,97,74,.18);color:var(--t);border:1px solid var(--t);border-radius:8px;padding:8px 14px;font-size:12px;cursor:pointer;font-family:inherit}
  .chips{display:flex;gap:10px;flex-wrap:wrap;padding:14px 22px;background:var(--m);border-bottom:1px solid var(--b)}
  .chip{background:var(--w);border:1px solid var(--b);border-radius:20px;padding:9px 18px;font-size:13px;cursor:pointer;font-family:inherit;color:var(--n);transition:.15s}
  .chip:hover{border-color:var(--t);color:var(--t)}
  #chat{flex:1;overflow-y:auto;padding:24px 22px;display:flex;flex-direction:column;gap:16px}
  .row{display:flex;gap:12px;align-items:flex-start;max-width:78%}
  .row.me{align-self:flex-end;flex-direction:row-reverse}
  .row .av{width:38px;height:38px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:15px;color:#fff;flex-shrink:0}
  .bubble{background:var(--w);border:1px solid var(--b);border-radius:16px;padding:14px 18px;font-size:15px;line-height:1.6;white-space:pre-wrap;word-wrap:break-word}
  .row.me .bubble{background:var(--t);color:#fff;border-color:var(--t)}
  .bubble b{font-weight:bold}.bubble i{font-style:italic}
  .bubble code{background:rgba(0,0,0,.06);padding:1px 5px;border-radius:4px;font-family:monospace;font-size:13px}
  .row.me .bubble code{background:rgba(255,255,255,.2)}
  .next-actions{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}
  .next-actions button{border:1px solid var(--b);background:var(--w);color:var(--n);border-radius:20px;padding:7px 12px;font-size:12px;font-family:inherit;cursor:pointer;white-space:nowrap;transition:.15s}
  .next-actions button:hover{border-color:var(--t);color:var(--t);background:rgba(196,97,74,.08)}
  .next-actions button:first-child{background:var(--t);color:#fff;border-color:var(--t);font-weight:bold}
  .next-actions button:first-child:hover{background:#A84026;border-color:#A84026}
  .typing{font-size:13px;color:var(--u);font-style:italic;padding:0 22px 8px}
  footer{border-top:1px solid var(--b);background:var(--c);padding:16px 22px}
  .inbar{display:flex;gap:12px;align-items:flex-end;max-width:1000px;margin:0 auto}
  #inp{flex:1;border:1px solid var(--b);border-radius:22px;padding:13px 20px;font-size:15px;font-family:inherit;resize:none;max-height:160px;outline:none;background:var(--w)}
  #inp:focus{border-color:var(--t)}
  #fileBtn{width:46px;height:46px;border-radius:50%;border:1px solid var(--b);background:var(--w);color:var(--t);font-size:20px;cursor:pointer;flex-shrink:0}
  #fileName{max-width:260px;color:var(--u);font-size:12px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
  #send{width:46px;height:46px;border-radius:50%;border:none;background:var(--t);color:#fff;font-size:18px;cursor:pointer;flex-shrink:0}
  #send:disabled{opacity:.4;cursor:default}
  .hint{text-align:center;font-size:11px;color:var(--u);margin-top:8px}
  .docLinkRow{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}
  .docLinkBtn{border:1px solid var(--b);background:var(--w);color:var(--n);border-radius:8px;padding:7px 12px;
    font-size:12px;font-family:inherit;cursor:pointer;text-decoration:none;display:inline-flex;align-items:center;gap:6px}
  .docLinkBtn.primary{background:var(--t);border-color:var(--t);color:#fff;font-weight:bold}
  .docLinkBtn:hover{filter:brightness(.98);border-color:var(--t)}

  /* Document modal styles */
  #docModal{display:none;position:fixed;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,.5);z-index:100;align-items:center;justify-content:center}
  #docModal.show{display:flex}
  .docModalContent{background:var(--c);border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.3);width:90%;max-width:900px;max-height:85vh;overflow-y:auto;padding:24px}
  .docModalContent .header{display:flex;justify-content:space-between;align-items:center;margin-bottom:20px;border-bottom:1px solid var(--b);padding-bottom:16px}
  .docModalContent .header h2{margin:0;color:var(--n);font-size:20px}
  .docModalContent .header .close{background:none;border:none;font-size:24px;cursor:pointer;color:var(--u);transition:.15s}
  .docModalContent .header .close:hover{color:var(--t)}
  .docStage{margin-bottom:20px;padding:14px;border:1px solid var(--b);border-radius:8px;background:var(--w)}
  .docStage .agent{font-weight:bold;color:var(--t);font-size:14px}
  .docStage .verdict{display:inline-block;margin-left:10px;font-size:11px;padding:2px 6px;border-radius:4px}
  .docStage .verdict.ready_next{background:#E3F0E6;color:#2C5F3A}
  .docStage .verdict.needs_revision{background:#FFF3E0;color:#E65100}
  .docStage .verdict.done{background:#E1F5FE;color:#01579B}
  .docStage .time{font-size:11px;color:#999;margin-top:4px}
  .docStage .content{margin-top:8px;font-size:13px;line-height:1.5;color:var(--n);white-space:pre-wrap;word-wrap:break-word}
  .docActions{display:flex;gap:10px;margin-top:20px;padding-top:16px;border-top:1px solid var(--b);flex-wrap:wrap}
  .docActions button{border:1px solid var(--b);background:var(--w);color:var(--n);border-radius:8px;padding:8px 14px;font-size:13px;font-family:inherit;cursor:pointer;transition:.15s}
  .docActions button:hover{border-color:var(--t);color:var(--t);background:rgba(196,97,74,.08)}
  .docActions button.primary{background:var(--t);color:#fff;border-color:var(--t);font-weight:bold}
  .docActions button.primary:hover{background:#A84026;border-color:#A84026}
  .feedbackBox{margin-top:12px;padding:12px;background:#F5F5F5;border-radius:6px;display:none}
  .feedbackBox.show{display:block}
  .feedbackBox textarea{width:100%;min-height:80px;padding:8px;border:1px solid var(--b);border-radius:4px;font-family:inherit;font-size:12px;resize:vertical}
  .feedbackBox .actions{display:flex;gap:8px;margin-top:8px}
  .feedbackBox .actions button{font-size:12px;padding:6px 12px}
</style>
</head>
<body>
  <div id="side"><div class="logo">MILA<br>OFFICE</div></div>
  <div id="main">
    <header>
      <div class="av" id="hav">M</div>
      <div class="ttl"><div class="nm" id="hname">…</div><div class="sub" id="hsub">@liudmyla.lykova · всегда онлайн</div></div>
      <button class="toggle" id="resetBtn">Очистить чат</button>
      <button class="toggle" id="resetSessBtn" title="Очистить переписку со всеми агентами">Очистить сессию</button>
    </header>
    <div class="chips" id="chips"></div>
    <div id="chat"></div>
    <div class="typing" id="typing" style="display:none"></div>
    <footer>
      <div class="inbar">
        <input id="fileInp" type="file" accept=".txt,.md,.csv,.json,.docx,.pdf,image/*" style="display:none">
        <button id="fileBtn" title="Прикрепить файл">📎</button>
        <div id="fileName"></div>
        <textarea id="inp" rows="1" placeholder="Напиши сообщение…"></textarea>
        <button id="send" title="Отправить">➤</button>
      </div>
      <div class="hint">Enter — отправить · Shift+Enter — новая строка</div>
    </footer>
  </div>

  <!-- Download popup (when Victoria approves - VERDICT: done) -->
  <div id="downloadPopup" style="display:none;position:fixed;top:50%;left:50%;transform:translate(-50%,-50%);background:#fff;border-radius:12px;box-shadow:0 20px 60px rgba(0,0,0,.3);padding:32px;z-index:200;max-width:400px;text-align:center">
    <div style="font-size:48px;margin-bottom:16px">✓</div>
    <h2 style="margin:0 0 8px;color:#2C1A12;font-size:20px">Воркбук одобрен!</h2>
    <p style="margin:0 0 24px;color:#7A5E54;font-size:14px">Финальная версия готова к скачиванию</p>
    <button onclick="downloadWorkbookTXT()" style="display:block;width:100%;background:#C46148;color:#fff;border:none;border-radius:8px;padding:12px;font-size:14px;font-family:inherit;cursor:pointer;margin-bottom:8px;font-weight:bold">Скачать TXT</button>
    <button onclick="closeDownloadPopup()" style="display:block;width:100%;background:#E0D0C8;color:#2C1A12;border:none;border-radius:8px;padding:12px;font-size:14px;font-family:inherit;cursor:pointer">Закрыть</button>
  </div>
  <div id="downloadOverlay" onclick="closeDownloadPopup()" style="display:none;position:fixed;top:0;left:0;width:100%;height:100%;background:rgba(0,0,0,.5);z-index:199"></div>

  <!-- Document detail modal -->
  <div id="docModal">
    <div class="docModalContent">
      <div class="header">
        <h2 id="docTitle">Документ</h2>
        <button class="close" onclick="closeDocModal()">✕</button>
      </div>
      <div id="docDetails"></div>
      <div id="feedbackChain"></div>
      <div class="docActions">
        <button class="primary" id="docFeedbackBtn" onclick="toggleFeedbackBox()">💬 Отправить правки назад</button>
        <button id="docArchiveBtn" onclick="archiveDoc()">📦 Архивировать</button>
        <button id="docExportBtn" onclick="exportDoc()">⬇️ Скачать историю</button>
      </div>
      <div class="feedbackBox" id="feedbackBox">
        <textarea id="feedbackText" placeholder="Опиши правки, которые нужны: что переписать, что добавить/убрать, какие исправления…"></textarea>
        <div class="actions">
          <button onclick="sendFeedback()" class="primary">Отправить</button>
          <button onclick="toggleFeedbackBox()">Отмена</button>
        </div>
      </div>
    </div>
  </div>
<script>
let AGENTS=[], cur=null, CSRF='', activeJob=null, pendingUpload=null, currentDocId=null, activeLoadAgent=null;

async function postJSON(url, payload){
  const body=JSON.stringify(payload||{});
  let r=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body});
  // 403 = протухший CSRF (напр. сервер перезапускали). Обновляем токен и пробуем ещё раз.
  if(r.status===403){
    try{ const m=await (await fetch('/api/meta')).json(); CSRF=m.csrf; }catch(e){}
    r=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body});
  }
  return r;
}

function esc(s){return String(s==null?'':s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}

// ─── Download & Modal Functions (must be defined early) ───
window.showDownloadPopup = function(content){
  document.getElementById('downloadPopup').style.display='block';
  document.getElementById('downloadOverlay').style.display='block';
  window.pendingDownloadContent = content;
};

window.closeDownloadPopup = function(){
  document.getElementById('downloadPopup').style.display='none';
  document.getElementById('downloadOverlay').style.display='none';
  window.pendingDownloadContent = null;
};

window.downloadWorkbookTXT = async function(){
  console.log('🔍 downloadWorkbookTXT вызвана');
  // Если документ зарегистрирован на сервере — скачиваем ЧИСТУЮ финальную версию
  // (с применёнными правками), а не обёртку вокруг текста реплики.
  if(currentDocId){
    try{
      const resp=await fetch('/api/document/'+encodeURIComponent(currentDocId)+'/download');
      if(resp.ok){
        const blob=await resp.blob();
        const cd=resp.headers.get('Content-Disposition')||'';
        const m=cd.match(/filename\*?=(?:UTF-8'')?"?([^\";]+)"?/i);
        const filename=m?decodeURIComponent(m[1]):('mila-готовый-'+currentDocId+'.txt');
        const url=URL.createObjectURL(blob);
        const a=document.createElement('a');
        a.href=url; a.download=filename; document.body.appendChild(a);
        a.click(); document.body.removeChild(a); URL.revokeObjectURL(url);
        setTimeout(()=>closeDownloadPopup(),300);
        return;
      }
    }catch(e){ console.warn('Серверное скачивание не удалось, фолбэк на текст реплики', e); }
  }
  console.log('📦 pendingDownloadContent:', window.pendingDownloadContent?.substring(0,50));

  if(!window.pendingDownloadContent) {
    const bubbles=Array.from(document.querySelectorAll('.bubble,.docStage,.card'));
    const source=bubbles.reverse().find(el=>{
      const t=el.innerText||'';
      return t.includes('ФИНАЛЬНЫЙ ЧЕК-ЛИСТ') || t.includes('Готово к загрузке') || t.includes('Скачать воркбук');
    }) || document.querySelector('.bubble:last-child');
    window.pendingDownloadContent = source ? (source.innerText||'').replace(/Скачать воркбук\s*\(TXT\)/g,'').trim() : '';
    console.log('📍 Извлечено из DOM:', window.pendingDownloadContent?.substring(0,50));
  }
  if(!window.pendingDownloadContent) {
    window.pendingDownloadContent = 'Финальный чек-лист перед публикацией\n\nДокумент готов к загрузке в GAMMA.';
  }
  if(!window.pendingDownloadContent) {
    alert('Нет содержимого для скачивания');
    return;
  }

  try {
    const timestamp = new Date().toLocaleString('ru-RU');
    const filename = 'workbook_approved_' + new Date().toISOString().split('T')[0] + '.txt';
    const txtContent = `════════════════════════════════════════════════════════════
ОДОБРЕННЫЙ ВОРКБУК
════════════════════════════════════════════════════════════

Дата одобрения: ${timestamp}
Статус: ✓ ОДОБРЕН (VERDICT: done)
Агент: Victoria (редактор)

════════════════════════════════════════════════════════════

СОДЕРЖИМОЕ:

${window.pendingDownloadContent}

════════════════════════════════════════════════════════════

Готово к:
✓ Загрузке в GAMMA (gamma.app)
✓ Отправке Marina для переформатирования
✓ Печати и распространению

════════════════════════════════════════════════════════════`;

    const blob = new Blob([txtContent], {type: 'text/plain; charset=utf-8'});
    const url = URL.createObjectURL(blob);
    const link = document.createElement('a');
    link.href = url;
    link.download = filename;
    link.style.display = 'none';
    document.body.appendChild(link);
    console.log('✅ Запускаю скачивание:', filename);
    link.click();
    document.body.removeChild(link);
    URL.revokeObjectURL(url);
    alert('✅ Файл скачан: ' + filename);
    setTimeout(() => closeDownloadPopup(), 500);
  } catch(e) {
    console.error('❌ Ошибка при скачивании:', e);
    alert('Ошибка при скачивании: ' + e.message);
  }
};

window.openDocModal = async function(docId){
  const safe=String(docId||currentDocId||'').replace(/[^A-Za-z0-9_-]/g,'');
  if(!safe){ alert('История документа пока недоступна: документ не выбран.'); return; }
  try{
    const resp=await fetch('/api/document/'+encodeURIComponent(safe));
    if(!resp.ok) throw new Error('Document not found');
    const data=await resp.json();
    if(!data.ok) throw new Error(data.error||'Error loading document');
    const doc=data.document;
    currentDocId=safe;
    saveSessionToStorage();
    document.getElementById('docTitle').textContent='📄 '+esc(doc.file_name);
    let html='<div style="margin-bottom:16px"><strong>Создан:</strong> '+esc(new Date(doc.created_at).toLocaleString('ru-RU'))+'<br><strong>Статус:</strong> '+esc(doc.status)+'</div>';
    html+='<div style="margin-bottom:16px;padding:12px;background:#F5F5F5;border-radius:6px;font-size:12px"><strong>Исходный материал:</strong><br>'+esc(doc.original_content||'[нет содержимого]')+'</div>';
    html+='<h3 style="margin:20px 0 12px;font-size:16px">История обработки</h3>';
    (doc.stages||[]).forEach((stage,idx)=>{
      const names={'victoria':'Виктория','rita':'Рита','marina':'Марина','lera':'Лера','producer':'Продюсер','manager':'Менеджер','vasya':'Вася','dima':'Дима','tyoma':'Тёма','olya':'Оля','alina':'Алина'};
      html+='<div class="docStage"><div class="agent">'+idx+'. '+esc(names[stage.agent]||stage.agent)+'<span class="verdict '+esc(stage.verdict)+'">'+esc(stage.verdict)+'</span></div>';
      html+='<div class="time">'+esc(new Date(stage.timestamp).toLocaleString('ru-RU'))+'</div>';
      if(stage.input) html+='<div style="margin-top:8px"><strong>Исходный текст:</strong><div class="content">'+esc(stage.input)+'</div></div>';
      if(stage.output) html+='<div style="margin-top:8px"><strong>Результат:</strong><div class="content">'+esc(stage.output)+'</div></div>';
      html+='</div>';
    });
    document.getElementById('docDetails').innerHTML=html;
    if(doc.feedback_chain && doc.feedback_chain.length){
      let fbHtml='<h3 style="margin:20px 0 12px;font-size:16px">Правки и комментарии</h3>';
      (doc.feedback_chain||[]).forEach((fb)=>{
        const names={'victoria':'Виктория','rita':'Рита','marina':'Марина','lera':'Лера','producer':'Продюсер','manager':'Менеджер','vasya':'Вася','dima':'Дима','tyoma':'Тёма','olya':'Оля','alina':'Алина'};
        fbHtml+='<div class="docStage"><div class="agent">'+esc(names[fb.from_agent]||fb.from_agent)+' → '+esc(names[fb.to_agent]||fb.to_agent)+'</div>';
        fbHtml+='<div class="time">'+esc(new Date(fb.timestamp).toLocaleString('ru-RU'))+'</div>';
        fbHtml+='<div class="content">'+esc(fb.feedback)+'</div></div>';
      });
      document.getElementById('feedbackChain').innerHTML=fbHtml;
    }
    document.getElementById('docModal').classList.add('show');
  }catch(e){
    alert('Ошибка: '+e.message);
  }
};

// #RRGGBB + alpha → rgba(...) для полупрозрачного фона пилюль-команд в тултипе.
function hexA(hex,a){const h=(hex||'#888').replace('#','');
  const r=parseInt(h.substr(0,2),16),g=parseInt(h.substr(2,2),16),b=parseInt(h.substr(4,2),16);
  return 'rgba('+r+','+g+','+b+','+a+')';}
function md(s){
  s=esc(s);
  s=s.replace(/\[doc_id:([A-Za-z0-9_-]{4,80})\]/g,function(_,id){return docButtons(id);});
  s=s.replace(/(^|\s)\/api\/document\/([A-Za-z0-9_-]{4,80})(?!\/[A-Za-z])/g,function(_,lead,id){return lead+docButtons(id);});
  s=s.replace(/\*\*([^*]+)\*\*/g,'<b>$1</b>');
  s=s.replace(/`([^`]+)`/g,'<code>$1</code>');
  s=s.replace(/(^|\n)\s*[-•]\s+/g,'$1• ');
  return s;
}
function docButtons(id){
  const safe=String(id||'').replace(/[^A-Za-z0-9_-]/g,'');
  if(!safe) return '';
  return '<span class="docLinkRow" data-doc-id="'+safe+'">'
    +'<button class="docLinkBtn primary" onclick="downloadDoc(event,\''+safe+'\')">Скачать документ</button>'
    +'<button class="docLinkBtn" onclick="openDocModal(\''+safe+'\')">История</button>'
    +'</span>';
}
function agent(){return AGENTS.find(a=>a.key===cur);}

const NEXT_ACTIONS={
  victoria:[
    {label:'Передать Рите',agent:'rita',prompt:'Возьми отзыв редактора и преврати его в конкретные правки структуры рабочей тетради: что переписать, что сократить, что добавить и какой порядок блоков сделать.',desc:'Структурные правки и переписка'},
    {label:'Упаковать с Мариной',agent:'marina',prompt:'На основе редакторского отзыва сформулируй упаковку продукта: кому нужна рабочая тетрадь, какую боль закрывает, 3-5 выгод и позиционирование.',desc:'Маркетинговая позиция и CTA'},
    {label:'План запуска',agent:'vasya',prompt:'Собери следующий план действий по рабочей тетради: кто что делает, сроки, порядок финализации и публикации.',desc:'Сроки и распределение задач'}
  ],
  rita:[
    {label:'Финальная редактура',agent:'victoria',prompt:'Проверь обновленную структуру рабочей тетради как финальный редактор: голос, ясность, CTA, лишнее и недостающее.',desc:'Финальная проверка голоса и ясности'},
    {label:'Оффер',agent:'lera',prompt:'Собери оффер для рабочей тетради: обещание, цена/ценность, кому подходит и почему купить сейчас.',desc:'Позиция и цена продукта'},
    {label:'Упаковка',agent:'marina',prompt:'Сформулируй публичную упаковку рабочей тетради для Instagram и описания продукта.',desc:'Маркетинговый текст для соцсетей'}
  ],
  marina:[
    {label:'Продажный текст',agent:'lera',prompt:'На основе упаковки продукта напиши продающий текст для запуска рабочей тетради.',desc:'Сильное торговое предложение'},
    {label:'Контент-план',agent:'vasya',prompt:'Разложи запуск рабочей тетради в календарь: посты, сторис, напоминания и дедлайны.',desc:'Расписание публикаций'},
    {label:'Проверка продюсера',agent:'producer',prompt:'Оцени упаковку рабочей тетради как продюсер: место в линейке, цена, следующий продукт и риски.',desc:'Стратегическая оценка проекта'}
  ],
  lera:[
    {label:'Проверка Кирилла',agent:'producer',prompt:'Проверь оффер и продажный текст рабочей тетради: сила предложения, цена, путь к консультации или следующему продукту.',desc:'Стратегическая оценка'},
    {label:'План публикаций',agent:'vasya',prompt:'Составь план публикаций и сторис для запуска по этому продажному тексту.',desc:'Расписание постов'},
    {label:'Редактура текста',agent:'victoria',prompt:'Отредактируй продажный текст: голос Людмилы, ясность, мягкость, CTA.',desc:'Проверка тона и ясности'}
  ],
  producer:[
    {label:'Поставить задачи',agent:'manager',prompt:'Разложи решение продюсера по рабочей тетради в задачи офиса: ответственный, следующий шаг, критерий готовности.',desc:'Фрагментация по задачам'},
    {label:'Календарь',agent:'vasya',prompt:'Поставь продюсерский план запуска рабочей тетради в календарь действий.',desc:'Сроки и вехи'},
    {label:'Финансы',agent:'dima',prompt:'Посчитай экономику рабочей тетради: цена, цель продаж, сценарии выручки, конверсия.',desc:'Расчет экономики'}
  ],
  manager:[
    {label:'Назначить задачи',agent:'manager',prompt:'Сделай короткий task-list для офиса по этому проекту: агент, задача, вход, выход, дедлайн.',desc:'Уточнение задач'},
    {label:'Календарь',agent:'vasya',prompt:'Преврати этот task-list в календарь выполнения.',desc:'Расписание выполнения'},
    {label:'Продолжить с Мариной',agent:'marina',prompt:'Возьми план офиса и подготовь следующий маркетинговый шаг.',desc:'Маркетинговый следующий шаг'}
  ],
  vasya:[
    {label:'Развернуть в задачи',agent:'manager',prompt:'Разложи этот план в задачи для офиса: кто, что, когда, критерий готовности.',desc:'Задачи с владельцами'},
    {label:'Согласовать с продюсером',agent:'producer',prompt:'Проверь этот календарь как продюсер: реалистичность, приоритеты, зависимости, риски.',desc:'Валидация плана'},
    {label:'Опубликовать в Telegram',agent:'tyoma',prompt:'Анонсируй этот план запуска в Telegram для команды и Людмилы.',desc:'Анонс для команды'}
  ],
  dima:[
    {label:'Одобрить план',agent:'producer',prompt:'На основе финансовых расчётов помоги принять решение: стоит ли идти в этот проект, цена правильная, риски.',desc:'GO/NO-GO решение'},
    {label:'Задачи по монетизации',agent:'manager',prompt:'Преврати финансовый план в конкретные задачи офиса: цена, условия, система продаж.',desc:'Операционные задачи'},
    {label:'Отчёт для Людмилы',agent:'marina',prompt:'Подготовь финансовый отчёт для Людмилы: инвестиции, цель прибыли, сценарии.',desc:'Финансовый отчет'}
  ],
  tyoma:[
    {label:'Проверить тон',agent:'victoria',prompt:'Отредактируй сообщение для Telegram: голос Людмилы, ясность, CTA.',desc:'Редактура тона'},
    {label:'Расширить текст',agent:'marina',prompt:'Расширь это сообщение для Telegram: добавь контекст, эмоцию, привлекательность.',desc:'Развитие текста'},
    {label:'Согласовать с планом',agent:'vasya',prompt:'Проверь, соответствует ли объявление в Telegram календарному плану и срокам.',desc:'Синхро со сроками'}
  ],
  olya:[
    {label:'Контент по тренду',agent:'marina',prompt:'Создай контент по этому тренду: формат, позиционирование, как вписать Людмилу и её услуги.',desc:'Создание контента'},
    {label:'Регулярный пост',agent:'marina',prompt:'Превратить этот тренд в регулярный контент для Instagram: серия постов, карусель или Reels.',desc:'Цикличный контент'},
    {label:'Добавить в план',agent:'vasya',prompt:'Запланируй контент по этому тренду в календарь и установи сроки публикации.',desc:'Планирование'}
  ],
  alina:[
    {label:'Встреча со своим стилем',agent:'marina',prompt:'Подготовь контент для лида, который интересуется услугами Людмилы: позиционирование, путь к консультации.',desc:'Контент для лида'},
    {label:'Письмо льду',agent:'lera',prompt:'Напиши письмо потенциальному клиенту на основе данных контакта: персонализация, оффер, CTA.',desc:'Персонализированное письмо'},
    {label:'Follow-up в Telegram',agent:'tyoma',prompt:'Составь follow-up сообщение для лида в Telegram с предложением консультации.',desc:'Follow-up в мессенджер'}
  ],
  default:[
    {label:'Что дальше',agent:'manager',prompt:'Определи следующий лучший шаг офиса по этому результату: кто должен продолжить, что сделать и какой результат получить.'},
    {label:'План действий',agent:'vasya',prompt:'Составь конкретный план следующих действий по этому результату.'}
  ]
};

function showAutoSuggestion(fromAgent, nextAction){
  // Показываем явное предложение перейти к следующему агенту
  const a=AGENTS.find(x=>x.key===fromAgent);
  const b=AGENTS.find(x=>x.key===nextAction.agent);
  if(!a||!b) return;

  setTimeout(()=>{
    const suggestion=document.createElement('div');
    suggestion.style.cssText='position:fixed;bottom:100px;right:20px;background:linear-gradient(135deg,'+a.color+' 0%,'+b.color+' 100%);color:#fff;padding:16px 20px;border-radius:12px;box-shadow:0 8px 24px rgba(0,0,0,.2);z-index:45;max-width:280px;font-size:13px;line-height:1.5;cursor:pointer;transition:transform 0.2s';
    suggestion.innerHTML=`
<div style="font-weight:bold;margin-bottom:8px">💡 Рекомендуемый следующий шаг</div>
<div style="margin-bottom:12px">${a.emoji} ${a.name} → ${b.emoji} <strong>${nextAction.label}</strong></div>
<div style="font-size:11px;opacity:0.9;margin-bottom:8px">${nextAction.desc||''}</div>
<button onclick="this.parentElement.remove()" style="background:rgba(255,255,255,.2);color:#fff;border:1px solid rgba(255,255,255,.3);padding:6px 12px;border-radius:6px;font-size:11px;cursor:pointer;width:100%">Использовать →</button>
    `;
    suggestion.onclick=(e)=>{ if(e.target.tagName==='BUTTON') runNextAction(nextAction,''); };
    document.body.appendChild(suggestion);

    // Автоматически удаляем через 10 сек
    setTimeout(()=>{ if(suggestion.parentElement) suggestion.remove(); }, 10000);
  }, 400);
}

function nextActionsFor(agentKey, verdict){
  // Условные actions: если needs_revision — вернуться к предыдущему, иначе обычные next actions
  if(verdict==='needs_revision'){
    // Для needs_revision показываем кнопку "вернуть с комментарием"
    const actions=(NEXT_ACTIONS[agentKey]||NEXT_ACTIONS.default).slice(0,3);
    return [{label:'Отправить с комментариями обратно',agent:agentKey,prompt:'Дай детальный комментарий к предыдущему этапу и предложи конкретные исправления.',desc:'Вернуться с правками'},...actions];
  }
  const actions = (NEXT_ACTIONS[agentKey]||NEXT_ACTIONS.default).slice(0,3);

  // Для ready_next добавляем явное предложение выбрать рекомендуемый агент
  if(verdict==='ready_next' && actions.length > 0){
    showAutoSuggestion(agentKey, actions[0]);
  }

  return actions;
}

// UI-переписка по агенту: {agentKey: [{text, me}, ...]}. Бэкенд хранит свою
// историю (для контекста модели), а это — то, что видно на экране. Переключение
// агентов больше НЕ стирает переписку: для каждого реплеим её из этого стора.
const TRANSCRIPTS = {};

// Рисует один пузырь в DOM (без записи в стор).
function drawMsg(text, me, actions, verdict){
  text = (text && typeof text === 'object') ? (text.text || text.content || text.message || JSON.stringify(text)) : String(text==null?'':text);
  actions = Array.isArray(actions) ? actions : [];
  const chat=document.getElementById('chat');
  const a=agent();
  const row=document.createElement('div'); row.className='row'+(me?' me':'');
  const av=document.createElement('div'); av.className='av';
  av.style.background=me?'#7A5E54':a.color; av.textContent=me?'Я':a.emoji;

  // Убираем [→ agent] подсказку из видимого текста, но парсим её для выделения.
  let displayText=text;
  let recommendedAgent=null;
  const recommended=text.match(/\[→\s*(\w+)\]$/);
  if(recommended){
    recommendedAgent=recommended[1];
    displayText=text.replace(/\s*\[→\s*\w+\]\s*$/, '');
  }

  const b=document.createElement('div'); b.className='bubble'; b.innerHTML=md(displayText);

  // Показываем verdict статус для agent-ответов
  if(!me && verdict){
    const vBadge=document.createElement('div');
    vBadge.style.cssText='margin-top:8px;font-size:11px;padding:4px 8px;border-radius:6px;display:inline-block;';
    const vColors={ready_next:'#E3F0E6;color:#2C5F3A',needs_revision:'#FFF3E0;color:#E65100',done:'#E1F5FE;color:#01579B'};
    const vText={ready_next:'✓ Готово для передачи',needs_revision:'⚠ Нужны правки',done:'✓ Завершено'};
    vBadge.style.background=vColors[verdict]||vColors.ready_next;
    vBadge.textContent=vText[verdict]||'Статус неизвестен';
    b.appendChild(vBadge);
  }

  if(!me && actions && actions.length){
    const hint=document.createElement('div');
    hint.style.cssText='margin-top:12px;margin-bottom:6px;font-size:11px;color:#7A5E54;text-transform:uppercase;letter-spacing:0.5px;font-weight:bold';
    hint.textContent='📋 Следующий шаг:';
    b.appendChild(hint);

    const wrap=document.createElement('div'); wrap.className='next-actions';
    wrap.style.cssText='display:flex;flex-wrap:wrap;gap:6px;margin-top:8px';
    actions.forEach((act, idx)=>{
      const btn=document.createElement('button'); btn.type='button'; btn.textContent=act.label;
      if(act.desc) btn.title=act.desc;
      if(idx===0){
        btn.style.cssText='background:var(--t);color:#fff;border:none;border-radius:6px;padding:8px 12px;font-size:13px;cursor:pointer;font-weight:bold;flex:1;min-width:140px';
      }else{
        btn.style.cssText='background:#E0D0C8;color:var(--n);border:none;border-radius:6px;padding:8px 12px;font-size:12px;cursor:pointer;flex:0.8;min-width:100px';
      }
      if(recommendedAgent && act.agent===recommendedAgent) {
        btn.style.outline='2px solid var(--t)'; btn.style.outlineOffset='2px';
      }
      btn.onclick=()=>runNextAction(act, displayText);
      wrap.appendChild(btn);
    });
    b.appendChild(wrap);
  }
  row.appendChild(av); row.appendChild(b); chat.appendChild(row);
  chat.scrollTop=chat.scrollHeight;
}

// Добавляет сообщение и в стор текущего агента, и на экран.
// Очищает [→ agent] подсказку перед сохранением.
function addMsg(text, me, actions, verdict){
  text = (text && typeof text === 'object') ? (text.text || text.content || text.message || JSON.stringify(text)) : String(text==null?'':text);
  const savedActions=actions||[];
  const cleanText=text.replace(/\s*\[→\s*\w+\]\s*$/, '');

  // Валидация
  if (!cleanText || !cur) {
    console.warn('Invalid message or agent');
    return;
  }

  (TRANSCRIPTS[cur] = TRANSCRIPTS[cur] || []).push({text:cleanText, me, actions:savedActions, verdict});
  drawMsg(text, me, savedActions, verdict);
  saveSessionToStorage();  // Сохраняем в localStorage

  // Если Victoria одобрила (VERDICT: done) - показываем popup для скачивания
  if (verdict === 'done' && cur === 'victoria') {
    setTimeout(() => showDownloadPopup(cleanText), 500);
  }

  // Сохраняем сообщение на сервер в истории агента (с retry logic)
  saveToServerWithRetry(cur, cleanText, me, verdict);
}

// Helper: retry logic для сохранения на сервер
async function saveToServerWithRetry(agent, text, isUser, verdict, attempt = 1) {
  const maxAttempts = 3;
  const delayMs = 1000 * attempt;  // exponential backoff: 1s, 2s, 4s

  try {
    await postJSON('/api/agent-message', {
      agent: agent,
      text: text,
      is_user: isUser,
      verdict: verdict || null
    });
    // Success - no action needed
  } catch (e) {
    if (attempt < maxAttempts) {
      console.warn(`Save to server failed (attempt ${attempt}), retrying in ${delayMs}ms...`);
      setTimeout(() => saveToServerWithRetry(agent, text, isUser, verdict, attempt + 1), delayMs);
    } else {
      console.error(`Could not save to server after ${maxAttempts} attempts:`, e);
      // Message is still in localStorage, so it's not lost
      // But warn user
      const msg = document.createElement('div');
      msg.style.cssText = 'position:fixed;bottom:80px;right:20px;background:#f6e1dc;color:#a8412c;' +
        'padding:10px 14px;border-radius:6px;font-size:12px;z-index:50;max-width:300px';
      msg.textContent = 'Не удалось сохранить на сервер (сообщение в памяти браузера)';
      document.body.appendChild(msg);
      setTimeout(() => msg.remove(), 5000);
    }
  }
}

async function runNextAction(action, context){
  if(activeJob) return;
  if(action.agent && action.agent!==cur) switchAgent(action.agent);
  const inp=document.getElementById('inp');
  let fullPrompt=action.prompt+'\n\nКонтекст предыдущего шага:\n'+context;
  // Если есть currentDocId (загруженный документ), передаём ссылку на его историю
  if(currentDocId){
    fullPrompt+='\n\n[doc_id:'+currentDocId+']\nЭтот материал уже прошёл обработку, см. /api/document/'+currentDocId;
  }
  inp.value=fullPrompt;
  inp.style.height='auto';
  inp.style.height=Math.min(inp.scrollHeight,160)+'px';
  inp.focus();
  send();
}

async function renderAgent(){
  // Prevent race conditions: don't load if another agent is already loading
  if (activeLoadAgent && activeLoadAgent !== cur) {
    console.warn('Render skipped: another agent is loading');
    return;
  }

  activeLoadAgent = cur;

  try {
    const a = agent();
    let titleText = a.name + ' — ' + a.role;
    if(currentDocId) titleText += ' · 📄 ' + currentDocId.substring(0,6) + '...';
    document.getElementById('hname').textContent = titleText;

    // Показываем кнопку просмотра документа если есть currentDocId
    const docBtn = document.getElementById('docViewBtn');
    if(docBtn) {
      if(currentDocId) {
        docBtn.style.display = 'block';
        docBtn.onclick = () => openDocModal(currentDocId);
      } else {
        docBtn.style.display = 'none';
      }
    }

    const hav = document.getElementById('hav');
    hav.textContent = a.emoji;
    hav.style.background = a.color;
    document.getElementById('inp').placeholder = 'Спроси ' + a.name + '…';
    const ch = document.getElementById('chips');
    ch.innerHTML = '';
    a.chips.forEach(c => {
      const el = document.createElement('button');
      el.className = 'chip';
      el.textContent = c.label;
      el.title = c.prompt;
      el.onclick = () => { document.getElementById('inp').value = c.prompt; send(); };
      ch.appendChild(el);
    });
    document.querySelectorAll('.apill').forEach(p => p.classList.toggle('active', p.dataset.k === cur));
    const chat = document.getElementById('chat');
    chat.innerHTML = '';

    // Проверяем есть ли история на сервере
    let hist = TRANSCRIPTS[cur];
    if (!hist || !hist.length) {
      try {
        const resp = await fetch('/api/agent-history/' + encodeURIComponent(cur));
        if (resp.ok) {
          const data = await resp.json();
          if (data.ok && data.history && data.history.messages && data.history.messages.length > 0) {
            // Загружаем историю с сервера и заполняем TRANSCRIPTS
            const serverMsgs = data.history.messages.map(m => ({
              text: m.text,
              me: m.role === 'user',
              actions: [],
              verdict: m.verdict
            }));
            TRANSCRIPTS[cur] = serverMsgs;
            hist = serverMsgs;
          }
        }
      } catch (e) {
        console.warn('Could not load server history:', e);
      }
    }

    if(hist && !Array.isArray(hist)) {
      hist = Object.values(hist).filter(Boolean);
      TRANSCRIPTS[cur] = hist;
    }
    if (hist && hist.length) {
      hist.forEach(m => {
        try{
          const item = (m && typeof m === 'object') ? m : {text:String(m||''), me:false};
          drawMsg(item.text, !!item.me, item.actions||[], item.verdict);
        }catch(e){
          console.error('Could not render restored message:', e, m);
        }
      });
    } else {
      drawMsg(a.intro, false);
    }
  } catch(e) {
    console.error('renderAgent failed:', e);
    const chat = document.getElementById('chat');
    if(chat){
      chat.innerHTML = '';
      const row=document.createElement('div'); row.className='row';
      const av=document.createElement('div'); av.className='av'; av.textContent='!';
      av.style.background='#A8412C';
      const b=document.createElement('div'); b.className='bubble';
      b.innerHTML='Не удалось восстановить старую переписку. <button class="docLinkBtn primary" onclick="resetChat()">Очистить этот чат</button>';
      row.appendChild(av); row.appendChild(b); chat.appendChild(row);
    }
  } finally {
    activeLoadAgent = null;
  }
}

async function switchAgent(k){ cur=k; await renderAgent(); saveSessionToStorage(); }

async function uploadSelectedFile(file){
  if(!file) return;
  const label=document.getElementById('fileName');
  label.textContent='Файл загружается';
  const fd=new FormData(); fd.append('file', file);
  let r=await fetch('/api/upload',{method:'POST',headers:{'X-CSRF-Token':CSRF},body:fd});
  if(r.status===403){
    try{ const m=await (await fetch('/api/meta')).json(); CSRF=m.csrf; }catch(e){}
    r=await fetch('/api/upload',{method:'POST',headers:{'X-CSRF-Token':CSRF},body:fd});
  }
  const d=await r.json();
  if(!r.ok || !d.ok){
    pendingUpload=null;
    label.textContent=d.error||'Файл не загрузился';
    return;
  }
  pendingUpload=d;
  label.textContent='Файл прикреплен';
}

async function handlePaste(e){
  const items=(e.clipboardData&&e.clipboardData.items)?Array.from(e.clipboardData.items):[];
  const img=items.find(it=>it.kind==='file' && it.type && it.type.startsWith('image/'));
  if(!img) return;
  const blob=img.getAsFile();
  if(!blob) return;
  e.preventDefault();
  const ext=(blob.type.split('/')[1]||'png').replace('jpeg','jpg').split(';')[0];
  const file=new File([blob], 'screenshot-'+new Date().toISOString().replace(/[:.]/g,'-')+'.'+ext, {type:blob.type||'image/png'});
  await uploadSelectedFile(file);
}

async function send(){
  const inp=document.getElementById('inp'); const text=inp.value.trim();
  if(!text && !pendingUpload) return;
  if(activeJob) return;
  const upload=pendingUpload;
  const shown=text || 'Дай фидбек по файлу';
  inp.value=''; inp.style.height='auto';
  addMsg(upload ? (shown+'\n\nПрикреплен файл') : shown,true);
  const t=document.getElementById('typing'); t.textContent=agent().name+' печатает…'; t.style.display='block';
  document.getElementById('send').disabled=true;
  try{
    const r=await postJSON('/api/chat',{agent:cur,message:text,upload_id:upload?upload.upload_id:null});
    if(!r.ok){ addMsg('⚠️ Сервер вернул '+r.status+' (попробуй обновить страницу).',false);
      t.style.display='none'; document.getElementById('send').disabled=false; return; }
    const j=await r.json();
    pendingUpload=null; document.getElementById('fileName').textContent=''; document.getElementById('fileInp').value='';
    if(j.error){ addMsg('⚠️ Ошибка: '+j.error,false); }
    else {
      // Агент думает в фоне — опрашиваем результат, пока не готов.
      const sleep=ms=>new Promise(r=>setTimeout(r,ms));
      const abort=new AbortController();
      const token={job:j.job, abort};
      activeJob=token;
      let d=null, tries=0;
      while(activeJob===token && tries<180){
        await sleep(1000);
        let rr=null;
        try{ rr=await fetch('/api/result?job='+encodeURIComponent(j.job)); d=await rr.json(); }
        catch(e){ d={error:'Сеть недоступна: '+e}; break; }
        if(d.status!=='pending') break;
        tries++;
      }
      if(activeJob!==token) return;
      activeJob=null;
      if(!d || d.status==='pending') { addMsg('Ответ ещё готовится. Обнови страницу или попробуй ещё раз через пару минут.',false); return; }
      if(d.error) addMsg('⚠️ Ошибка: '+d.error,false);
      else {
        // Сохраняем doc_id и verdict (для document workflow tracking и conditional actions)
        if(d.doc_id){
          currentDocId=d.doc_id;
          saveSessionToStorage();
        }
        const verdict=d.verdict||'ready_next';
        addMsg(d.reply,false,nextActionsFor(cur, verdict), verdict);

        // Показываем попап для скачивания если Victoria одобрила (done)
        if(verdict==='done' && cur==='victoria'){
          setTimeout(()=>{
            showDownloadPopup(d.reply);
          }, 600);
        }

        // Добавляем финальный чек-лист если Victoria одобрила (done)
        if(verdict==='done' && cur==='victoria' && currentDocId){
          setTimeout(()=>{
            // Очищаем VERDICT и другие теги из контента
            const cleanContent=d.reply.replace(/\[VERDICT:\s*\w+\]/g, '').replace(/\s*\[→\s*\w+\]\s*$/g, '').trim();
            window.pendingDownloadContent=cleanContent;
            console.log('✅ Сохранено для скачивания:', cleanContent.substring(0,100)+'...');

            const downloadBtn='<div style="margin-top:16px;display:flex;gap:8px"><button onclick="downloadWorkbookTXT()" style="background:#C46148;color:#fff;border:none;border-radius:8px;padding:12px 24px;font-size:14px;font-family:inherit;cursor:pointer;font-weight:bold;flex:1">📥 Скачать TXT</button><button onclick="openDocModal(\''+currentDocId+'\')" style="background:#4A7A5E;color:#fff;border:none;border-radius:8px;padding:12px 24px;font-size:14px;font-family:inherit;cursor:pointer;font-weight:bold;flex:1">📄 История</button></div>';
            const checklist=`
✅ ФИНАЛЬНЫЙ ЧЕК-ЛИСТ ПЕРЕД ПУБЛИКАЦИЕЙ

☑ Содержимое одобрено финальным редактором
☑ Голос и тон соответствуют бренду
☑ Все CTA работают и направляют в воронку
☑ Нет запрещённых формулировок
☑ Дизайн готов к GAMMA (или другой платформе)
☑ Метаданные заполнены (заголовок, описание, теги)
☑ Тестовая публикация пройдена
☑ Ссылка на купить/консультацию активна

🎯 Готово к загрузке в GAMMA
            `;
            const chat=document.getElementById('chat');
            const row=document.createElement('div'); row.className='row';
            const av=document.createElement('div'); av.className='av'; av.style.background=agent().color; av.textContent=agent().emoji;
            const b=document.createElement('div'); b.className='bubble'; b.innerHTML=md(checklist.trim())+downloadBtn;
            row.appendChild(av); row.appendChild(b); chat.appendChild(row);
            chat.scrollTop=chat.scrollHeight;
          }, 1200);
        }
      }
    }
  }catch(e){ addMsg('⚠️ Сеть недоступна: '+e,false); }
  t.style.display='none'; document.getElementById('send').disabled=false; inp.focus();
}

async function openDocModal(docId){
  const safe=String(docId||currentDocId||'').replace(/[^A-Za-z0-9_-]/g,'');
  if(!safe){ alert('История документа пока недоступна: документ не выбран.'); return; }
  currentDocId=safe;
  saveSessionToStorage();
  try{
    const r=await fetch('/api/document/'+encodeURIComponent(safe));
    if(!r.ok) throw new Error('Документ не найден');
    const d=await r.json();
    const doc=d.document||d;
    const lines=[];
    if(doc.file_name) lines.push('Документ: '+doc.file_name);
    if(doc.status) lines.push('Статус: '+doc.status);
    if(doc.created_at) lines.push('Создан: '+rel(doc.created_at));
    if(doc.original_content) lines.push('\\nИсходный материал:\\n'+doc.original_content);
    (doc.stages||[]).forEach((stage,idx)=>{
      lines.push('\\nШаг '+(idx+1)+': '+(stage.agent||'агент'));
      if(stage.verdict) lines.push('Вердикт: '+stage.verdict);
      if(stage.output) lines.push(stage.output);
    });
    alert(lines.join('\\n')||('История документа '+safe));
  }catch(e){
    alert('История документа '+safe+' пока недоступна.\\n'+e.message);
  }
}
window.openDocModal=openDocModal;

async function resetChat(){
  // Чистит переписку ТОЛЬКО текущего агента (UI + localStorage + server).
  if(!confirm('Очистить чат с '+agent().name+'?')) return;

  // Очищаем на сервере
  try{
    await postJSON('/api/agent-history/'+encodeURIComponent(cur)+'/clear',{});
  }catch(e){
    console.warn('Could not clear server history:',e);
  }

  // Очищаем UI
  await postJSON('/api/reset',{agent:cur});
  TRANSCRIPTS[cur]=[];
  await renderAgent();
}

async function resetSession(){
  // Чистит переписку со ВСЕМИ агентами (UI + localStorage + server history).
  if(!confirm('Очистить переписку со всеми агентами?')) return;

  // Очищаем на сервере (новые API)
  try{
    await postJSON('/api/agent-histories/clear-all',{});
  }catch(e){
    console.warn('Could not clear server history:',e);
  }

  // Очищаем UI и localStorage
  try{
    await postJSON('/api/reset',{all:true});
  }catch(e){
    console.warn('Could not reset server session:',e);
  }
  for(const k in TRANSCRIPTS) delete TRANSCRIPTS[k];
  currentDocId=null;
  localStorage.removeItem('mila_transcripts');
  localStorage.removeItem('mila_current_agent');
  localStorage.removeItem('mila_current_doc_id');
  location.reload();
}

// ─── localStorage persistence for session ───
function saveSessionToStorage(){
  // Сохраняет TRANSCRIPTS и текущего агента в localStorage
  try{
    const json = JSON.stringify(TRANSCRIPTS);
    // Проверяем примерный размер перед сохранением
    if (json.length > 8000000) {  // ~8 МБ - близко к лимиту
      console.warn('localStorage nearly full, clearing old messages');
      // Очищаем самые старые сообщения из каждого агента
      for (const agent in TRANSCRIPTS) {
        if (TRANSCRIPTS[agent] && TRANSCRIPTS[agent].length > 100) {
          TRANSCRIPTS[agent] = TRANSCRIPTS[agent].slice(-100);  // Keep last 100
        }
      }
      saveSessionToStorage();  // Retry with trimmed data
      return;
    }
    localStorage.setItem('mila_transcripts', json);
    if(cur) localStorage.setItem('mila_current_agent', cur);
    if(currentDocId) localStorage.setItem('mila_current_doc_id', currentDocId);
  }catch(e){
    if (e.name === 'QuotaExceededError') {
      console.error('localStorage is full, clearing old messages');
      // Emergency cleanup: keep only last 20 messages per agent
      for (const agent in TRANSCRIPTS) {
        if (TRANSCRIPTS[agent] && TRANSCRIPTS[agent].length > 20) {
          TRANSCRIPTS[agent] = TRANSCRIPTS[agent].slice(-20);
        }
      }
      try {
        localStorage.setItem('mila_transcripts', JSON.stringify(TRANSCRIPTS));
      } catch (e2) {
        console.error('Still cannot save to localStorage:', e2);
      }
    } else {
      console.warn('Could not save to localStorage:', e);
    }
  }
}

function loadSessionFromStorage(){
  // Загружает TRANSCRIPTS из localStorage
  try{
    const saved=localStorage.getItem('mila_transcripts');
    if(saved){
      Object.assign(TRANSCRIPTS, JSON.parse(saved));
      console.log('Session restored from localStorage');
      // Показываем уведомление что сессия восстановлена
      const msgCount=Object.values(TRANSCRIPTS).reduce((sum,msgs)=>sum+(msgs||[]).length,0);
      if(msgCount>0){
        setTimeout(()=>{
          const ch=document.getElementById('chat');
          if(ch){
            const notice=document.createElement('div');
            notice.style.cssText='text-align:center;color:#7A5E54;font-size:12px;padding:8px;margin:8px 0;border-top:1px solid #E0D0C8;border-bottom:1px solid #E0D0C8';
            notice.textContent='Сессия восстановлена ('+msgCount+' сообщений)';
            ch.insertBefore(notice, ch.firstChild);
          }
        },100);
      }
    }
    const savedAgent=localStorage.getItem('mila_current_agent');
    if(savedAgent) cur=savedAgent;
    const savedDocId=localStorage.getItem('mila_current_doc_id');
    if(savedDocId) currentDocId=savedDocId;
  }catch(e){
    console.warn('Could not load from localStorage:', e);
  }
}

window.onload=async()=>{
  const d=await (await fetch('/api/meta')).json();
  AGENTS=d.agents;
  CSRF=d.csrf;
  const side=document.getElementById('side');
  AGENTS.forEach(a=>{
    const p=document.createElement('div'); p.className='apill'; p.dataset.k=a.key;
    p.innerHTML='<div class="em">'+a.emoji+'</div><div class="nm">'+a.name+'</div>';
    p.onclick=()=>switchAgent(a.key);

    // Тултип: заголовок (emoji·имя·роль) + описание + до 3 быстрых команд.
    // Лежит в body (position:fixed) — иначе overflow сайдбара его обрежет.
    const tip=document.createElement('div'); tip.className='atip';
    tip.style.borderColor=a.color;
    const cmds=(a.chips||[]).slice(0,3).map(c=>
      '<span class="cmd" style="background:'+hexA(a.color,.15)+'">/'+
      esc((c.label||'').toLowerCase())+'</span>').join('');
    tip.innerHTML='<div class="h" style="color:'+a.color+'">'+a.emoji+' '+esc(a.name)+
      ' · '+esc(a.role)+'</div>'+
      '<div class="d">'+esc(a.tagline||'')+'</div>'+
      (cmds?'<div class="cmds">'+cmds+'</div>':'');
    document.body.appendChild(tip);

    let timer=null;
    p.addEventListener('mouseenter',()=>{
      timer=setTimeout(()=>{
        const r=p.getBoundingClientRect();
        tip.style.left=(r.right+8)+'px';
        // не вылезать за низ окна
        tip.style.top=Math.min(r.top, window.innerHeight-tip.offsetHeight-8)+'px';
        tip.classList.add('show');
      },300);
    });
    p.addEventListener('mouseleave',()=>{ clearTimeout(timer); tip.classList.remove('show'); });

    side.appendChild(p);
  });
  const sp=document.createElement('div'); sp.className='apill'; sp.title='Настройки и подключения';
  sp.innerHTML='<div class="em">⚙</div><div class="nm">Настройки</div>';
  sp.onclick=()=>window.open('/settings','_blank'); side.appendChild(sp);
  const op=document.createElement('div'); op.className='apill'; op.title='Очередь задач (оператор)';
  op.innerHTML='<div class="em">Q</div><div class="nm">Очередь</div>';
  op.onclick=()=>window.open('/operator','_blank'); side.appendChild(op);
  const inp=document.getElementById('inp');
  const fileInp=document.getElementById('fileInp');
  document.getElementById('fileBtn').onclick=()=>fileInp.click();
  fileInp.addEventListener('change',()=>uploadSelectedFile(fileInp.files[0]));
  inp.addEventListener('keydown',e=>{ if(e.key==='Enter'&&!e.shiftKey){e.preventDefault();send();} });
  inp.addEventListener('input',()=>{ inp.style.height='auto'; inp.style.height=Math.min(inp.scrollHeight,160)+'px'; });
  inp.addEventListener('paste',handlePaste);
  document.getElementById('send').onclick=send;
  document.getElementById('resetBtn').onclick=resetChat;
  document.getElementById('resetSessBtn').onclick=resetSession;

  // Загружаем сохранённую сессию из localStorage
  loadSessionFromStorage();

  // Если агент был выбран ранее, переходим к нему, иначе первый
  const initialAgent=cur||AGENTS[0].key;
  switchAgent(initialAgent);
};
</script>
</body>
</html>"""


SETTINGS_HTML = r"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1"><title>MILA OFFICE · Настройки</title>
<style>
  body{margin:0;font-family:Georgia,'Times New Roman',serif;background:#FAF6F1;color:#1E140F}
  .wrap{max-width:720px;margin:0 auto;padding:28px 20px}
  h1{color:#C4614A;font-size:24px;margin:0 0 4px}
  .sub{color:#7A5E54;font-size:13px;margin-bottom:22px}
  .card{background:#fff;border:1px solid #E0D0C8;border-radius:14px;padding:18px 20px;margin-bottom:14px}
  .card h3{margin:0 0 6px;font-size:17px}
  .badge{display:inline-block;font-size:12px;padding:3px 10px;border-radius:20px;margin-left:8px;vertical-align:middle}
  .ok{background:#E3F0E6;color:#2C5F3A}.no{background:#F6E1DC;color:#A8412C}
  .meta{font-size:13px;color:#7A5E54;line-height:1.6;margin:6px 0 0}
  .btn{display:inline-block;margin-top:12px;background:#C4614A;color:#fff;border:none;border-radius:8px;
       padding:10px 18px;font-size:14px;font-family:inherit;cursor:pointer;text-decoration:none}
  .btn.dim{background:#B7A89F;cursor:default}
  .btn.sec{background:#7A5E54;padding:7px 13px;font-size:13px}
  .tres{font-size:13px;margin-left:6px}
  code{background:rgba(0,0,0,.06);padding:1px 5px;border-radius:4px;font-family:monospace;font-size:12px}
  a.back{color:#7A5E54;font-size:13px}
</style></head><body><div class="wrap">
  <p><a class="back" href="/">← В чат</a></p>
  <h1>Настройки и подключения</h1>
  <div class="sub">Чек-лист подключений. Нажми «Проверить» — увидишь живой статус. Секреты не отображаются.</div>
  <div id="cards">Загрузка…</div>
<script>
let CSRF='';
function esc(s){return String(s==null?'':s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}
function badge(ok){return '<span class="badge '+(ok?'ok':'no')+'">'+(ok?'подключено':'не настроено')+'</span>';}
function testCell(svc){return ' <button class="btn sec" data-svc="'+svc+'" onclick="runTest(this)">Проверить</button><span class="tres" id="tr-'+svc+'"></span>';}
async function runTest(btn){
  var svc=btn.getAttribute('data-svc'), s=document.getElementById('tr-'+svc);
  s.textContent='Проверяю…'; btn.disabled=true;
  try{ var d=await (await fetch('/api/test/'+svc)).json();
    s.innerHTML=d.ok?('<span style="color:#2C5F3A">✓ '+esc(d.detail||'ок')+'</span>')
                    :('<span style="color:#A8412C">⚠️ '+esc(d.error||'ошибка')+'</span>');
  }catch(e){ s.textContent='⚠️ сеть'; }
  btn.disabled=false;
}
async function saveToken(){
  const tok=document.getElementById('igtok').value.trim();
  const msg=document.getElementById('igmsg');
  if(!tok){ msg.textContent='Вставь токен'; return; }
  msg.textContent='Проверяю…';
  try{
    const r=await fetch('/api/settings/instagram-token',{method:'POST',
      headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body:JSON.stringify({token:tok})});
    const d=await r.json();
    if(d.ok){ msg.innerHTML='✓ Сохранено: @'+(d.username||'?')+'. Перезапусти приложение, чтобы агенты подхватили токен.'; }
    else { msg.textContent='⚠️ '+(d.error||'ошибка'); }
  }catch(e){ msg.textContent='⚠️ сеть: '+e; }
}
async function load(){
  const s = await (await fetch('/api/settings')).json();
  CSRF=s.csrf;
  let h={}; try{ h = await (await fetch('/api/health')).json(); }catch(e){}
  const el = document.getElementById('cards'); el.innerHTML='';
  // Claude
  el.innerHTML += '<div class="card"><h3>Claude (Anthropic)'+badge(s.claude.configured)+'</h3>'
    + '<p class="meta">Режим: <code>'+s.claude.mode+'</code><br>'+s.claude.note+'</p>'+testCell('claude')+'</div>';
  // Gemini
  el.innerHTML += '<div class="card"><h3>Gemini'+badge(s.gemini.configured)+'</h3>'
    + '<p class="meta">Provider: <code>'+s.gemini.provider+'</code> / model: <code>'+s.gemini.model+'</code><br>'
    + 'Heavy work: <code>'+(s.gemini.heavy_lifting?'Gemini':'Claude')+'</code> / Claude agents: <code>'
    + (s.gemini.anthropic_agents||[]).join(', ')+'</code></p>'+testCell('gemini')+'</div>';
  // Instagram + OAuth
  const ig = s.instagram;
  let igc = '<div class="card"><h3>Instagram'+badge(ig.configured)+'</h3>'
    + '<p class="meta">Flow: <code>'+ig.flow+'</code> · node: <code>'+(ig.node||'—')+'</code><br>'
    + 'Redirect URI: <code>'+ig.redirect_uri+'</code> — зарегистрируй его в кабинете Meta '
    + '(часто нужен HTTPS; localhost-http может быть отклонён).</p>';
  if(ig.oauth){ igc += '<a class="btn" href="/auth/instagram/login">🔗 Подключить через OAuth</a>'; }
  else { igc += '<span class="btn dim">OAuth недоступен</span>'
    + '<p class="meta">Задай <code>IG_APP_ID</code> и <code>IG_APP_SECRET</code> в tools/.env.</p>'; }
  igc += '<hr style="border:none;border-top:1px solid #E0D0C8;margin:16px 0">'
    + '<p class="meta"><b>Или вставь токен вручную</b> — надёжно на localhost (без redirect URI). '
    + 'Сгенерируй в кабинете Meta токен с правом <code>instagram_business_manage_comments</code> и вставь сюда:</p>'
    + '<input id="igtok" placeholder="IG access token" '
    + 'style="width:100%;padding:9px 12px;border:1px solid #E0D0C8;border-radius:8px;'
    + 'font-family:monospace;font-size:12px;margin-bottom:8px">'
    + '<button class="btn" onclick="saveToken()">Проверить и сохранить</button> '
    + '<span id="igmsg" class="meta"></span>';
  igc += '<div style="margin-top:12px">'+testCell('instagram')+'</div></div>'; el.innerHTML += igc;
  // Telegram
  el.innerHTML += '<div class="card"><h3>Telegram'+badge(s.telegram.configured)+'</h3>'
    + '<p class="meta">Токен бота в <code>TELEGRAM_BOT_TOKEN</code> / <code>TELEGRAM_API</code>.</p>'+testCell('telegram')+'</div>';
  // Supabase (статус из /api/health)
  const sb=h.supabase||{};
  el.innerHTML += '<div class="card"><h3>Supabase'+badge(sb.configured)+'</h3>'
    + '<p class="meta">База продаж и продуктов. '+(sb.can_write?'Доступна запись.':'Только чтение или не настроено.')+'</p>'+testCell('supabase')+'</div>';
  // n8n + мост
  const br=h.bridge||{}, nn=h.n8n||{};
  el.innerHTML += '<div class="card"><h3>n8n + мост'+badge(br.up)+'</h3>'
    + '<p class="meta">Автоматизация: расписания и очередь задач. n8n: <code>'+(nn.up?'up':'down')+'</code> · мост: <code>'+(br.up?'up':'down')+'</code></p>'+testCell('bridge')+'</div>';
  // Gumroad
  el.innerHTML += '<div class="card"><h3>Gumroad'+badge(s.gumroad.configured)+'</h3>'
    + '<p class="meta">Токен в <code>GUMROAD_ACCESS_TOKEN</code>. OAuth не применяется.</p></div>';
}
load();
</script></div></body></html>"""


# ─── Дашборд Людмилы (одна страница, кнопка «Одобрить всё») ──
DASHBOARD_HTML = r"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>MILA — Дашборд</title>
<style>
  :root{--t:#C4614A;--n:#1E140F;--c:#FAF6F1;--m:#F2EAE2;--u:#7A5E54;--b:#E0D0C8;--w:#fff;--g:#4A7A5E;--r:#A8412C}
  *{box-sizing:border-box}
  html,body{height:100%}
  body{margin:0;font-family:Georgia,'Times New Roman',serif;background:var(--c);color:var(--n);
       height:100vh;display:flex;flex-direction:column;overflow:hidden}
  header{background:var(--n);padding:16px 24px;flex-shrink:0}
  header .k{font-size:11px;color:var(--t);letter-spacing:2px}
  header .h{font-size:22px;color:var(--w);margin-top:4px;font-weight:bold}
  header .s{font-size:12px;color:#9a8278;margin-top:3px}
  .wrap{flex:1;max-width:1100px;width:100%;margin:0 auto;padding:18px 22px;display:flex;
        flex-direction:column;min-height:0;overflow-y:auto}
  .top{display:flex;align-items:center;gap:10px;margin-bottom:16px;flex-wrap:wrap;flex-shrink:0}
  .top a,.top button{font-family:inherit;font-size:13px;color:var(--t);background:rgba(196,97,74,.08);
       border:1px solid var(--b);border-radius:8px;padding:8px 14px;text-decoration:none;cursor:pointer;transition:.15s}
  .top a:hover,.top button:hover{border-color:var(--t);background:rgba(196,97,74,.16)}
  .top .spacer{flex:1}
  .top .auto{color:var(--u);font-size:12px;background:none;border:0;cursor:default}
  /* KPI */
  .kpis{display:grid;grid-template-columns:repeat(5,1fr);gap:12px;margin-bottom:18px;flex-shrink:0}
  .kpi{background:var(--w);border:1px solid var(--b);border-radius:14px;padding:14px 16px}
  .kpi .v{font-size:24px;font-weight:bold;color:var(--n)}
  .kpi .l{font-size:11px;color:var(--u);margin-top:4px;text-transform:uppercase;letter-spacing:.5px}
  .kpi .d{font-size:11px;color:var(--g);margin-top:2px}
  @media(max-width:860px){.kpis{grid-template-columns:repeat(2,1fr)}}
  /* hero */
  .hero{background:var(--w);border:2px solid var(--t);border-radius:16px;padding:18px 20px;margin-bottom:18px;
        display:flex;align-items:center;gap:18px;flex-wrap:wrap;flex-shrink:0}
  .hero .sum{flex:1;min-width:220px;font-size:14px;color:var(--u);line-height:1.5}
  .hero .sum b{color:var(--n)}
  .approve{background:var(--t);color:#fff;border:none;border-radius:12px;padding:14px 28px;font-size:16px;
       font-family:inherit;cursor:pointer;font-weight:bold;transition:.15s}
  .approve:hover{filter:brightness(1.08)}
  .approve:disabled{opacity:.45;cursor:default}
  .grid{display:grid;grid-template-columns:1fr 1fr;gap:16px}
  @media(max-width:760px){.grid{grid-template-columns:1fr}}
  .card{background:var(--w);border:1px solid var(--b);border-radius:14px;padding:18px 20px;margin-bottom:16px;
        display:flex;flex-direction:column;min-height:0}
  .card h2{font-size:15px;margin:0 0 12px;color:var(--n);flex-shrink:0}
  .card h2 .n{color:var(--t);font-size:13px;margin-left:6px}
  .card .body{overflow-y:auto;max-height:300px;min-height:0}
  .row{padding:9px 0;border-bottom:1px solid var(--m);font-size:13px;color:var(--n);line-height:1.5;
       display:flex;justify-content:space-between;gap:10px;align-items:flex-start}
  .row:last-child{border:0}
  .row .txt{flex:1;min-width:0;word-break:break-word}
  .row .meta{color:var(--u);font-size:11px;margin-top:3px}
  .empty{color:var(--u);font-size:13px;font-style:italic}
  .pill{display:inline-block;font-size:10px;padding:1px 8px;border-radius:8px;background:var(--m);color:var(--u);margin-right:6px}
  .pill.p1{background:#f6dcd6;color:#a23a28}
  .ib{display:flex;gap:5px;flex-shrink:0}
  .ib button{border:1px solid var(--b);background:var(--w);border-radius:8px;padding:4px 9px;font-size:12px;
       font-family:inherit;cursor:pointer;color:var(--n);transition:.15s}
  .ib button:hover{border-color:var(--t);color:var(--t)}
  .ib button.ok:hover{border-color:var(--g);color:var(--g)}
  .ib button.no:hover{border-color:var(--r);color:var(--r)}
  .body::-webkit-scrollbar{width:10px}.body::-webkit-scrollbar-thumb{background:var(--b);border-radius:6px}
  .body::-webkit-scrollbar-track{background:transparent}
  #toast{position:fixed;bottom:24px;left:50%;transform:translateX(-50%);background:var(--g);color:#fff;
       padding:12px 22px;border-radius:10px;font-size:14px;display:none;z-index:50}
</style></head>
<body>
<header>
  <div class="k">УТРЕННИЙ ОБЗОР · MILA OFFICE</div>
  <div class="h">Дашборд Людмилы</div>
  <div class="s" id="sub">@liudmyla.lykova</div>
</header>
<div class="wrap">
  <div class="top">
    <a href="/">Агенты</a><a href="/operator">Очередь</a><a href="/settings">Настройки</a>
    <button onclick="load()">Обновить</button>
    <span class="spacer"></span>
    <span class="auto" id="auto">авто-обновление: 30с</span>
  </div>

  <div class="kpis" id="kpis"></div>

  <div class="hero">
    <div class="sum" id="summary">Загружаю…</div>
    <button class="approve" id="approveAll" disabled>Одобрить всё</button>
  </div>

  <div class="grid">
    <div class="card"><h2>📸 Посты на одобрение <span class="n" id="cPosts"></span></h2><div class="body" id="posts"></div></div>
    <div class="card"><h2>✅ Задачи офиса <span class="n" id="cActions"></span></h2><div class="body" id="actions"></div></div>
  </div>
  <div class="grid">
    <div class="card"><h2>⚙️ Улучшения от Стаса <span class="n" id="cImpr"></span></h2><div class="body" id="impr"></div></div>
    <div class="card"><h2>🕑 Что было ночью <span class="n" id="cEv"></span></h2><div class="body" id="events"></div></div>
  </div>
</div>
<div id="toast"></div>
<script>
let CSRF="";
function esc(s){return (s||"").replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}
function toast(m){const t=document.getElementById('toast');t.textContent=m;t.style.display='block';setTimeout(()=>t.style.display='none',3500);}
function nf(n){return (n==null)?'—':String(n).replace(/\B(?=(\d{3})+(?!\d))/g,' ');}
// «Утренний пульт» без техножаргона: человеческие имена задач, фаз и событий.
const PNAME={content_week:'Контент на неделю',new_client:'Новая клиентка',monday_brief:'Утренний бриф',
  weekly_report:'Недельный отчёт',competitive_analysis:'Анализ конкурентов',product_research:'Исследование продукта',
  new_product:'Новый продукт'};
const PHASE_RU={cold_start:'запуск',learning:'накопление данных',analysis:'анализ'};
function phaseRu(p){return PHASE_RU[p]||p||'—';}
function rel(ts){if(!ts)return '';var d=new Date(ts).getTime();if(isNaN(d))return esc(ts);var s=(Date.now()-d)/1000;
  if(s<0)s=0;if(s<60)return 'только что';if(s<3600)return Math.floor(s/60)+' мин назад';
  if(s<86400)return Math.floor(s/3600)+' ч назад';return Math.floor(s/86400)+' дн назад';}
function evRu(k){
  k=k||'';
  if(k.indexOf('pipeline:done:')===0)return '✅ Готово: '+(PNAME[k.slice(14)]||k.slice(14));
  if(k.indexOf('pipeline:start:')===0)return '▶ Запущено: '+(PNAME[k.slice(15)]||k.slice(15));
  if(k.indexOf('pipeline:fail:')===0)return '⚠️ Сбой: '+(PNAME[k.slice(14)]||k.slice(14));
  if(k.indexOf('context:')===0)return 'Событие из мира';
  const M={'task:queued':'Поставлено в очередь','task:dequeued':'Взято в работу','task:complete':'Задача завершена',
    'task:recovered':'Задача восстановлена','task:cancelled':'Задача отменена','published':'📸 Пост опубликован',
    'measured':'📊 Замерены метрики поста','approval:set':'Решение по одобрению','profile:update':'Профиль обновлён',
    'handoff':'Передача между агентами','queue:cleared':'Очередь очищена'};
  return M[k]||k;
}

async function item(url){
  const r=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body:'{}'});
  const d=await r.json();
  if(d.ok){ toast(d.result==='needs_media'?'⚠️ Нет медиа — нужен файл':'Готово'); load(); }
  else toast('Ошибка: '+(d.error||'?'));
}
async function approveAll(){
  const btn=document.getElementById('approveAll'); btn.disabled=true;
  try{
    const r=await fetch('/api/approve-all',{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body:'{}'});
    const d=await r.json();
    if(d.ok){ toast('Одобрено постов: '+d.approved+'. Опубликуются по расписанию.'); load(); }
    else { toast('Ошибка: '+(d.error||'?')); btn.disabled=false; }
  }catch(e){ toast('Сеть недоступна'); btn.disabled=false; }
}

async function load(){
  let d; try{ d=await (await fetch('/api/dashboard')).json(); }catch(e){ return; }
  CSRF=d.csrf;
  const ph=d.profile||{}, k=d.kpi||{};
  document.getElementById('sub').textContent='@liudmyla.lykova · этап: '+phaseRu(k.phase||ph.phase);

  // KPI карточки
  document.getElementById('kpis').innerHTML=[
    {v:nf(k.followers),l:'подписчики'},
    {v:nf(k.avg_reach),l:'ср. охват'},
    {v:(k.er!=null?k.er+'%':'—'),l:'вовлечённость'},
    {v:nf(k.sales),l:'продажи'},
    {v:phaseRu(k.phase),l:'этап',d:k.goal?('цель: '+esc(k.goal)):''},
  ].map(x=>'<div class="kpi"><div class="v">'+x.v+'</div><div class="l">'+x.l+'</div>'+(x.d?'<div class="d">'+x.d+'</div>':'')+'</div>').join('');

  const np=d.pending_posts.length, na=d.open_actions.length, ni=d.improvements.length;
  document.getElementById('summary').innerHTML=
    'Доброе утро! На одобрении: <b>'+np+'</b> пост(ов), <b>'+na+'</b> задач(и), <b>'+ni+'</b> улучшений агентов. '
    +(np?'Жми «Одобрить всё» или решай по каждому посту.':'Новых постов на одобрение нет.');
  document.getElementById('approveAll').disabled = np===0;

  // посты + кнопки на каждый
  document.getElementById('cPosts').textContent=np||'';
  document.getElementById('posts').innerHTML = np ? d.pending_posts.map(p=>
    '<div class="row"><div class="txt">'+esc(p.caption||'(без текста)')
    +'<div class="meta">'+(p.when?('⏰ '+esc(p.when)+' · '):'')+(p.has_media?'медиа ✓':'⚠️ нет медиа')+'</div></div>'
    +'<div class="ib"><button class="ok" onclick="item(\'/api/dashboard/post/'+encodeURIComponent(p.id)+'/approve\')">одобрить</button>'
    +'<button class="no" onclick="item(\'/api/dashboard/post/'+encodeURIComponent(p.id)+'/reject\')">скрыть</button></div></div>'
  ).join('') : '<div class="empty">Пусто</div>';

  // задачи + кнопка закрыть
  document.getElementById('cActions').textContent=na||'';
  document.getElementById('actions').innerHTML = na ? d.open_actions.map(a=>
    '<div class="row"><div class="txt"><span class="pill '+((a.priority||'').toLowerCase()==='p1'?'p1':'')+'">'+esc(a.priority||'P?')+'</span>'
    +esc(a.title||'')+'<div class="meta">'+esc(a.assignee||'')+(a.due?(' · до '+esc(a.due)):'')+'</div></div>'
    +'<div class="ib"><button class="ok" onclick="item(\'/api/dashboard/action/'+encodeURIComponent(a.id)+'/close\')">закрыть</button></div></div>'
  ).join('') : '<div class="empty">Открытых задач нет</div>';

  // улучшения
  document.getElementById('cImpr').textContent=ni||'';
  document.getElementById('impr').innerHTML = ni ? d.improvements.map(i=>
    '<div class="row"><div class="txt"><b>'+esc(i.agent)+'</b><div class="meta">темы: '+i.topics.map(esc).join(', ')+'</div></div></div>'
  ).join('') : '<div class="empty">Активных улучшений нет</div>';

  // события
  document.getElementById('cEv').textContent=d.events.length||'';
  document.getElementById('events').innerHTML = d.events.length ? d.events.map(e=>
    '<div class="row"><div class="txt">'+esc(evRu(e.kind))+'<div class="meta">'+esc(rel(e.ts))+'</div></div></div>'
  ).join('') : '<div class="empty">Событий нет</div>';
}

document.getElementById('approveAll').onclick=approveAll;
load();
setInterval(load, 30000);  // авто-обновление каждые 30с
</script></body></html>"""


OPERATOR_HTML = r"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>MILA OFFICE · Очередь</title>
<style>
  :root{--t:#C4614A;--n:#1E140F;--c:#FAF6F1;--m:#F2EAE2;--u:#7A5E54;--b:#E0D0C8;--w:#fff;--g:#4A7A5E;--r:#A8412C}
  *{box-sizing:border-box}
  html,body{height:100%}
  body{margin:0;font-family:Georgia,'Times New Roman',serif;background:var(--c);color:var(--n);
       height:100vh;display:flex;flex-direction:column;overflow:hidden}
  header{background:var(--n);padding:16px 24px;color:#fff;flex-shrink:0}
  header .k{font-size:11px;color:var(--t);letter-spacing:2px}
  header .h{font-size:22px;margin-top:4px;font-weight:bold}
  .wrap{flex:1;max-width:1180px;width:100%;margin:0 auto;padding:18px 22px;display:flex;
        flex-direction:column;min-height:0}
  .top{display:flex;align-items:center;gap:10px;margin-bottom:14px;flex-wrap:wrap;flex-shrink:0}
  .top a,.top button{font-family:inherit;font-size:13px;color:var(--t);background:rgba(196,97,74,.08);
       border:1px solid var(--b);border-radius:8px;padding:8px 14px;text-decoration:none;cursor:pointer;transition:.15s}
  .top a:hover,.top button:hover{border-color:var(--t);background:rgba(196,97,74,.16)}
  .tabs{display:flex;gap:8px;flex-wrap:wrap;margin-bottom:16px;flex-shrink:0}
  .tab{background:var(--w);border:1px solid var(--b);border-radius:20px;padding:8px 16px;font-family:inherit;
       font-size:13px;cursor:pointer;color:var(--n);transition:.15s}
  .tab:hover{border-color:var(--t);color:var(--t)}
  .tab.on{border-color:var(--t);color:#fff;background:var(--t)}
  .grid{display:grid;grid-template-columns:1.8fr 1fr;gap:16px;flex:1;min-height:0}
  .col{display:flex;flex-direction:column;gap:16px;min-height:0}
  .card{background:var(--w);border:1px solid var(--b);border-radius:14px;padding:18px 20px;
        display:flex;flex-direction:column;min-height:0}
  .card.scroll{overflow:hidden}
  .card .body{overflow-y:auto;min-height:0}
  h2{font-size:16px;margin:0 0 12px;flex-shrink:0}
  table{width:100%;border-collapse:collapse;font-size:13px}
  thead th{position:sticky;top:0;background:var(--w);z-index:1}
  th{text-align:left;color:var(--u);font-weight:normal;border-bottom:1px solid var(--b);padding:8px 7px}
  td{border-bottom:1px solid var(--m);padding:9px 7px;vertical-align:top}
  .pill{display:inline-block;border-radius:8px;padding:2px 8px;background:var(--m);font-size:11px;color:var(--u)}
  .st{display:inline-block;border-radius:8px;padding:2px 9px;font-size:12px}
  .pending{color:#8B6B10;background:rgba(139,107,16,.10)}
  .running{color:#2B5278;background:rgba(43,82,120,.10)}
  .failed{color:var(--r);background:rgba(168,65,44,.10)}
  .done{color:var(--g);background:rgba(74,122,94,.10)}
  .cancelled{color:#777;background:rgba(119,119,119,.10)}
  .awaiting_approval{color:#8B4513;background:rgba(139,69,19,.10)}
  .actions{display:flex;gap:6px;flex-wrap:wrap}
  .actions button{border:1px solid var(--b);background:var(--w);border-radius:8px;padding:5px 10px;
       font-size:12px;font-family:inherit;cursor:pointer;color:var(--n);transition:.15s}
  .actions button:hover{border-color:var(--t);color:var(--t)}
  .muted{color:var(--u);font-size:12px}
  .event{border-bottom:1px solid var(--m);padding:9px 0;font-size:13px;word-break:break-word}
  .event:last-child{border:0}
  .toast{position:fixed;bottom:20px;left:50%;transform:translateX(-50%);background:var(--g);color:#fff;
       border-radius:8px;padding:10px 16px;display:none;z-index:50}
  /* Скроллбары в тон теме */
  .card .body::-webkit-scrollbar{width:10px}
  .card .body::-webkit-scrollbar-thumb{background:var(--b);border-radius:6px}
  .card .body::-webkit-scrollbar-track{background:transparent}
  @media(max-width:760px){.grid{grid-template-columns:1fr;overflow-y:auto} th:nth-child(4),td:nth-child(4){display:none}}
  /* Карточки задач: человеческий вид + результат/артефакт + раскрытие деталей */
  .tcard{border:1px solid var(--b);border-radius:12px;padding:12px 14px;margin-bottom:10px;background:var(--w)}
  .tc-head{display:flex;align-items:center;gap:10px;flex-wrap:wrap}
  .tc-name{font-weight:bold;font-size:14px}
  .tc-time{color:var(--u);font-size:12px;margin-left:auto}
  .tc-result{margin:10px 0 2px;padding:10px 12px;background:var(--c);border-radius:9px;font-size:13px;line-height:1.5}
  .tc-sum{color:var(--n);white-space:pre-wrap;word-break:break-word}
  .art-f{color:var(--u);font-size:12px;margin-top:6px;word-break:break-all}
  .art-u{margin-top:6px}.art-u a{color:var(--t);font-size:12px;word-break:break-all;text-decoration:none}
  .tc-next{margin-top:8px;color:#8B4513;font-size:12px}
  .tc-det{margin-top:8px}
  .tc-det summary{cursor:pointer;color:var(--u);font-size:12px;list-style:none}
  .tc-det summary::-webkit-details-marker{display:none}
  .tc-det summary:before{content:'\25B8 ';color:var(--t)}
  .tc-det[open] summary:before{content:'\25BE '}
  .tc-det .kv{font-size:12px;color:var(--u);margin-top:6px;line-height:1.8;word-break:break-word}
  .tc-det .kv code{background:var(--m);padding:1px 5px;border-radius:4px;font-size:11px}
  .tcard .actions{margin-top:10px}
</style></head><body>
<header><div class="k">MILA OFFICE · ОПЕРАТОР</div><div class="h">Управление очередью</div></header>
<div class="wrap">
  <div class="top"><a href="/">Агенты</a><a href="/dashboard">Дашборд</a><a href="/settings">Настройки</a><button onclick="load()">Обновить</button></div>
  <div class="tabs" id="tabs"></div>
  <div class="grid">
    <div class="col" style="display:flex;flex-direction:column;gap:16px">
      <div class="card scroll" style="flex:1 1 auto"><h2>Задачи</h2><div class="body" id="tasks"></div></div>
      <div class="card scroll" style="flex:0 1 auto;max-height:35%"><h2>Документы в обработке</h2><div class="body" id="documents"></div></div>
    </div>
    <div class="col">
      <div class="card" style="flex:0 0 auto"><h2>Supervisor</h2><div id="supervisor" class="muted"></div></div>
      <div class="card scroll" style="flex:0 1 auto;max-height:25%"><h2>Ждут одобрения</h2><div class="body" id="approvals"></div></div>
      <div class="card scroll" style="flex:0 1 auto;max-height:25%"><h2>Очередь ответов</h2><div class="body" id="reply_queue"></div></div>
      <div class="card scroll" style="flex:1 1 auto"><h2>Последние события</h2><div class="body" id="events"></div></div>
    </div>
  </div>
</div>
<div class="toast" id="toast"></div>
<script>
let CSRF='', FILTER='';
// status-фильтры: значение для API → русская подпись
const STATUSES=[['','все'],['pending','в очереди'],['running','выполняется'],
  ['awaiting_approval','ждут одобрения'],['failed','ошибка'],['done','готово'],['cancelled','отменено']];
const ST_RU={pending:'в очереди',running:'выполняется',awaiting_approval:'ждёт одобрения',
  failed:'ошибка',done:'готово',cancelled:'отменено'};
const ACT_RU={retry:'повторить',unblock:'разблокировать',cancel:'отменить'};
function esc(s){return String(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}
function toast(s){const t=document.getElementById('toast');t.textContent=s;t.style.display='block';setTimeout(()=>t.style.display='none',2200);}
function renderTabs(){
  document.getElementById('tabs').innerHTML=STATUSES.map(([v,ru])=>
    '<button class="tab '+(v===FILTER?'on':'')+'" onclick="FILTER=\''+v+'\';load()">'+ru+'</button>').join('');
}
async function act(id, action){
  const body=action==='retry'?{reset_attempts:false}:action==='cancel'?{reason:'operator'}:{};
  const r=await fetch('/api/operator/task/'+encodeURIComponent(id)+'/'+action,{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF},body:JSON.stringify(body)});
  const d=await r.json();
  toast(d.ok ? (ACT_RU[action]||action)+': '+id : (d.error||'ошибка'));
  load();
}
function actions(t){
  if(t.status==='running') return '';
  return '<div class="actions"><button onclick="act(\''+esc(t.id)+'\',\'retry\')">повторить</button>'+
    '<button onclick="act(\''+esc(t.id)+'\',\'unblock\')">разблок.</button>'+
    '<button onclick="act(\''+esc(t.id)+'\',\'cancel\')">отменить</button></div>';
}
async function sendReplyOne(){
  const btn=event.target;
  btn.disabled=true; btn.textContent='Отправляю…';
  try{
    const r=await fetch('/api/reply-send-one',{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF}});
    const d=await r.json();
    toast(d.ok?('Отправлен: '+d.detail):('Ошибка: '+(d.error||'?')));
  }catch(e){ toast('Сеть: '+e); }
  btn.disabled=false; btn.textContent='Отправить 1';
  load();
}
async function sendReplyAll(){
  const btn=event.target;
  btn.disabled=true; btn.textContent='Отправляю все…';
  try{
    const r=await fetch('/api/reply-send-all',{method:'POST',headers:{'Content-Type':'application/json','X-CSRF-Token':CSRF}});
    const d=await r.json();
    toast(d.ok?(d.sent+' отправлено, '+d.failed+' ошибок'):('Ошибка: '+(d.error||'?')));
  }catch(e){ toast('Сеть: '+e); }
  btn.disabled=false; btn.textContent='Отправить все';
  load();
}
async function delReply(id){
  if(!confirm('Удалить ответ из очереди?')) return;
  try{
    const r=await fetch('/api/reply-delete/'+encodeURIComponent(id),{method:'POST',headers:{'X-CSRF-Token':CSRF}});
    const d=await r.json();
    toast(d.ok?'Удалено':'Ошибка: '+(d.error||'?'));
  }catch(e){ toast('Сеть: '+e); }
  load();
}
// Человеческие имена пайплайнов — чтобы оператор видел «Контент на неделю», а не content_week.
const PNAME={content_week:'Контент на неделю',new_client:'Новая клиентка',monday_brief:'Утренний бриф',
  weekly_report:'Недельный отчёт',competitive_analysis:'Анализ конкурентов',product_research:'Исследование продукта',
  new_product:'Новый продукт'};
function pname(p){return PNAME[p]||p||'—';}
function rel(ts){if(!ts)return '';var d=new Date(ts).getTime();if(isNaN(d))return '';var s=(Date.now()-d)/1000;
  if(s<0)s=0;if(s<60)return 'только что';if(s<3600)return Math.floor(s/60)+' мин назад';
  if(s<86400)return Math.floor(s/3600)+' ч назад';return Math.floor(s/86400)+' дн назад';}
function artHtml(a){
  if(!a)return '';
  var h='<div class="tc-result">';
  if(a.summary)h+='<div class="tc-sum">'+esc(a.summary)+'</div>';
  (a.files||[]).forEach(function(f){h+='<div class="art-f">📄 '+esc(f)+'</div>';});
  (a.urls||[]).forEach(function(u){h+='<div class="art-u"><a href="'+esc(u)+'" target="_blank" rel="noopener">🔗 '+esc(u)+'</a></div>';});
  (a.next_actions||[]).forEach(function(n){h+='<div class="tc-next">Дальше: '+esc(n)+'</div>';});
  return h+'</div>';
}
function details(t){
  var b=['id <code>'+esc(t.id)+'</code>','попыток: '+esc(t.attempts||0)];
  if(t.dedupe_key)b.push('dedupe <code>'+esc(t.dedupe_key)+'</code>');
  if(t.worker_id)b.push('worker <code>'+esc(t.worker_id)+'</code>');
  if(t.lease_expires_at)b.push('lease до '+esc(t.lease_expires_at));
  if(t.next_run_at)b.push('след. запуск '+esc(t.next_run_at));
  if(t.priority!=null)b.push('приоритет '+esc(t.priority));
  return '<details class="tc-det"><summary>технические детали</summary><div class="kv">'+b.join(' · ')+'</div></details>';
}
function taskCard(t){
  var a=(t.result&&t.result.artifact)||(t.last_result&&t.last_result.artifact)||null;
  var when=t.finished_at||t.started_at||t.next_run_at||t.created_at;
  return '<div class="tcard"><div class="tc-head">'
    +'<span class="st '+esc(t.status)+'">'+esc(ST_RU[t.status]||t.status)+'</span>'
    +'<span class="tc-name">'+esc(pname(t.pipeline))+'</span>'
    +'<span class="tc-time">'+esc(rel(when))+'</span></div>'
    +artHtml(a)+details(t)+actions(t)+'</div>';
}
async function load(){
  renderTabs();
  const url='/api/operator'+(FILTER?'?status='+encodeURIComponent(FILTER):'');
  const d=await (await fetch(url)).json();
  CSRF=d.csrf;
  const sv=d.supervisor||{};
  const svc=sv.services||{};
  document.getElementById('supervisor').innerHTML=
    '<span class="st '+(sv.status==='ok'?'done':'failed')+'">'+esc(sv.status||'missing')+'</span>'+
    '<div class="muted">webapp: '+esc((svc.webapp&&svc.webapp.up)?'up':'down')+
    ' · bridge: '+esc((svc.bridge&&svc.bridge.up)?'up':'down')+
    ' · n8n: '+esc((svc.n8n&&svc.n8n.up)?'up':'down')+'</div>'+
    '<div class="muted">worker: '+esc((sv.last_worker&&sv.last_worker.ok)?'ok':((sv.last_worker&&sv.last_worker.error)||'—'))+'</div>';
  const tasks=d.tasks||[];
  document.getElementById('tasks').innerHTML=tasks.length?tasks.map(taskCard).join(''):'<div class="muted">Задач нет</div>';
  const approvals=d.pending_approvals||{};
  const ak=Object.keys(approvals);
  document.getElementById('approvals').innerHTML=ak.length?ak.map(k=>'<div class="event"><b>'+esc(k)+'</b><div class="muted">'+esc(ST_RU[approvals[k].status]||approvals[k].status)+(approvals[k].comment?' · '+esc(approvals[k].comment):'')+'</div></div>').join(''):'<div class="muted">Нет задач на одобрении</div>';
  // Reply queue (ответы на Instagram-комментарии)
  const rq=d.reply_queue||{pending:0,sent:0,failed:0,items_pending:[]};
  const rqTot=(rq.pending||0)+(rq.sent||0)+(rq.failed||0);
  let rqHtml='<div class="muted">в очереди: '+esc(rq.pending||0)+' · отправлено: '+esc(rq.sent||0)+' · ошибок: '+esc(rq.failed||0)+'</div>';
  const items=(rq.items_pending||[]).slice(0,5);
  if(items.length){rqHtml+='<div style="margin-top:8px;font-size:12px">'+items.map(item=>'<div style="border-bottom:1px solid #E0D0C8;padding:6px 0;display:flex;justify-content:space-between;align-items:center;gap:8px"><div><span style="color:var(--u)">@'+esc(item.username||'?')+'</span><br><span style="color:#666;font-size:11px">'+esc((item.comment_text||'').slice(0,40)+(item.comment_text&&item.comment_text.length>40?'...':''))+'</span></div><button class="actions" style="display:inline-block;border:1px solid #ddd;background:var(--w);border-radius:6px;padding:3px 7px;font-size:11px;font-family:inherit;cursor:pointer;color:var(--r);flex-shrink:0" onclick="delReply(\''+esc(item.id)+'\');" title="Удалить из очереди">✕</button></div>').join('')+'</div>';
    if(rq.pending>0){rqHtml+='<div style="margin-top:10px;display:flex;gap:8px"><button class="actions" style="display:inline-block;border:1px solid var(--b);background:var(--w);border-radius:8px;padding:5px 10px;font-size:12px;font-family:inherit;cursor:pointer;color:var(--n);transition:.15s" onclick="sendReplyOne()" title="Отправить один ответ из очереди">Отправить 1</button><button class="actions" style="display:inline-block;border:1px solid var(--b);background:var(--w);border-radius:8px;padding:5px 10px;font-size:12px;font-family:inherit;cursor:pointer;color:var(--n);transition:.15s" onclick="sendReplyAll()" title="Отправить все ответы из очереди подряд">Отправить все</button></div>';}
  }
  document.getElementById('reply_queue').innerHTML=rqTot?rqHtml:'<div class="muted">Очередь пуста</div>';

  // Документ workflows (timeline)
  try{
    const docs=await (await fetch('/api/documents')).json();
    const inProg=(docs.in_progress||[]).slice(0,5);
    let docHtml=inProg.length?inProg.map(d=>{
      const stages=(d.stages||[]);
      const names={'victoria':'Виктория','rita':'Рита','marina':'Марина','lera':'Лера',
        'producer':'Продюсер','manager':'Менеджер','vasya':'Вася','dima':'Дима',
        'tyoma':'Тёма','olya':'Оля','alina':'Алина'};
      const vBadges={'ready_next':'✓','needs_revision':'⚠','done':'✓'};
      let tl='<div style="font-size:12px;margin:8px 0;cursor:pointer;padding:8px;border-radius:6px;transition:.15s;border:1px solid transparent" onmouseover="this.style.background=\'rgba(196,97,74,.08);border-color=#E0D0C8\'" onmouseout="this.style.background=\'transparent\'" onclick="openDocModal(\''+esc(d.id)+'\')" title="Нажми для просмотра полной истории"><b>'+esc(d.file_name)+'</b><div style="color:var(--u);margin-top:3px">';
      stages.forEach((s,i)=>{
        tl+='<span>'+esc(names[s.agent]||s.agent)+'</span><span style="color:#888">'+vBadges[s.verdict]+'</span>';
        if(i<stages.length-1) tl+=' → ';
      });
      tl+='</div><div style="color:#999;font-size:11px;margin-top:4px">'+esc(rel(d.created_at))+'</div></div>';
      return tl;
    }).join(''):'<div class="muted">Нет активных документов</div>';
    document.getElementById('documents').innerHTML=docHtml;
  }catch(e){ document.getElementById('documents').innerHTML='<div class="muted">Ошибка загрузки: '+esc(e)+'</div>'; }

  const ev=d.events||[];
  document.getElementById('events').innerHTML=ev.length?ev.map(e=>'<div class="event">'+esc(e.kind)+'<div class="muted">'+esc(e.ts)+' · '+esc(JSON.stringify(e.payload||{}))+'</div></div>').join(''):'<div class="muted">Событий нет</div>';
}

// Download popup functions are defined early in the script

// ─── Document management functions ───
let currentViewingDocId=null;

document.addEventListener('click',function(e){
  const btn=e.target&&e.target.closest?e.target.closest('button,a'):null;
  if(!btn) return;
  if(btn.getAttribute('onclick')) return;
  const txt=(btn.innerText||btn.textContent||'').trim();
  if(txt.includes('Скачать воркбук') || txt.includes('Скачать TXT')){
    e.preventDefault();
    downloadWorkbookTXT();
  }
});

window.downloadDoc = async function(ev, docId){
  if(ev) ev.stopPropagation();
  const safe=String(docId||'').replace(/[^A-Za-z0-9_-]/g,'');
  if(!safe) return;
  try{
    const resp=await fetch('/api/document/'+encodeURIComponent(safe)+'/download');
    if(!resp.ok) throw new Error('download failed');
    const blob=await resp.blob();
    const url=URL.createObjectURL(blob);
    const a=document.createElement('a');
    a.href=url;
    a.download='mila-document-'+safe+'.txt';
    document.body.appendChild(a);
    a.click();
    document.body.removeChild(a);
    URL.revokeObjectURL(url);
  }catch(e){
    const bubble=ev&&ev.target?ev.target.closest('.bubble,.msg,.message'):null;
    const text=bubble?bubble.innerText:('Документ '+safe);
    const blob=new Blob([text],{type:'text/plain; charset=utf-8'});
    const url=URL.createObjectURL(blob);
    const a=document.createElement('a');
    a.href=url;
    a.download='mila-document-'+safe+'.txt';
    document.body.appendChild(a);
    a.click();
    document.body.removeChild(a);
    URL.revokeObjectURL(url);
  }
};

// openDocModal is defined early in script

function closeDocModal(){
  document.getElementById('docModal').classList.remove('show');
  currentViewingDocId=null;
}

function toggleFeedbackBox(){
  const box=document.getElementById('feedbackBox');
  box.classList.toggle('show');
  if(box.classList.contains('show')){
    document.getElementById('feedbackText').focus();
  }
}

async function sendFeedback(){
  const feedback=document.getElementById('feedbackText').value.trim();
  if(!feedback) { alert('Введи текст правок'); return; }
  if(!currentViewingDocId) { alert('Документ не выбран'); return; }

  if(feedback.length > 10000) {
    alert('Текст правок слишком длинный (макс 10000 символов)');
    return;
  }

  try{
    // Определяем агентов для обратной связи
    const docResp=await fetch('/api/document/'+encodeURIComponent(currentViewingDocId));
    if(!docResp.ok) throw new Error('Документ не найден');

    const docData=await docResp.json();
    if(!docData.ok) throw new Error(docData.error||'Ошибка загрузки документа');

    const stages=docData.document.stages||[];
    if(stages.length<1) {
      alert('Нет этапов для отправки правок');
      return;
    }

    // Validation: последний stage должен быть валидный
    const lastStage=stages[stages.length-1];
    if(!lastStage.agent) {
      alert('Последний этап некорректен');
      return;
    }

    const fromAgent=lastStage.agent;
    const prevStage=stages.length>1?stages[stages.length-2]:null;
    const toAgent=prevStage?prevStage.agent:fromAgent;

    const r=await postJSON('/api/document/'+encodeURIComponent(currentViewingDocId)+'/feedback',{
      from_agent:fromAgent,
      to_agent:toAgent,
      feedback:feedback
    });
    if(!r.ok) throw new Error('HTTP '+r.status);

    const res=await r.json();
    if(res.ok){
      alert('Правки отправлены!');
      toggleFeedbackBox();
      document.getElementById('feedbackText').value='';
      await openDocModal(currentViewingDocId);
    }else {
      throw new Error(res.error||'Ошибка при отправке');
    }
  }catch(e){
    alert('Ошибка: '+e.message);
    console.error('sendFeedback error:', e);
  }
}

async function archiveDoc(){
  if(!currentViewingDocId) return;
  if(!confirm('Архивировать документ?')) return;

  try{
    const r=await postJSON('/api/document/'+encodeURIComponent(currentViewingDocId)+'/archive',{});
    if(!r.ok) throw new Error('Error archiving');
    const res=await r.json();
    if(res.ok){
      alert('Документ архивирован!');
      closeDocModal();
      load();
    }else throw new Error(res.error);
  }catch(e){
    alert('Ошибка: '+e.message);
  }
}

async function exportDoc(){
  if(!currentViewingDocId) return;
  try{
    const a=document.createElement('a');
    a.href='/api/document/'+encodeURIComponent(currentViewingDocId)+'/export';
    a.download='document_history.json';
    a.click();
  }catch(e){
    alert('Ошибка: '+e.message);
  }
}

// Клик вне модала закрывает его
document.addEventListener('click',function(e){
  const modal=document.getElementById('docModal');
  if(e.target===modal) closeDocModal();
});

load();
</script></body></html>"""


# ─── Общий статус-бар системы (Healthy / Degraded / Action needed) ──────────
# Один компонент на все страницы. Данные берёт из /api/health (level + reasons +
# attention). На странице чата (есть #side) рендерится компактной «пилюлей»
# справа сверху, на остальных — полосой в самом верху (flex-/обычный поток).
_STATUS_BAR = r"""
<style>
#mila-sb{font-family:Georgia,'Times New Roman',serif;font-size:13px;flex-shrink:0;display:flex;
  align-items:center;gap:10px;padding:6px 16px;color:#fff;cursor:pointer;user-select:none;
  border-bottom:1px solid rgba(0,0,0,.12)}
#mila-sb.healthy{background:#4A7A5E}#mila-sb.degraded{background:#B7791F}
#mila-sb.action_needed{background:#A8412C}#mila-sb.unknown{background:#7A5E54}
#mila-sb .dot{width:9px;height:9px;border-radius:50%;background:rgba(255,255,255,.92);flex-shrink:0}
#mila-sb .lbl{font-weight:bold;letter-spacing:.3px}
#mila-sb .rsn{opacity:.92;font-size:12px}
#mila-sb .sp{flex:1}#mila-sb .chev{opacity:.85;font-size:11px}
#mila-sb.pill{position:fixed;top:18px;right:18px;z-index:60;border-radius:20px;padding:6px 13px;
  border:0;box-shadow:0 2px 8px rgba(0,0,0,.18);max-width:min(320px,calc(100vw - 36px))}
body:has(#side) header{padding-right:360px}
#mila-sb.pill .rsn,#mila-sb.pill .sp{display:none}
#mila-sb-panel{position:fixed;z-index:61;background:#fff;color:#1E140F;border:1px solid #E0D0C8;
  border-radius:12px;box-shadow:0 8px 28px rgba(0,0,0,.18);padding:12px 14px;font-size:13px;
  min-width:250px;display:none}
#mila-sb-panel.bar{left:16px;top:38px}#mila-sb-panel.pill{right:14px;top:46px}
#mila-sb-panel h4{margin:0 0 8px;font-size:11px;color:#7A5E54;text-transform:uppercase;letter-spacing:.5px}
#mila-sb-panel .svc{display:flex;justify-content:space-between;gap:14px;padding:4px 0;border-bottom:1px solid #F2EAE2}
#mila-sb-panel .svc:last-child{border:0}
#mila-sb-panel .up{color:#4A7A5E}#mila-sb-panel .down{color:#A8412C}
#mila-sb-panel .lk{margin-top:10px;display:flex;gap:12px;flex-wrap:wrap}
#mila-sb-panel .lk a{color:#C4614A;text-decoration:none;font-size:12px}
@media(max-width:900px){
  body:has(#side) header{padding-right:18px;padding-bottom:52px}
  #mila-sb.pill{top:76px;right:14px}
  #mila-sb-panel.pill{right:14px;top:122px}
}
</style>
<script>
(function(){
  var LBL={healthy:'Система в норме',degraded:'Сниженный режим',action_needed:'Нужно вмешательство',unknown:'Статус неизвестен'};
  function row(name,ok,extra){return '<div class="svc"><span>'+name+'</span><span class="'+(ok?'up':'down')+'">'+(ok?('✓ '+(extra||'ок')):('✕ '+(extra||'нет')))+'</span></div>';}
  function panelHtml(h){
    var a=h.attention||{}, g=h.gemini||{}, c=h.claude||{}, ig=h.instagram||{}, sb=h.supabase||{};
    var llmOk=g.configured||c.configured, llmName=g.configured?'Gemini':(c.configured?'Claude':'нет');
    return '<h4>Состояние системы</h4>'
      +row('LLM',llmOk,llmName)
      +row('Telegram',(h.telegram||{}).configured)
      +row('Instagram',ig.configured,ig.flow)
      +row('Supabase',sb.configured,sb.can_write?'запись':'')
      +row('n8n',(h.n8n||{}).up)
      +row('n8n-мост',(h.bridge||{}).up)
      +'<div class="svc"><span>Очередь</span><span>'+(a.failed||0)+' ошибок · '+(a.awaiting_approval||0)+' на одобрении · '+(a.running||0)+' в работе</span></div>'
      +'<div class="lk"><a href="/dashboard">Дашборд</a><a href="/operator">Очередь</a><a href="/settings">Настройки</a></div>';
  }
  var mode=document.getElementById('side')?'pill':'bar';
  var bar=document.createElement('div');bar.id='mila-sb';bar.className=(mode==='pill'?'pill ':'')+'unknown';
  bar.innerHTML='<span class="dot"></span><span class="lbl">…</span><span class="rsn"></span><span class="sp"></span><span class="chev">▾</span>';
  var panel=document.createElement('div');panel.id='mila-sb-panel';panel.className=mode;
  if(mode==='pill'){document.body.appendChild(bar);}else{document.body.insertBefore(bar,document.body.firstChild);}
  document.body.appendChild(panel);
  bar.addEventListener('click',function(e){e.stopPropagation();panel.style.display=panel.style.display==='block'?'none':'block';});
  document.addEventListener('click',function(){panel.style.display='none';});
  async function refresh(){
    try{
      var h=await (await fetch('/api/health')).json();
      var lvl=h.level||'unknown';
      bar.className=(mode==='pill'?'pill ':'')+lvl;
      bar.querySelector('.lbl').textContent=LBL[lvl]||lvl;
      bar.querySelector('.rsn').textContent=(h.reasons&&h.reasons.length)?('· '+h.reasons.join(' · ')):'';
      panel.innerHTML=panelHtml(h);
    }catch(e){
      bar.className=(mode==='pill'?'pill ':'')+'unknown';
      bar.querySelector('.lbl').textContent='Статус недоступен';
    }
  }
  refresh();setInterval(refresh,30000);
})();
</script>
"""

INDEX_HTML = INDEX_HTML.replace("</body>", _STATUS_BAR + "</body>")
DASHBOARD_HTML = DASHBOARD_HTML.replace("</body>", _STATUS_BAR + "</body>")
OPERATOR_HTML = OPERATOR_HTML.replace("</body>", _STATUS_BAR + "</body>")
SETTINGS_HTML = SETTINGS_HTML.replace("</body>", _STATUS_BAR + "</body>")


def _open_browser():
    import time
    time.sleep(1.3)
    scheme = "https" if _HTTPS else "http"
    try:
        webbrowser.open(f"{scheme}://localhost:5000")
    except Exception:
        pass


# ─── Глобальные error handlers с Telegram alerts ─────────────────────────────────────────

@app.errorhandler(500)
def handle_500(error):
    """Обработка 500 ошибок: логируем, отправляем alert, возвращаем безопасный ответ."""
    context = {
        "source": "webapp",
        "endpoint": request.path,
        "method": request.method,
        "user_agent": request.headers.get("User-Agent", "unknown")[:50]
    }
    error_monitor.log_error(error, context=context, alert=True, level="CRITICAL")
    logger.error(f"500 Error: {error}", exc_info=True)
    return jsonify({
        "ok": False,
        "error": "Внутренняя ошибка сервера. Техническая команда уведомлена."
    }), 500


@app.errorhandler(404)
def handle_404(error):
    """404 — не alert, просто логируем."""
    logger.warning(f"404 Not Found: {request.path}")
    return jsonify({"ok": False, "error": "Not found"}), 404


@app.errorhandler(Exception)
def handle_generic_error(error):
    """Перехват всех необработанных исключений."""
    context = {
        "source": "webapp",
        "endpoint": request.path,
        "method": request.method,
        "error_type": type(error).__name__
    }
    error_monitor.log_error(error, context=context, alert=isinstance(error, (RuntimeError, ValueError)), level="ERROR")
    logger.error(f"Unhandled error: {error}", exc_info=True)
    return jsonify({
        "ok": False,
        "error": "Ошибка при обработке запроса"
    }), 500


if __name__ == "__main__":
    scheme = "https" if _HTTPS else "http"
    logger.info("MILA OFFICE web -> %s://localhost:5000  (Ctrl+C to stop)", scheme)
    if _HTTPS:
        logger.info("HTTPS включён (self-signed). Браузер один раз предупредит о сертификате.")
    logger.info("Логи: %s", _LOG_FILE)
    threading.Thread(target=_open_browser, daemon=True).start()
    # debug=False: traceback идёт в лог-файл, не в браузер. threaded=True —
    # чтобы /api/result опрашивался, пока фоновая задача ещё считается.
    # ssl_context='adhoc' (нужен пакет cryptography) — самоподписанный сертификат
    # для localhost: позволяет валидный https-redirect для Instagram OAuth.
    app.run(host="127.0.0.1", port=5000, debug=False, threaded=True,
            ssl_context="adhoc" if _HTTPS else None)


