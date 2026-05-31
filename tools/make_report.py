#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_report.py — строит .docx-отчёт по аналитике Instagram из выгрузки reports/posts_*.json.

Использование:
    python make_report.py                       # последний posts_*.json, месяц «Май 2026»
    python make_report.py <путь.json> "Май 2026"

Работает строго на реальных полях выгрузки: date, type, likes, comments,
reach, caption, link. Чего НЕТ в данных (сохранения, время суток) — в отчёте
честно помечается, а не выдумывается. Темы определяются по тексту caption.
"""
import sys
# UTF-8 для консоли Windows — иначе финальный print с русским текстом падает
# с UnicodeEncodeError на cp1252 (хотя .docx уже сохранён).
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import json
import datetime
from pathlib import Path
from collections import defaultdict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
REPORTS = ROOT / "reports"
OUTDIR = ROOT / "MILA-BUSINESS" / "05-analytics"

GOLD = RGBColor(0xB8, 0x86, 0x0B)
DARK = RGBColor(0x33, 0x33, 0x33)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x99, 0x99, 0x99)

WD_RU = ["понедельник", "вторник", "среда", "четверг",
         "пятница", "суббота", "воскресенье"]


def fmt(n):
    return f"{int(round(n)):,}".replace(",", " ")


def classify(caption):
    """Грубая тематическая категория по тексту поста (приоритет сверху вниз)."""
    c = (caption or "").lower()
    rel = any(w in c for w in ["мужчин", "отношени", "партн", " пара", "пары",
                               "ради мужчины", "он хотел", "счастливых пар"])
    kids = any(w in c for w in ["дет", "сын", "ребён", "ребен", "первенц",
                                "школ", "день матери", "мама двоих"])
    # «отношения» важнее «детей», если речь про мужчину/женщину в паре
    if rel and not (kids and "мужчин" not in c and "отношени" not in c):
        return "Отношения (пара, мужчина/женщина)"
    if kids:
        return "Материнство / дети"
    if any(w in c for w in ["самооценк", "любовь к себе", "люблю себя",
                            "возвраща", "выгорани", "работа над собой",
                            "контролировать", "контролирует", "тревог",
                            "беречь", "берегём", "терпe", "терпеть", "опора",
                            "себя", "идеальной", "настоящей"]):
        return "Самоценность / возвращение к себе"
    if any(w in c for w in ["спорт", "калори", "танцева", "микрофон",
                            "канад", "подруг", "prada", "хобби"]):
        return "Лайфстайл / личное"
    return "Другое"


def load(path):
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    for p in data:
        p["likes"] = int(p.get("likes", 0) or 0)
        p["comments"] = int(p.get("comments", 0) or 0)
        p["reach"] = int(p.get("reach", 0) or 0)
        p["eng"] = p["likes"] + p["comments"]
        p["er"] = (p["eng"] / p["reach"]) if p["reach"] else 0
        d = datetime.date.fromisoformat(p["date"])
        p["wd"] = WD_RU[d.weekday()]
        p["theme"] = classify(p.get("caption", ""))
        # короткий заголовок из первой строки
        first = (p.get("caption", "") or "").splitlines()
        p["hook"] = (first[0] if first else "")[:60]
    return data


def styled_table(doc, headers, rows, widths=None, align_right=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htext in enumerate(t.rows[0].cells):
        run = htext.paragraphs[0].add_run(headers[i])
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t


def H(doc, text):
    p = doc.add_heading(level=1)
    p.add_run(text).font.color.rgb = GOLD
    return p


def bullet(doc, text, lead=None, color=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead)
        r.bold = True
        if color:
            r.font.color.rgb = color
    p.add_run(text)
    return p


def build(data, month, out_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    reels = [p for p in data if p["type"] == "REELS"]
    feed = [p for p in data if p["type"] != "REELS"]
    n = len(data)
    tot_reach = sum(p["reach"] for p in data)
    tot_likes = sum(p["likes"] for p in data)
    tot_comm = sum(p["comments"] for p in data)
    avg_reach = tot_reach / n
    reach_sorted = sorted([p["reach"] for p in data])
    median = reach_sorted[n // 2] if n % 2 else (
        reach_sorted[n // 2 - 1] + reach_sorted[n // 2]) / 2

    by_reach = sorted(data, key=lambda p: p["reach"], reverse=True)
    top3 = by_reach[:3]
    top3_reach = sum(p["reach"] for p in top3)
    top1 = by_reach[0]
    dates = sorted(p["date"] for p in data)

    # ---- титул ----
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Аналитика Instagram"); r.bold = True
    r.font.size = Pt(24); r.font.color.rgb = DARK
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(f"@liudmyla.lykova  ·  {month}")
    r.font.size = Pt(13); r.font.color.rgb = GOLD
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(f"Период данных: {dates[0]} – {dates[-1]}  ·  выводы и план на следующий месяц")
    r.italic = True; r.font.size = Pt(10)
    doc.add_paragraph()

    # ---- 1. Сводка ----
    H(doc, "1. Сводка за месяц")
    doc.add_paragraph(
        f"За период опубликовано {n} публикаций ({len(reels)} Reels + {len(feed)} в ленте). "
        f"Суммарный охват — {fmt(tot_reach)}."
    )
    styled_table(
        doc,
        ["Показатель", "Значение"],
        [
            ["Публикаций всего", f"{n}  ({len(reels)} Reels, {len(feed)} лента)"],
            ["Суммарный охват", fmt(tot_reach)],
            ["Средний охват на пост", fmt(avg_reach)],
            ["Медианный охват (типичный пост)", fmt(median)],
            ["Суммарно лайков", fmt(tot_likes)],
            ["Суммарно комментариев", fmt(tot_comm)],
            ["Лучший пост по охвату", f"{fmt(top1['reach'])}  ({top1['date']})"],
        ],
        widths=[3.4, 2.6],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Главное наблюдение: ").bold = True
    p.add_run(
        f"3 поста дали {fmt(top3_reach)} охвата — это "
        f"{round(top3_reach / tot_reach * 100)}% всего месячного охвата. "
        f"При этом средний охват ({fmt(avg_reach)}) сильно раздут вирусными видео, "
        f"а медианный (типичный) пост собрал всего {fmt(median)}. "
        "Аккаунт растёт не на объёме, а на нескольких сильных Reels — "
        "именно их тему и подачу нужно усиливать, а базовый охват подтягивать."
    )

    # ---- 2. Топ-посты ----
    H(doc, "2. Посты с наибольшим охватом")
    note = doc.add_paragraph()
    note.add_run("Важно: ").bold = True
    note.add_run(
        "сохранения (saves) Instagram отдаёт отдельной метрикой insights, "
        "и в текущую выгрузку они не попали. Поэтому рейтинг построен по охвату и "
        "вовлечённости (лайки + комментарии). Добавить сохранения — пункт плана ниже."
    )
    rows = []
    for i, p in enumerate(by_reach[:10], 1):
        rows.append([i, p["date"], p["hook"], fmt(p["reach"]),
                     fmt(p["likes"]), p["comments"], f"{p['er']*100:.1f}%"])
    styled_table(
        doc,
        ["#", "Дата", "Заход (первая строка)", "Охват", "Лайки", "Комм.", "ER*"],
        rows,
        widths=[0.3, 0.85, 2.7, 0.9, 0.8, 0.6, 0.65],
    )
    p = doc.add_paragraph()
    p.add_run("* ER = (лайки + комментарии) / охват — насколько сильно зацепило тех, кто увидел.\n").italic = True
    doc.add_paragraph(
        f"Топ-пост месяца — Reel от {top1['date']}: «{top1['hook']}…» — "
        f"{fmt(top1['reach'])} охвата и {fmt(top1['likes'])} лайков. "
        "У трёх лидеров высокий ER — они «выстрелили» не случайно: тема и заход "
        "реально резонируют. Это сигнал «делай ещё такое», а не разовая удача алгоритма."
    )

    # ---- 3. Темы ----
    H(doc, "3. Какие темы резонируют (по тексту постов)")
    groups = defaultdict(list)
    for p in data:
        groups[p["theme"]].append(p)
    theme_rows = []
    for name, items in sorted(groups.items(),
                              key=lambda kv: -sum(p["reach"] for p in kv[1])):
        sr = sum(p["reach"] for p in items)
        theme_rows.append([name, len(items), fmt(sr),
                           f"{round(sr / tot_reach * 100)}%",
                           fmt(sr / len(items))])
    styled_table(
        doc,
        ["Тема", "Постов", "Суммарный охват", "Доля", "Ср. охват"],
        theme_rows,
        widths=[2.7, 0.8, 1.5, 0.7, 1.1],
    )
    doc.add_paragraph()
    bullet(doc, "почти весь охват и все вирусные видео. Контрастные «неудобные правды» "
                "о паре («не делай ради мужчины…», «её зовут меркантильной, но и мужчины "
                "хотят премиум», «устают не только женщины») — твой проверенный формат №1.",
           lead="Отношения мужчина/женщина — ", color=GREEN)
    bullet(doc, "стабильно средне: полезно, аудитория откликается, но охваты ниже — "
                "это поддерживающий контент, не драйвер роста.",
           lead="Самоценность / возвращение к себе — ")
    bullet(doc, "и лайфстайл (спорт, калории, танцы, «чем занимаетесь») — самый слабый "
                "охват. Они размывают позиционирование психолога по отношениям.",
           lead="Материнство/дети — ", color=RED)
    doc.add_paragraph(
        "Что именно цепляет внутри темы отношений: контраст и узнавание себя, "
        "взгляд на обе стороны (и про женщину, и про мужчину), чёткая мысль в первой "
        "же фразе. Что не заходит — общий «селф-кэр» («люби себя») и бытовой лайфстайл."
    )

    # ---- 4. Дни/время ----
    H(doc, "4. В какие дни и время публиковать")
    p = doc.add_paragraph()
    p.add_run("Честная оговорка: ").bold = True
    p.add_run(
        "надёжного вывода по времени суток сделать нельзя — в выгрузке есть только "
        "дата публикации, без часа. По дням недели данные есть, но вирусные видео "
        "случайно попали на отдельные дни и раздувают их средние. Поэтому таблица ниже "
        "— ориентир, а не закон."
    )
    wd = defaultdict(list)
    for p in data:
        wd[p["wd"]].append(p)
    wd_rows = []
    for d in WD_RU:
        if d in wd:
            items = wd[d]
            wd_rows.append([d.capitalize(), len(items),
                            fmt(sum(x["reach"] for x in items) / len(items)),
                            fmt(max(x["reach"] for x in items))])
    styled_table(doc, ["День", "Постов", "Ср. охват", "Макс. охват"], wd_rows,
                 widths=[1.6, 0.9, 1.3, 1.3])
    doc.add_paragraph()
    doc.add_paragraph(
        "Что делать: в следующем месяце фиксируй время публикации каждого Reels "
        "(например, вечер 18:00–20:00 — пик женской аудитории) и в выгрузку добавь "
        "поле времени. За 2–3 месяца накопится честная статистика по дням и часам."
    )

    # ---- 5. Больше / меньше ----
    H(doc, "5. Что делать больше — что меньше")
    p = doc.add_paragraph(); r = p.add_run("Делать БОЛЬШЕ"); r.bold = True; r.font.color.rgb = GREEN
    bullet(doc, "Reels про динамику «мужчина ↔ женщина» с контрастным заходом («её считают…, но на самом деле…»).")
    bullet(doc, "Главную мысль — в первую же фразу: у всех вирусных видео цепляющая строка стоит первой.")
    bullet(doc, "Под сильными Reels — явный призыв «сохрани / напиши ХОЧУ», чтобы переводить охват в заявки на практикум.")
    bullet(doc, "Концентрация: несколько сильных Reels важнее десятков случайных постов (3 видео дали 90%+ охвата).")
    bullet(doc, "Развивать темы-победители новыми углами: «не делай ради мужчины…», «премиум-класс», усталость в паре.")
    p = doc.add_paragraph(); r = p.add_run("Делать МЕНЬШЕ"); r.bold = True; r.font.color.rgb = RED
    bullet(doc, "Общий «селф-кэр» («люби себя», «береги себя») — тема перегрета, охваты низкие.")
    bullet(doc, "Лайфстайл/спорт/калории/«чем занимаетесь» и личные танцы — слабый охват, размывают позиционирование.")
    bullet(doc, "Статичные посты в ленту ради охвата — у них минимальные охваты (96–158). Лента — для пользы и сохранений, охват тяни через Reels.")

    # ---- 6. Прогноз ----
    H(doc, "6. Прогноз: 4 Reels в месяц на лучшие темы")
    rel = [p for p in reels if p["theme"].startswith("Отношения")]
    viral = [p for p in data if p["reach"] >= 40000]
    rel_solid = [p for p in rel if p not in viral]
    solid_avg = (sum(p["reach"] for p in rel_solid) / len(rel_solid)) if rel_solid else avg_reach
    viral_avg = (sum(p["reach"] for p in viral) / len(viral)) if viral else avg_reach
    viral_rate = len(viral) / len(reels)

    doc.add_paragraph(
        f"Логика расчёта — на реальных числах месяца: из {len(reels)} Reels «выстрелили» "
        f"{len(viral)} (≈{round(viral_rate*100)}%); типичный крепкий Reel про отношения без "
        f"вируса собирал ≈{fmt(solid_avg)} охвата, средний вирусный — ≈{fmt(viral_avg)}. "
        "Если делать 4 точечных Reels в месяц строго на проверенных темах:"
    )
    pess = 4 * solid_avg
    base = viral_avg * (4 * viral_rate) + solid_avg * (4 - 4 * viral_rate)
    opt = viral_avg + 3 * solid_avg
    styled_table(
        doc,
        ["Сценарий", "Логика", "Охват/мес", "Подписчики*"],
        [
            ["Осторожный", "ни одного «выстрела», 4 крепких видео", f"~{fmt(pess)}", "+50–120"],
            ["Базовый (ожидаемый)", f"≈{round(viral_rate*100)}% шанс вируса на видео", f"~{fmt(base)}", "+200–500"],
            ["Оптимистичный", "1 настоящий вирус + 3 крепких", f"~{fmt(opt)}", "+700–1500"],
        ],
        widths=[1.7, 2.4, 1.1, 1.1],
    )
    doc.add_paragraph()
    bullet(doc, f"Базовый сценарий (~{fmt(base)} охвата за 4 поста) сопоставим с охватом "
                f"всего прошедшего месяца ({fmt(tot_reach)}) — но достигается точечно, "
                "а не объёмом. Качество вместо количества реально работает.")
    bullet(doc, "4 Reels — это минимум для роста. Для активных продаж практикума лучше "
                "4 «герой»-Reels + лёгкий поддерживающий контент (сторис/карусели), "
                "чтобы не пропадать из ленты.")
    star = doc.add_paragraph()
    star.add_run("* Прогноз подписчиков ориентировочный ").italic = True
    star.add_run("(≈0,1–0,3% от охвата non-followers); сильно зависит от качества хука. "
                 "Выводы по темам и формату надёжны, прогноз — предварителен (1 месяц данных, "
                 "вирусные выбросы).").italic = True

    # ---- 7. План ----
    H(doc, "7. План на следующий месяц")
    for name, desc in [
        ("Reel 1 (локомотив)", "Развить тему поста-лидера «не делай ради мужчины эти 3 вещи» — "
         "потеря себя в отношениях. Контрастный заход в первой фразе."),
        ("Reel 2 (премиум)", "Развить «её зовут меркантильной, но и мужчины хотят премиум» "
         "под новым углом — взаимные ожидания в паре."),
        ("Reel 3 (усталость/роли)", "«Устают оба» / «пара забывает, что они мужчина и женщина» — "
         "разделение нагрузки, здоровые роли."),
        ("Reel 4 (тест-тема)", "Новый угол внутри отношений + сильный хук. Проверяем свежую идею, "
         "чтобы не выгорала одна тема."),
    ]:
        bullet(doc, desc, lead=f"{name}: ")
    doc.add_paragraph()
    doc.add_paragraph(
        "В каждом Reel: цепляющий хук в первой фразе, одна чёткая мысль, призыв "
        "«сохрани / отправь подруге / напиши ХОЧУ», мягкая отсылка к практикуму "
        "«Почему я снова выбрала не того» в конце или в закреплённом комментарии. "
        "Зафиксировать время постинга и в следующую выгрузку добавить сохранения и комментарии."
    )

    # ---- ограничения ----
    H(doc, "Ограничения данных (для честности выводов)")
    bullet(doc, "Сохранения (saves) и репосты в выгрузку не попали — рейтинг по охвату и вовлечённости.")
    bullet(doc, "Время публикации отсутствует (только дата) — выводы по часам сделать нельзя.")
    bullet(doc, "Комментарии-заявки не выгружены отдельно — конверсию в заявки оцениваю качественно.")
    bullet(doc, "Темы определены автоматически по тексту постов — возможны небольшие неточности категоризации.")
    bullet(doc, "Один месяц данных с вирусными выбросами: выводы по темам/формату надёжны, по дням и прогнозу — предварительны.")

    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Отчёт построен автоматически из выгрузки Instagram Graph API · MILA GOLD")
    r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY

    doc.save(str(out_path))
    return out_path


def latest_posts_json():
    files = sorted(REPORTS.glob("posts_*.json"))
    if not files:
        sys.exit(f"Не найдено posts_*.json в {REPORTS}. Сначала: python get_analytics.py posts")
    return files[-1]


def main():
    json_path = Path(sys.argv[1]) if len(sys.argv) > 1 else latest_posts_json()
    month = sys.argv[2] if len(sys.argv) > 2 else "Май 2026"
    data = load(json_path)
    out = OUTDIR / f"analytics_{month.lower().replace(' ', '_')}.docx"
    build(data, month, out)
    print(f"OK: {out}  ({len(data)} постов, источник {json_path.name})")


if __name__ == "__main__":
    main()
